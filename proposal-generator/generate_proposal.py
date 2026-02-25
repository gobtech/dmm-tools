#!/usr/bin/env python3
"""
Proposal Generator
==================
Generates client proposal .docx documents from artist/genre/timeline inputs.
Pre-fills country-specific press targets, radio targets, DSP pitching strategies,
and budget breakdown based on genre and selected services.

Called from the web UI via /api/proposal/generate.
"""

import csv
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

ROOT_DIR = Path(__file__).parent.parent.resolve()
sys.path.insert(0, str(ROOT_DIR))

REPORT_DIR = ROOT_DIR / 'reports'
REPORT_DIR.mkdir(exist_ok=True)

DATA_DIR = ROOT_DIR / 'data'

# ── Country display ordering ────────────────────────────────────────────

COUNTRY_ORDER = [
    'MÉXICO', 'BRAZIL', 'ARGENTINA', 'CHILE', 'COLOMBIA',
    'ECUADOR', 'PERU', 'URUGUAY', 'VENEZUELA', 'LATAM',
]

COUNTRY_DISPLAY = {
    'MÉXICO': 'Mexico',
    'BRAZIL': 'Brazil',
    'ARGENTINA': 'Argentina',
    'CHILE': 'Chile',
    'COLOMBIA': 'Colombia',
    'ECUADOR': 'Ecuador',
    'PERU': 'Peru',
    'URUGUAY': 'Uruguay',
    'VENEZUELA': 'Venezuela',
    'LATAM': 'LATAM',
}

# ── Genre mapping for filtering ─────────────────────────────────────────

GENRE_KEYWORDS = {
    'electronic': ['electronic', 'dance', 'edm', 'techno', 'house', 'electronica', 'dj'],
    'indie': ['indie', 'alternative', 'lo-fi', 'dream pop', 'shoegaze'],
    'rock': ['rock', 'punk', 'metal', 'post-punk', 'garage', 'grunge'],
    'pop': ['pop', 'synth-pop', 'electropop', 'pop rock'],
    'urban': ['reggaeton', 'urbano', 'trap', 'hip-hop', 'rap', 'r&b', 'urban'],
    'general': [],
}


# ═════════════════════════════════════════════════════════════════════════
#  DATA LOADERS
# ═════════════════════════════════════════════════════════════════════════

def _load_press_targets(genre='general'):
    """Load press outlets, optionally filtered by genre."""
    path = DATA_DIR / 'press_database.csv'
    if not path.exists():
        return {}

    keywords = GENRE_KEYWORDS.get(genre, [])
    outlets_by_country = {}

    with open(path, encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        for row in reader:
            name = row.get('NAME OF MEDIA', '').strip()
            territory = row.get('Territory', '').strip().upper()
            if not name or not territory:
                continue
            # Normalize territory names
            if territory in ('PENDING', 'CANCELLED', ''):
                continue

            genre_tags = row.get('Genre', '').strip().lower()
            description = row.get('DESCRIPTION & SM', '').strip()
            website = row.get('WEBSITE', '').strip()
            reach = row.get('REACH', '').strip()

            # Genre filtering: include if outlet has matching genre tag,
            # or if genre is 'general', or if outlet has no genre tag (include all)
            genre_match = False
            if genre == 'general' or not keywords:
                genre_match = True
            elif not genre_tags:
                genre_match = False  # skip untagged when filtering by genre
            else:
                for kw in keywords:
                    if kw in genre_tags:
                        genre_match = True
                        break

            if not genre_match:
                continue

            outlets_by_country.setdefault(territory, []).append({
                'name': name,
                'description': description,
                'website': website,
                'reach': reach,
                'genre_tags': genre_tags,
            })

    return outlets_by_country


def _load_radio_targets(genre='general', countries=None):
    """Load radio station targets filtered by genre and country."""
    path = DATA_DIR / 'radio_targets.csv'
    if not path.exists():
        return {}

    keywords = GENRE_KEYWORDS.get(genre, [])
    stations_by_country = {}

    with open(path, encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        for row in reader:
            station = row.get('Station', '').strip()
            country = row.get('Country', '').strip()
            station_genre = row.get('Genre', '').strip().lower()
            fmt = row.get('Format', '').strip()
            price = row.get('Price USD', '').strip()
            duration = row.get('Duration', '').strip()
            notes = row.get('Notes', '').strip()

            if not station or not country:
                continue

            if countries and country.upper() not in [c.upper() for c in countries]:
                continue

            # Genre filtering
            genre_match = False
            if genre == 'general' or not keywords:
                genre_match = True
            else:
                for kw in keywords:
                    if kw in station_genre:
                        genre_match = True
                        break

            if not genre_match:
                continue

            stations_by_country.setdefault(country, []).append({
                'station': station,
                'genre': row.get('Genre', '').strip(),
                'format': fmt,
                'price': price,
                'duration': duration,
                'notes': notes,
            })

    return stations_by_country


def _load_dsp_strategies(genre='general'):
    """Load DSP pitching strategies for the given genre."""
    path = DATA_DIR / 'dsp_strategy.json'
    if not path.exists():
        return {}

    with open(path, encoding='utf-8') as f:
        data = json.load(f)

    # Map genre to strategy key
    genre_key = genre
    if genre_key not in ('electronic', 'indie', 'rock', 'pop', 'urban', 'general'):
        genre_key = 'general'

    platforms = data.get('platforms', {})
    result = {}
    for name, info in platforms.items():
        playlists = info.get('genre_playlists', {}).get(genre_key, [])
        if not playlists:
            playlists = info.get('genre_playlists', {}).get('general', [])
        result[name] = {
            'description': info.get('description', ''),
            'strategy': info.get('strategy', ''),
            'pitch_timing': info.get('pitch_timing', ''),
            'playlists': playlists,
        }

    return result


def _load_pricing():
    """Load pricing data."""
    path = DATA_DIR / 'pricing.json'
    if not path.exists():
        return {}

    with open(path, encoding='utf-8') as f:
        return json.load(f)


# ═════════════════════════════════════════════════════════════════════════
#  MAIN FUNCTION
# ═════════════════════════════════════════════════════════════════════════

def generate_proposal(
    artist,
    genre='general',
    timeline=None,
    collaborators='',
    campaign_duration=3,
    goal_strategy='',
    digital_marketing='',
    countries=None,
    radio_stations=None,
    influencer_tier='mid',
    dj_markets=None,
    digital_package='standard',
    output_path=None,
    log_fn=None,
):
    """
    Generate a client proposal .docx.

    Parameters
    ----------
    artist : str
        Artist name.
    genre : str
        Genre key: electronic, indie, rock, pop, urban, general.
    timeline : list of dict, optional
        Release timeline entries: [{'title': ..., 'date': ..., 'format': ...}, ...]
    collaborators : str
        Key collaborators or hooks (free text).
    campaign_duration : int
        Campaign duration in months (default 3).
    goal_strategy : str
        Custom goal/strategy text. If empty, a template is generated.
    digital_marketing : str
        Custom digital marketing text. If empty, a template is generated.
    countries : list of str, optional
        Target countries. Default: all LATAM.
    radio_stations : list of str, optional
        Selected radio station names. If None, all genre-relevant stations.
    influencer_tier : str
        'micro', 'mid', 'macro', or None to exclude.
    dj_markets : list of str, optional
        DJ/club servicing market selections. None to exclude.
    digital_package : str
        'starter', 'standard', 'premium', or None to exclude.
    output_path : str, optional
        Output .docx path.
    log_fn : callable, optional
        Logging callback.

    Returns
    -------
    dict with output_path and metadata.
    """
    if log_fn is None:
        log_fn = print

    if timeline is None:
        timeline = []
    if countries is None:
        countries = list(COUNTRY_DISPLAY.keys())

    safe_artist = re.sub(r'[^\w\-]', '_', artist.lower())
    if not output_path:
        output_path = str(REPORT_DIR / f'{safe_artist}_proposal.docx')

    log_fn(f'Generating proposal for {artist} ({genre})...')

    # ── Load data ──────────────────────────────────────────────
    log_fn('Loading press targets...')
    press_targets = _load_press_targets(genre)

    log_fn('Loading radio targets...')
    radio_targets = _load_radio_targets(genre, countries)

    log_fn('Loading DSP strategies...')
    dsp_strategies = _load_dsp_strategies(genre)

    log_fn('Loading pricing...')
    pricing = _load_pricing()

    # ── Count data for summary ─────────────────────────────────
    total_press = sum(len(v) for v in press_targets.values())
    total_radio = sum(len(v) for v in radio_targets.values())
    log_fn(f'  Press targets: {total_press} outlets across {len(press_targets)} countries')
    log_fn(f'  Radio targets: {total_radio} stations across {len(radio_targets)} countries')
    log_fn(f'  DSP platforms: {len(dsp_strategies)}')

    # ── Generate .docx ─────────────────────────────────────────
    log_fn('Building proposal document...')
    _build_proposal_docx(
        artist=artist,
        genre=genre,
        timeline=timeline,
        collaborators=collaborators,
        campaign_duration=campaign_duration,
        goal_strategy=goal_strategy,
        digital_marketing=digital_marketing,
        countries=countries,
        press_targets=press_targets,
        radio_targets=radio_targets,
        radio_stations=radio_stations,
        dsp_strategies=dsp_strategies,
        pricing=pricing,
        influencer_tier=influencer_tier,
        dj_markets=dj_markets,
        digital_package=digital_package,
        output_path=output_path,
    )

    log_fn(f'Proposal saved → {output_path}')
    return {
        'output_path': output_path,
        'press_count': total_press,
        'radio_count': total_radio,
        'dsp_platforms': len(dsp_strategies),
    }


# ═════════════════════════════════════════════════════════════════════════
#  DOCX BUILDER
# ═════════════════════════════════════════════════════════════════════════

def _build_proposal_docx(
    artist, genre, timeline, collaborators, campaign_duration,
    goal_strategy, digital_marketing, countries,
    press_targets, radio_targets, radio_stations, dsp_strategies,
    pricing, influencer_tier, dj_markets, digital_package,
    output_path,
):
    """Build the proposal .docx file."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Default styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    RED = RGBColor(0xC4, 0x30, 0x30)
    BLUE = RGBColor(0x00, 0x56, 0xD2)
    GRAY = RGBColor(0x66, 0x66, 0x66)
    BLACK = RGBColor(0x22, 0x22, 0x22)

    GENRE_LABELS = {
        'electronic': 'Electronic / Dance',
        'indie': 'Indie / Alternative',
        'rock': 'Rock',
        'pop': 'Pop',
        'urban': 'Urban / Reggaeton',
        'general': 'General',
    }
    genre_label = GENRE_LABELS.get(genre, genre.title())

    # ─── Title Page ────────────────────────────────────────────
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(artist.upper())
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(28)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run('LATAM Campaign Proposal')
    run.font.size = Pt(16)
    run.font.color.rgb = GRAY
    sub_para.paragraph_format.space_after = Pt(8)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime('%B %Y'))
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY

    # Genre line
    genre_para = doc.add_paragraph()
    genre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = genre_para.add_run(genre_label)
    run.font.size = Pt(11)
    run.font.color.rgb = RED
    run.italic = True
    genre_para.paragraph_format.space_after = Pt(24)

    # DMM attribution
    attr_para = doc.add_paragraph()
    attr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = attr_para.add_run('Prepared by Dorado Music Marketing')
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    # Page break after title
    doc.add_page_break()

    # ─── 1. Campaign Overview ─────────────────────────────────
    _section_header(doc, '1. Campaign Overview', RED)

    _label_value(doc, 'Artist', artist, GRAY)
    _label_value(doc, 'Genre', genre_label, GRAY)
    _label_value(doc, 'Campaign Duration', f'{campaign_duration} months', GRAY)

    if collaborators:
        _label_value(doc, 'Key Collaborators', collaborators, GRAY)

    target_countries_display = ', '.join(
        COUNTRY_DISPLAY.get(c, c) for c in countries
        if c in COUNTRY_DISPLAY
    )
    _label_value(doc, 'Target Markets', target_countries_display, GRAY)

    # ─── 2. Goal & Strategy ───────────────────────────────────
    _section_header(doc, '2. Goal & Strategy', RED)

    if goal_strategy:
        p = doc.add_paragraph()
        run = p.add_run(goal_strategy)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(8)
    else:
        _default_goal_strategy(doc, artist, genre_label, campaign_duration, collaborators)

    # ─── 3. Release Timeline ──────────────────────────────────
    if timeline:
        _section_header(doc, '3. Release Timeline', RED)
        for i, release in enumerate(timeline):
            p = doc.add_paragraph()
            title_run = p.add_run(release.get('title', f'Release {i+1}'))
            title_run.bold = True
            title_run.font.size = Pt(10)

            parts = []
            if release.get('date'):
                parts.append(release['date'])
            if release.get('format'):
                parts.append(release['format'])
            if parts:
                detail_run = p.add_run(f"  —  {' · '.join(parts)}")
                detail_run.font.size = Pt(10)
                detail_run.font.color.rgb = GRAY

            p.paragraph_format.space_before = Pt(4)

    # ─── 4. Digital Marketing ─────────────────────────────────
    _section_header(doc, '4. Digital Marketing', RED)

    if digital_marketing:
        p = doc.add_paragraph()
        run = p.add_run(digital_marketing)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(8)
    else:
        _default_digital_marketing(doc, artist, genre_label, countries, GRAY)

    # ─── 5. Retail / DSP Strategy ─────────────────────────────
    _section_header(doc, '5. Retail / DSP Strategy', RED)

    _body_text(doc, (
        f'DMM will pitch {artist} to editorial teams at all major DSPs across LATAM, '
        f'targeting genre-relevant playlists for maximum exposure. The pitching timeline '
        f'follows a 3-phase approach: pre-release (2 weeks before), release week, and '
        f'post-release repitch (1-2 weeks after).'
    ))

    for platform, info in dsp_strategies.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run(platform)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BLUE

        # Strategy description
        if info.get('strategy'):
            sp = doc.add_paragraph()
            run = sp.add_run(info['strategy'])
            run.font.size = Pt(10)
            sp.paragraph_format.space_before = Pt(2)

        # Pitch timing
        if info.get('pitch_timing'):
            tp = doc.add_paragraph()
            run = tp.add_run(f"Pitch timing: {info['pitch_timing']}")
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
            run.italic = True
            tp.paragraph_format.space_before = Pt(2)

        # Target playlists
        if info.get('playlists'):
            pp = doc.add_paragraph()
            run = pp.add_run('Target playlists: ')
            run.font.size = Pt(10)
            run.bold = True
            pl_run = pp.add_run(', '.join(info['playlists']))
            pl_run.font.size = Pt(10)
            pp.paragraph_format.space_before = Pt(2)

    # ─── 6. Press Targets ─────────────────────────────────────
    _section_header(doc, '6. Press Targets', RED)

    if genre != 'general':
        _body_text(doc, (
            f'The following {genre_label.lower()} outlets have been identified as key targets '
            f'for {artist}. DMM will service press releases, coordinate interviews, and '
            f'pursue feature coverage across these outlets.'
        ))
    else:
        _body_text(doc, (
            f'DMM will target the following press outlets across LATAM for {artist}. '
            f'The team will service press releases, coordinate interviews, and '
            f'pursue feature coverage.'
        ))

    # Sort countries by predefined order
    sorted_countries = sorted(
        press_targets.keys(),
        key=lambda c: COUNTRY_ORDER.index(c) if c in COUNTRY_ORDER else 99
    )

    for country in sorted_countries:
        outlets = press_targets[country]
        display = COUNTRY_DISPLAY.get(country, country.title())

        # Country subheader
        cp = doc.add_paragraph()
        cr = cp.add_run(f'{display} ({len(outlets)} outlets)')
        cr.bold = True
        cr.font.size = Pt(10)
        cr.underline = True
        cr.font.color.rgb = GRAY
        cp.paragraph_format.space_before = Pt(12)
        cp.paragraph_format.space_after = Pt(2)

        # Limit to top outlets (by reach or alphabetical)
        shown = sorted(outlets, key=lambda o: _parse_reach(o.get('reach', '')), reverse=True)[:15]
        for outlet in shown:
            op = doc.add_paragraph()
            nr = op.add_run(f"• {outlet['name']}")
            nr.bold = True
            nr.font.size = Pt(10)

            if outlet.get('reach'):
                rr = op.add_run(f"  ({outlet['reach']} reach)")
                rr.font.size = Pt(9)
                rr.font.color.rgb = GRAY

            op.paragraph_format.space_before = Pt(1)

            if outlet.get('description'):
                # Truncate long descriptions
                desc = outlet['description']
                if len(desc) > 150:
                    desc = desc[:147] + '...'
                dp = doc.add_paragraph()
                dr = dp.add_run(f'  {desc}')
                dr.font.size = Pt(9)
                dr.font.color.rgb = GRAY
                dr.italic = True

        if len(outlets) > 15:
            mp = doc.add_paragraph()
            mr = mp.add_run(f'  + {len(outlets) - 15} more outlets in {display}')
            mr.font.size = Pt(9)
            mr.font.color.rgb = GRAY
            mr.italic = True

    # ─── 7. Radio Targets ─────────────────────────────────────
    if radio_targets:
        _section_header(doc, '7. Radio Targets', RED)

        _body_text(doc, (
            f'DMM will service {artist} to the following radio stations across LATAM. '
            f'Stations are selected based on genre affinity ({genre_label.lower()}) '
            f'and market reach.'
        ))

        sorted_radio = sorted(
            radio_targets.keys(),
            key=lambda c: COUNTRY_ORDER.index(c.upper()) if c.upper() in COUNTRY_ORDER else 99
        )

        for country in sorted_radio:
            stations = radio_targets[country]

            # Filter by selected stations if specified
            if radio_stations:
                stations = [s for s in stations if s['station'] in radio_stations]
                if not stations:
                    continue

            cp = doc.add_paragraph()
            cr = cp.add_run(country)
            cr.bold = True
            cr.font.size = Pt(10)
            cr.underline = True
            cr.font.color.rgb = GRAY
            cp.paragraph_format.space_before = Pt(12)
            cp.paragraph_format.space_after = Pt(2)

            for s in stations:
                sp = doc.add_paragraph()
                sr = sp.add_run(f"• {s['station']}")
                sr.bold = True
                sr.font.size = Pt(10)

                detail_parts = []
                if s.get('genre'):
                    detail_parts.append(s['genre'])
                if s.get('format'):
                    detail_parts.append(s['format'])
                if detail_parts:
                    dr = sp.add_run(f"  ({' · '.join(detail_parts)})")
                    dr.font.size = Pt(9)
                    dr.font.color.rgb = GRAY

                if s.get('price'):
                    pr = sp.add_run(f"  — ${s['price']}")
                    pr.font.size = Pt(9)
                    pr.font.color.rgb = BLUE

                sp.paragraph_format.space_before = Pt(2)

                if s.get('notes'):
                    np = doc.add_paragraph()
                    nr = np.add_run(f"  {s['notes']}")
                    nr.font.size = Pt(9)
                    nr.font.color.rgb = GRAY
                    nr.italic = True

    # ─── 8. DJ & Club Servicing ───────────────────────────────
    if dj_markets:
        _section_header(doc, '8. DJ & Club Servicing', RED)

        _body_text(doc, (
            f'DMM will distribute {artist} material to DJs and venue playlists '
            f'in the following markets:'
        ))

        dj_pricing = pricing.get('dj_club', {}).get('markets', {})
        for market in dj_markets:
            info = dj_pricing.get(market, {})
            mp = doc.add_paragraph()
            mr = mp.add_run(f'• {market}')
            mr.bold = True
            mr.font.size = Pt(10)

            if info.get('cities'):
                cr = mp.add_run(f"  — {info['cities']}")
                cr.font.size = Pt(9)
                cr.font.color.rgb = GRAY

            if info.get('price'):
                pr = mp.add_run(f"  (${info['price']:,})")
                pr.font.size = Pt(9)
                pr.font.color.rgb = BLUE

            mp.paragraph_format.space_before = Pt(2)

    # ─── 9. Budget ────────────────────────────────────────────
    _section_header(doc, '9. Budget Breakdown', RED)

    budget_items = []
    total = 0

    # Base fee
    base = pricing.get('base_fee', {})
    base_monthly = base.get('monthly', 2000)
    base_total = base_monthly * campaign_duration
    budget_items.append(('DMM Campaign Management', f'{campaign_duration} months × ${base_monthly:,}/mo', base_total))
    total += base_total

    # Radio
    if radio_targets:
        radio_total = 0
        for country_stations in radio_targets.values():
            for s in country_stations:
                if radio_stations and s['station'] not in radio_stations:
                    continue
                price_str = s.get('price', '0').replace(',', '')
                try:
                    radio_total += int(price_str)
                except ValueError:
                    pass
        if radio_total > 0:
            budget_items.append(('Radio Servicing', f'{sum(len(v) for v in radio_targets.values())} stations, 3 months', radio_total))
            total += radio_total

    # Influencer
    if influencer_tier:
        inf_data = pricing.get('influencers', {}).get('tiers', {}).get(influencer_tier)
        if inf_data:
            inf_total = inf_data.get('estimated_total', 0) * campaign_duration
            budget_items.append(('Influencer Campaign', f"{inf_data['label']} — {inf_data['recommended_count']}", inf_total))
            total += inf_total

    # DJ/Club
    if dj_markets:
        dj_pricing = pricing.get('dj_club', {}).get('markets', {})
        dj_total = sum(dj_pricing.get(m, {}).get('price', 0) for m in dj_markets)
        if dj_total > 0:
            budget_items.append(('DJ & Club Servicing', ', '.join(dj_markets), dj_total))
            total += dj_total

    # Digital marketing
    if digital_package:
        dm_data = pricing.get('digital_marketing', {}).get('packages', {}).get(digital_package)
        if dm_data:
            dm_total = dm_data.get('monthly', 0) * campaign_duration
            budget_items.append(('Digital Marketing', f"{dm_data['label']} — {dm_data['includes']}", dm_total))
            total += dm_total

    # Build budget table
    for label, description, amount in budget_items:
        bp = doc.add_paragraph()
        lr = bp.add_run(f'{label}')
        lr.bold = True
        lr.font.size = Pt(10)

        bp.add_run('\n')

        dr = bp.add_run(f'  {description}')
        dr.font.size = Pt(9)
        dr.font.color.rgb = GRAY

        ar = bp.add_run(f'\n  ${amount:,} USD')
        ar.font.size = Pt(10)
        ar.font.color.rgb = BLUE
        ar.bold = True

        bp.paragraph_format.space_before = Pt(8)

    # Total
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(16)
    sep = tp.add_run('─' * 50)
    sep.font.size = Pt(8)
    sep.font.color.rgb = GRAY

    total_para = doc.add_paragraph()
    tr = total_para.add_run(f'TOTAL CAMPAIGN INVESTMENT')
    tr.bold = True
    tr.font.size = Pt(12)
    tr.font.color.rgb = RED

    amount_para = doc.add_paragraph()
    ar = amount_para.add_run(f'${total:,} USD')
    ar.bold = True
    ar.font.size = Pt(16)
    ar.font.color.rgb = RED
    amount_para.paragraph_format.space_before = Pt(4)

    monthly_para = doc.add_paragraph()
    mr = monthly_para.add_run(f'(${total // campaign_duration:,} USD / month over {campaign_duration} months)')
    mr.font.size = Pt(10)
    mr.font.color.rgb = GRAY
    monthly_para.paragraph_format.space_before = Pt(2)

    # ─── Notes ────────────────────────────────────────────────
    doc.add_paragraph()
    np = doc.add_paragraph()
    nr = np.add_run('Note: ')
    nr.bold = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = GRAY
    nr2 = np.add_run(
        'Press outreach and DSP playlist pitching are included in the base DMM fee. '
        'Radio station pricing is based on current rate cards and may be subject to negotiation. '
        'Budget items can be adjusted based on campaign priorities.'
    )
    nr2.font.size = Pt(9)
    nr2.font.color.rgb = GRAY
    nr2.italic = True

    doc.save(output_path)


# ═════════════════════════════════════════════════════════════════════════
#  HELPERS
# ═════════════════════════════════════════════════════════════════════════

def _section_header(doc, text, color):
    """Add a styled section header."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    return p


def _label_value(doc, label, value, gray_color):
    """Add a label: value line."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    lr = p.add_run(f'{label}: ')
    lr.bold = True
    lr.font.size = Pt(10)
    vr = p.add_run(value)
    vr.font.size = Pt(10)
    vr.font.color.rgb = gray_color
    p.paragraph_format.space_before = Pt(4)
    return p


def _body_text(doc, text):
    """Add a body text paragraph."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    return p


def _default_goal_strategy(doc, artist, genre_label, duration, collaborators):
    """Generate default goal/strategy text."""
    from docx.shared import Pt

    text = (
        f'The primary goal of this campaign is to establish and grow {artist}\'s '
        f'presence across Latin America over a {duration}-month period. '
        f'Leveraging {artist}\'s {genre_label.lower()} sound, DMM will execute a '
        f'multi-pronged strategy combining DSP playlist pitching, radio servicing, '
        f'press outreach, digital marketing, and influencer activations across '
        f'key LATAM markets.'
    )

    if collaborators:
        text += (
            f'\n\nKey campaign hooks include: {collaborators}. These will be '
            f'leveraged as talking points across all press and DSP pitches to '
            f'maximize editorial interest.'
        )

    text += (
        f'\n\nThe campaign follows a phased approach aligned with the release timeline: '
        f'pre-release pitching (2 weeks before each release), release-week activation, '
        f'and post-release sustain (repitch and press follow-up for 2 weeks after).'
    )

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)


def _default_digital_marketing(doc, artist, genre_label, countries, gray_color):
    """Generate default digital marketing section."""
    from docx.shared import Pt

    strategies = [
        f'Geo-targeted social media advertising (Meta, TikTok, YouTube) across key LATAM markets',
        f'Creator campaign with genre-aligned influencers for content seeding',
        f'Content strategy: behind-the-scenes, lyric videos, short-form clips for IG Reels/TikTok/Shorts',
        f'Retargeting campaigns to convert playlist listeners into followers',
        f'A/B creative testing across markets to optimize ad spend',
    ]

    _body_text(doc, (
        f'DMM will execute a targeted digital marketing campaign for {artist} '
        f'focused on audience growth and release awareness across LATAM:'
    ))

    for strategy in strategies:
        p = doc.add_paragraph()
        run = p.add_run(f'• {strategy}')
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(2)


def _parse_reach(reach_str):
    """Parse reach string to integer for sorting."""
    if not reach_str:
        return 0
    cleaned = reach_str.replace(',', '').replace('.', '').strip()
    try:
        return int(cleaned)
    except ValueError:
        return 0


if __name__ == '__main__':
    result = generate_proposal(
        artist='Test Artist',
        genre='electronic',
        timeline=[
            {'title': 'Single 1', 'date': '2026-03-15', 'format': 'Single'},
            {'title': 'EP', 'date': '2026-05-01', 'format': 'EP'},
        ],
        collaborators='ft. Producer X, remix by DJ Y',
        campaign_duration=3,
        dj_markets=['Mexico', 'Brazil'],
    )
    print(f"Generated: {result['output_path']}")

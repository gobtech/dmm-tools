#!/usr/bin/env python3
"""
DSP Pickup Tool
===============
Checks LATAM editorial playlists for artist releases from the release schedule.
Supports Spotify (embed scraping), Deezer (API), Apple Music (page scraping),
Amazon Music (embed scraping), Claro Música (anonymous login + SSR scraping),
and YouTube Music (innertube API).

Usage:
  python dsp_pickup.py --week current              # Check this week's releases
  python dsp_pickup.py --week 2026-02-21           # Check specific week
  python dsp_pickup.py --artist "Djo"              # Check one artist across all playlists
  python dsp_pickup.py --all                        # Check all artists from schedule

Requirements:
  pip install requests
"""

import argparse
import csv
import io
import json
import os
import re
import sys
from datetime import datetime, timedelta
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
from shared.database import load_playlist_database, load_release_schedule

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

PLAYLIST_DB_PATH = os.environ.get(
    'PLAYLIST_DB_PATH',
    str(Path(__file__).parent.parent / 'data' / 'playlist_database.csv')
)

RELEASE_SCHEDULE_URL = os.environ.get(
    'RELEASE_SCHEDULE_URL',
    'https://docs.google.com/spreadsheets/d/e/2PACX-1vTSd9mhkVibb7AwXtsZjRgBfuRT9sLY_qhhu-rB_P35CX2vFk_fZw_f31AJyW84KrCzWLLMUcTzzgqU/pub?gid=497066221&single=true&output=csv'
)


def get_spotify_playlist_tracks(playlist_id):
    """
    Fetch tracks from a Spotify playlist via the public embed endpoint.
    Returns (tracks_list, playlist_cover_url).
    """
    import requests
    import time

    url = f"https://open.spotify.com/embed/playlist/{playlist_id}"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    }

    try:
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
    except Exception as e:
        print(f"    Error fetching playlist {playlist_id}: {e}")
        return [], ''

    # Extract __NEXT_DATA__ JSON from the embed page
    match = re.search(r'<script id="__NEXT_DATA__"[^>]*>(.*?)</script>', resp.text, re.DOTALL)
    if not match:
        print(f"    Could not parse playlist {playlist_id}")
        return [], ''

    try:
        data = json.loads(match.group(1))
        entity = data['props']['pageProps']['state']['data']['entity']
        track_list = entity.get('trackList', [])
    except (json.JSONDecodeError, KeyError) as e:
        print(f"    Error parsing playlist data: {e}")
        return [], ''

    # Playlist cover art
    cover_url = ''
    try:
        sources = entity.get('coverArt', {}).get('sources', [])
        if sources:
            cover_url = sources[0].get('url', '')
    except (KeyError, IndexError, TypeError):
        pass

    tracks = []
    for i, item in enumerate(track_list):
        subtitle = item.get('subtitle', '')
        artists = [a.strip() for a in subtitle.split(',')] if subtitle else []
        track_uri = item.get('uri', '')
        tracks.append({
            'artist': subtitle,
            'artists_list': artists,
            'track': item.get('title', ''),
            'album': '',
            'position': i + 1,
            'artwork_url': '',
            'spotify_uri': track_uri,
        })

    time.sleep(0.5)
    return tracks, cover_url


def get_deezer_playlist_tracks(playlist_id):
    """
    Fetch tracks from a Deezer playlist using their public API.
    Returns (tracks_list, playlist_cover_url).
    """
    import requests

    # Fetch playlist metadata (cover art)
    cover_url = ''
    try:
        pl_resp = requests.get(f"https://api.deezer.com/playlist/{playlist_id}", timeout=10)
        pl_data = pl_resp.json()
        cover_url = pl_data.get('picture_xl', '') or pl_data.get('picture_big', '') or pl_data.get('picture_medium', '') or pl_data.get('picture', '')
    except Exception:
        pass

    tracks = []
    url = f"https://api.deezer.com/playlist/{playlist_id}/tracks"

    while url:
        try:
            response = requests.get(url, timeout=10)
            data = response.json()
        except Exception as e:
            print(f"  Error fetching Deezer playlist {playlist_id}: {e}")
            break

        for i, item in enumerate(data.get('data', [])):
            album_obj = item.get('album', {})
            artwork_url = album_obj.get('cover_big', '') or album_obj.get('cover_xl', '') or album_obj.get('cover_medium', '') or album_obj.get('cover', '')
            tracks.append({
                'artist': item.get('artist', {}).get('name', ''),
                'artists_list': [item.get('artist', {}).get('name', '')],
                'track': item.get('title', ''),
                'album': album_obj.get('title', ''),
                'position': len(tracks) + 1,
                'artwork_url': artwork_url,
            })

        url = data.get('next')

    return tracks, cover_url


def get_apple_music_playlist_tracks(playlist_id):
    """
    Fetch tracks from an Apple Music playlist by scraping the public page.
    Returns (tracks_list, playlist_cover_url).
    """
    import requests
    import time

    url = f"https://music.apple.com/us/playlist/-/{playlist_id}"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    }

    try:
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
    except Exception as e:
        print(f"    Error fetching Apple Music playlist {playlist_id}: {e}")
        return [], ''

    tracks = []
    cover_url = ''

    # Try hydration JSON first: <script type="application/json">
    json_blocks = re.findall(
        r'<script[^>]*type="application/json"[^>]*>(.*?)</script>', resp.text, re.DOTALL
    )
    for block in json_blocks:
        try:
            data = json.loads(block)
        except (json.JSONDecodeError, ValueError):
            continue

        # Navigate to sections — handle both structures:
        #   dict: data['data'][0]['data']['sections']
        #   list: data[0]['data']['sections']
        try:
            root = data['data'] if isinstance(data, dict) and 'data' in data else data
            if isinstance(root, list) and len(root) > 0:
                sections = root[0]['data']['sections']
                # Extract playlist cover from header section (first section, first item)
                if not cover_url:
                    for sec in sections:
                        if 'header' in sec.get('id', '').lower() or sec.get('id', '').startswith('playlist-detail'):
                            header_items = sec.get('items', [])
                            if header_items:
                                art_obj = header_items[0].get('artwork', {})
                                # Apple uses artwork.dictionary.url or artwork.url
                                art_dict = art_obj.get('dictionary', art_obj)
                                art_raw = art_dict.get('url', '')
                                if art_raw:
                                    cover_url = art_raw.replace('{w}', '600').replace('{h}', '600').replace('{f}', 'jpg').replace('{c}', '')
                                    break
            else:
                continue
        except (KeyError, IndexError, TypeError):
            continue

        for section in sections:
            # Prefer section with 'track-list' in its id
            sec_id = section.get('id', '')
            items = section.get('items', [])
            if not items or len(items) < 2:
                continue
            first = items[0]
            if not isinstance(first, dict) or 'title' not in first:
                continue
            # Must be a track section (has artist info or is explicitly track-list)
            is_track_section = (
                'track-list' in sec_id or
                'subtitleLinks' in first or
                'artistName' in first
            )
            if not is_track_section:
                continue

            for i, item in enumerate(items):
                artist_name = item.get('artistName', '')
                # Build artists_list from subtitleLinks (each artist is a separate entry)
                subtitle_links = item.get('subtitleLinks', [])
                if subtitle_links and isinstance(subtitle_links, list):
                    artists_list = [link.get('title', '') for link in subtitle_links if link.get('title')]
                else:
                    # Fallback: split artistName on comma and ampersand
                    artists_list = [a.strip() for a in re.split(r'[,&]', artist_name) if a.strip()]
                # Build display artist from subtitleLinks if artistName is empty
                if not artist_name and artists_list:
                    artist_name = ', '.join(artists_list)
                album = ''
                tertiary = item.get('tertiaryLinks', [])
                if tertiary and isinstance(tertiary, list):
                    album = tertiary[0].get('title', '')

                # Artwork: Apple Music uses artwork.dictionary.url or artwork.url with {w}x{h} template
                artwork_url = ''
                artwork_obj = item.get('artwork', {})
                if isinstance(artwork_obj, dict):
                    art_dict = artwork_obj.get('dictionary', artwork_obj)
                    art_raw = art_dict.get('url', '')
                    if art_raw:
                        artwork_url = art_raw.replace('{w}', '300').replace('{h}', '300').replace('{f}', 'jpg').replace('{c}', '')
                if not artwork_url and isinstance(item.get('artworkUrl'), str):
                    artwork_url = item['artworkUrl']

                tracks.append({
                    'artist': artist_name,
                    'artists_list': artists_list,
                    'track': item.get('title', ''),
                    'album': album,
                    'position': i + 1,
                    'artwork_url': artwork_url,
                })

            if tracks:
                break  # Found the track section, stop searching

        if tracks:
            break  # Found data in this JSON block

    # Fallback: try ld+json for partial data (track names, no artist)
    if not tracks:
        ld_blocks = re.findall(
            r'<script[^>]*type="application/ld\+json"[^>]*>(.*?)</script>', resp.text, re.DOTALL
        )
        for block in ld_blocks:
            try:
                ld_data = json.loads(block)
            except (json.JSONDecodeError, ValueError):
                continue

            # MusicPlaylist schema has 'track' as dict with 'itemListElement' or as a list
            track_list = None
            playlist_obj = None
            if isinstance(ld_data, dict) and ld_data.get('@type') == 'MusicPlaylist':
                playlist_obj = ld_data
            elif isinstance(ld_data, list):
                for item in ld_data:
                    if isinstance(item, dict) and item.get('@type') == 'MusicPlaylist':
                        playlist_obj = item
                        break

            if playlist_obj:
                track_obj = playlist_obj.get('track', [])
                if isinstance(track_obj, dict):
                    track_list = track_obj.get('itemListElement', [])
                elif isinstance(track_obj, list):
                    track_list = track_obj

            if track_list:
                for i, entry in enumerate(track_list):
                    if isinstance(entry, dict):
                        item = entry.get('item', entry)
                    else:
                        continue
                    by_artist = item.get('byArtist', {})
                    artist_name = ''
                    if isinstance(by_artist, dict):
                        artist_name = by_artist.get('name', '')
                    elif isinstance(by_artist, str):
                        artist_name = by_artist
                    artists_list = [a.strip() for a in re.split(r'[,&]', artist_name) if a.strip()] if artist_name else []
                    img = item.get('image', '')
                    artwork_url = img if isinstance(img, str) else (img.get('url', '') if isinstance(img, dict) else '')
                    tracks.append({
                        'artist': artist_name,
                        'artists_list': artists_list,
                        'track': item.get('name', ''),
                        'album': '',
                        'position': i + 1,
                        'artwork_url': artwork_url,
                    })
                break

    if not tracks:
        print(f"    Could not parse Apple Music playlist {playlist_id}")

    # Apple Music is stricter — longer delay between requests
    time.sleep(1)
    return tracks, cover_url


def get_ytmusic_playlist_tracks(playlist_id):
    """
    Fetch tracks from a YouTube Music playlist via the innertube API.
    Returns (tracks_list, playlist_cover_url).
    """
    import requests
    import time

    url = "https://music.youtube.com/youtubei/v1/browse?prettyPrint=false"
    payload = {
        "context": {
            "client": {
                "clientName": "WEB_REMIX",
                "clientVersion": "1.20240101.01.00",
                "hl": "en",
                "gl": "US",
            }
        },
        "browseId": f"VL{playlist_id}",
    }
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Content-Type': 'application/json',
        'Referer': 'https://music.youtube.com/',
        'Origin': 'https://music.youtube.com',
    }

    try:
        resp = requests.post(url, json=payload, headers=headers, timeout=15)
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        print(f"    Error fetching YouTube Music playlist {playlist_id}: {e}")
        return [], ''

    tracks = []
    cover_url = ''

    # Extract playlist cover from header
    try:
        header = data.get('header', {}).get('musicImmersiveHeaderRenderer', {})
        if not header:
            header = data.get('header', {}).get('musicDetailHeaderRenderer', {})
        thumbs = header.get('thumbnail', {}).get('musicThumbnailRenderer', {}).get('thumbnail', {}).get('thumbnails', [])
        if thumbs:
            cover_url = thumbs[-1].get('url', '')
    except (KeyError, IndexError, TypeError):
        pass

    try:
        renderer = data['contents']['twoColumnBrowseResultsRenderer']
        shelf = renderer['secondaryContents']['sectionListRenderer']['contents'][0]['musicPlaylistShelfRenderer']
        items = shelf['contents']
    except (KeyError, IndexError, TypeError):
        # Try alternative structure
        try:
            renderer = data['contents']['singleColumnBrowseResultsRenderer']
            tabs = renderer['tabs'][0]['tabRenderer']['content']
            shelf = tabs['sectionListRenderer']['contents'][0]['musicPlaylistShelfRenderer']
            items = shelf['contents']
        except (KeyError, IndexError, TypeError):
            print(f"    Could not parse YouTube Music playlist {playlist_id}")
            return [], cover_url

    for i, item in enumerate(items):
        try:
            mrlir = item['musicResponsiveListItemRenderer']
            flex_cols = mrlir.get('flexColumns', [])
            col_texts = []
            for fc in flex_cols:
                col = fc['musicResponsiveListItemFlexColumnRenderer']
                runs = col.get('text', {}).get('runs', [])
                # Filter out separator characters
                texts = [r['text'] for r in runs if r['text'] not in (' \u2022 ', ' & ', ', ')]
                col_texts.append(texts)

            title = col_texts[0][0] if col_texts and col_texts[0] else ''
            # Artist is in the second column — may have multiple artists
            artists_list = col_texts[1] if len(col_texts) > 1 else []
            artist_display = ', '.join(artists_list)
            album = col_texts[2][0] if len(col_texts) > 2 and col_texts[2] else ''

            # Artwork: thumbnail.musicThumbnailRenderer.thumbnail.thumbnails (pick largest)
            artwork_url = ''
            try:
                thumbs = mrlir['thumbnail']['musicThumbnailRenderer']['thumbnail']['thumbnails']
                if thumbs:
                    artwork_url = thumbs[-1].get('url', '')  # last = highest res
            except (KeyError, IndexError, TypeError):
                pass

            tracks.append({
                'artist': artist_display,
                'artists_list': artists_list,
                'track': title,
                'album': album,
                'position': i + 1,
                'artwork_url': artwork_url,
            })
        except (KeyError, IndexError, TypeError):
            continue

    time.sleep(0.5)
    return tracks, cover_url


def get_amazon_music_playlist_tracks(playlist_id, domain='com.mx'):
    """
    Fetch tracks from an Amazon Music playlist via the public embed page.
    Returns (tracks_list, playlist_cover_url).
    """
    import requests
    import time

    url = f"https://music.amazon.{domain}/embed/{playlist_id}"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    }

    try:
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
    except Exception as e:
        print(f"    Error fetching Amazon Music playlist {playlist_id}: {e}")
        return [], ''

    tracks = []
    import html as html_module

    # Extract playlist cover from header image
    # The cover is an <img> inside <a class="headerImg playlistHeader ...">
    cover_url = ''
    cover_match = re.search(r'class="[^"]*(?:headerImg|playlistHeader)[^"]*"[^>]*>\s*<img[^>]*src="([^"]+)"', resp.text)
    if not cover_match:
        # Fallback: img with alt containing "portada" or "cover"
        cover_match = re.search(r'<img[^>]*alt="[^"]*(?:portada|cover)[^"]*"[^>]*src="([^"]+)"', resp.text, re.IGNORECASE)
    if cover_match:
        cover_url = html_module.unescape(cover_match.group(1))

    # Extract artwork URLs from <li> elements — background-image in style or <img> tags
    # Build a map of position -> artwork_url
    # Each track <li> has aria-posinset="N" and contains a div.imageBackground with background-image
    artwork_map = {}
    art_matches = re.findall(
        r'aria-posinset="(\d+)".*?(?:background-image:\s*url\([\'"]?([^\)\'\"]+)[\'"]?\)|<img[^>]*class="[^"]*(?:trackListImage|albumArt|trackIndexImg)[^"]*"[^>]*src="([^"]+)")',
        resp.text, re.DOTALL
    )
    for pos, bg_url, img_url in art_matches:
        artwork_map[int(pos)] = html_module.unescape(bg_url or img_url)

    # Each track is in an <li> with aria-posinset, artist in .trackListArtist, title in .trackListTitle
    # aria-label format: "cancion, {title}" or "song, {title}" / "artista, {artist}" or "artist, {artist}"
    track_matches = re.findall(
        r'aria-posinset="(\d+)".*?'
        r'class="trackListTitle truncate">\s*<a[^>]*aria-label="[^,]*,\s*(.*?)"[^>]*>.*?'
        r'class="trackListArtist truncate">\s*<a[^>]*aria-label="[^,]*,\s*(.*?)"[^>]*>',
        resp.text, re.DOTALL
    )

    for position, title, artist in track_matches:
        title = html_module.unescape(title).strip()
        artist = html_module.unescape(artist).strip()
        artists_list = [a.strip() for a in re.split(r'[,&]', artist) if a.strip()]
        tracks.append({
            'artist': artist,
            'artists_list': artists_list,
            'track': title,
            'album': '',
            'position': int(position),
            'artwork_url': artwork_map.get(int(position), ''),
        })

    time.sleep(0.5)
    return tracks, cover_url


def get_claro_playlist_tracks(playlist_id, country='MX'):
    """
    Fetch tracks from a Claro Música playlist via anonymous login + SSR state.
    Returns (tracks_list, playlist_cover_url).
    Returns list of { artist, artists_list, track, album, position }.
    """
    import requests
    import time

    playlist_path = f"/systemPlaylist/{playlist_id}"
    if country != 'MX':
        playlist_path += f"/{country}"

    login_url = f"https://www.claromusica.com/anonymousLogin/{country}?redirectTo={playlist_path}"
    page_url = f"https://www.claromusica.com{playlist_path}"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    }

    try:
        session = requests.Session()
        session.headers.update(headers)
        # Step 1: anonymous login to get session cookies (don't follow redirect)
        session.get(login_url, allow_redirects=False, timeout=15)
        # Step 2: fetch the playlist page with cookies set
        resp = session.get(page_url, timeout=15)
        resp.raise_for_status()
    except Exception as e:
        print(f"    Error fetching Claro Música playlist {playlist_id}: {e}")
        return [], ''

    # Extract window.__PRELOADED_STATE__ JSON
    marker = 'window.__PRELOADED_STATE__ = '
    idx = resp.text.find(marker)
    if idx == -1:
        print(f"    Could not find preloaded state for Claro playlist {playlist_id}")
        return [], ''

    # Extract JSON: starts after marker, ends at </script>
    json_start = idx + len(marker)
    json_end = resp.text.find('</script>', json_start)
    if json_end == -1:
        print(f"    Could not find end of preloaded state for Claro playlist {playlist_id}")
        return [], ''

    raw_json = resp.text[json_start:json_end].rstrip().rstrip(';')

    try:
        state = json.loads(raw_json)
    except (json.JSONDecodeError, ValueError) as e:
        print(f"    Error parsing Claro Música state: {e}")
        return [], ''

    playlist_detail = state.get('playlistDetail', {})
    track_list = playlist_detail.get('tracks', [])

    # Playlist cover
    cover_url = ''
    for key in ('image', 'imageUrl', 'cover', 'artwork'):
        val = playlist_detail.get(key, '')
        if val and isinstance(val, str) and val.startswith('http'):
            cover_url = val
            break
    if not cover_url:
        pl_images = playlist_detail.get('images', [])
        if isinstance(pl_images, list) and pl_images:
            first = pl_images[0]
            cover_url = first.get('url', '') if isinstance(first, dict) else str(first)

    tracks = []
    for i, item in enumerate(track_list):
        # Artist can be a list or nested objects
        artist_names = item.get('artist', [])
        if isinstance(artist_names, list):
            artists_list = [a if isinstance(a, str) else a.get('name', '') for a in artist_names]
        else:
            artists_list = [str(artist_names)]
        artist_display = ', '.join(artists_list)

        # Artwork: try common field names in Claro's SSR state
        artwork_url = ''
        for key in ('image', 'imageUrl', 'albumImage', 'cover', 'artwork'):
            val = item.get(key, '')
            if val and isinstance(val, str) and val.startswith('http'):
                artwork_url = val
                break
        if not artwork_url:
            images = item.get('images', [])
            if isinstance(images, list) and images:
                first = images[0]
                artwork_url = first.get('url', '') if isinstance(first, dict) else str(first)

        tracks.append({
            'artist': artist_display,
            'artists_list': artists_list,
            'track': item.get('name', ''),
            'album': item.get('albumName', ''),
            'position': i + 1,
            'artwork_url': artwork_url,
        })

    time.sleep(1)
    return tracks, cover_url


def normalize_name(name):
    """Normalize artist/track name for fuzzy matching."""
    name = name.lower().strip()
    # Remove common suffixes/prefixes
    name = re.sub(r'\s*\(.*?\)\s*', '', name)  # Remove parenthetical
    name = re.sub(r'\s*\[.*?\]\s*', '', name)  # Remove brackets
    name = re.sub(r'\s*-\s*.*$', '', name)       # Remove after dash (for "Artist - feat. X")
    name = re.sub(r'\s*feat\.?\s.*$', '', name, flags=re.IGNORECASE)
    name = re.sub(r'\s*ft\.?\s.*$', '', name, flags=re.IGNORECASE)
    # Normalize unicode
    name = name.replace('é', 'e').replace('á', 'a').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
    name = name.replace('ñ', 'n').replace('ü', 'u').replace('ö', 'o')
    return name.strip()


def check_release_in_playlist(release, playlist_tracks):
    """
    Check if a release appears in a playlist's tracks.
    Returns match info or None.
    """
    release_artist = normalize_name(release['artist'])
    release_title = normalize_name(release['title'])

    # Clean focus_track: treat placeholders like "-", "–", "N/A", "TBD" as empty
    raw_focus = (release.get('focus_track', '') or '').strip()
    if raw_focus in ('-', '–', '—', 'N/A', 'n/a', 'TBD', 'tbd', ''):
        focus_track = ''
    else:
        focus_track = normalize_name(raw_focus)

    artist_only_mode = not release_title and not focus_track

    for track in playlist_tracks:
        # Check if any artist matches
        artist_match = False
        for pl_artist in track.get('artists_list', [track.get('artist', '')]):
            pl_artist_norm = normalize_name(pl_artist)
            if not pl_artist_norm:
                continue
            if pl_artist_norm == release_artist:
                artist_match = True
                break
            # Partial match — require minimum length to avoid false positives
            # Skip partial matching in artist-only mode (no title to cross-validate)
            if not artist_only_mode and len(release_artist) >= 3 and len(pl_artist_norm) >= 3:
                if release_artist in pl_artist_norm or pl_artist_norm in release_artist:
                    artist_match = True
                    break

        if not artist_match:
            continue

        # If release has no title and no focus track, exact artist match alone is enough
        if artist_only_mode:
            return {
                'playlist_track': track.get('track', ''),
                'playlist_artist': track.get('artist', ''),
                'position': track.get('position', '?'),
                'added_at': track.get('added_at', ''),
                'artwork_url': track.get('artwork_url', ''),
                'spotify_uri': track.get('spotify_uri', ''),
            }

        # Check track/album title match
        pl_track = normalize_name(track.get('track', ''))
        pl_album = normalize_name(track.get('album', ''))

        # Skip empty values in substring checks to prevent "" matching everything
        title_match = False
        if release_title and pl_track:
            if release_title == pl_track:
                title_match = True
            # Substring match only if the shorter string is long enough (>=4 chars)
            elif len(release_title) >= 4 and len(pl_track) >= 4:
                if release_title in pl_track or pl_track in release_title:
                    title_match = True
        if not title_match and focus_track and pl_track:
            if focus_track == pl_track:
                title_match = True
            elif len(focus_track) >= 4 and len(pl_track) >= 4:
                if focus_track in pl_track or pl_track in focus_track:
                    title_match = True
        if not title_match and release_title and pl_album:
            if release_title == pl_album:
                title_match = True

        if title_match:
            return {
                'playlist_track': track.get('track', ''),
                'playlist_artist': track.get('artist', ''),
                'position': track.get('position', '?'),
                'added_at': track.get('added_at', ''),
                'artwork_url': track.get('artwork_url', ''),
                'spotify_uri': track.get('spotify_uri', ''),
            }
    
    return None


def generate_proof_image(match, output_dir):
    """
    Generate a composite 'proof' image for a DSP playlist match.
    Two-section layout: playlist header (with playlist cover art) + highlighted track row.
    Returns the path to the generated image, or None on failure.
    """
    import requests as _requests
    import unicodedata

    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        print("    Pillow not installed — skipping proof image generation")
        return None

    track_artwork_url = match.get('artwork_url', '')
    playlist_cover_url = match.get('playlist_cover_url', '')
    # Spotify: resolve track artwork from URI via oEmbed (lazy, only for matches)
    if not track_artwork_url and match.get('spotify_uri', ''):
        try:
            track_id = match['spotify_uri'].split(':')[-1]
            oembed_url = f"https://open.spotify.com/oembed?url=https://open.spotify.com/track/{track_id}"
            oembed_resp = _requests.get(oembed_url, timeout=8)
            if oembed_resp.ok:
                track_artwork_url = oembed_resp.json().get('thumbnail_url', '')
        except Exception:
            pass

    track_name = match.get('playlist_track', '')
    artist_name = match.get('playlist_artist', '')
    playlist_name = match.get('playlist_name', '')
    position = match.get('position', '?')
    platform = match.get('platform', '')
    country = match.get('playlist_country', '')
    followers = match.get('playlist_followers', '')

    # Platform accent colors
    platform_colors = {
        'Spotify': (30, 185, 84),
        'Deezer': (162, 56, 255),
        'Apple Music': (252, 60, 68),
        'Amazon Music': (0, 168, 225),
        'Claro Música': (229, 9, 20),
        'YouTube Music': (255, 0, 0),
    }
    accent = platform_colors.get(platform, (196, 48, 48))

    # Dimensions — render at 3× for sharp output in .docx
    S = 3
    W = 640 * S
    HEADER_H = 120 * S
    TRACK_H = 60 * S
    H = HEADER_H + TRACK_H
    PAD = 16 * S
    PL_COVER_SIZE = 88 * S
    TRACK_ART_SIZE = 40 * S

    # Create image
    img = Image.new('RGB', (W, H), (18, 18, 18))
    draw = ImageDraw.Draw(img)

    # Fonts
    try:
        font_pl_name = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 16 * S)
        font_pl_meta = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12 * S)
        font_track = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 14 * S)
        font_artist = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12 * S)
        font_pos = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12 * S)
        font_badge = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 10 * S)
    except (IOError, OSError):
        font_pl_name = ImageFont.load_default()
        font_pl_meta = font_pl_name
        font_track = font_pl_name
        font_artist = font_pl_name
        font_pos = font_pl_name
        font_badge = font_pl_name

    def truncate(text, font, max_w):
        if not text:
            return ''
        bbox = draw.textbbox((0, 0), text, font=font)
        if bbox[2] - bbox[0] <= max_w:
            return text
        while len(text) > 3:
            text = text[:-1]
            bbox = draw.textbbox((0, 0), text + '…', font=font)
            if bbox[2] - bbox[0] <= max_w:
                return text + '…'
        return text

    def _fetch_image(url, size):
        """Fetch and resize an image from URL."""
        try:
            resp = _requests.get(url, timeout=8)
            resp.raise_for_status()
            im = Image.open(io.BytesIO(resp.content)).convert('RGB')
            return im.resize(size, Image.LANCZOS)
        except Exception:
            return None

    # ─── SECTION 1: Playlist header ───────────────────────────────
    # Background: slightly lighter dark
    draw.rectangle([0, 0, W, HEADER_H], fill=(28, 28, 28))

    # Accent stripe on top
    draw.rectangle([0, 0, W, 3 * S], fill=accent)

    # Playlist cover art
    pl_x = PAD
    pl_y = PAD + 4 * S
    if playlist_cover_url:
        pl_img = _fetch_image(playlist_cover_url, (PL_COVER_SIZE, PL_COVER_SIZE))
        if pl_img:
            img.paste(pl_img, (pl_x, pl_y))
        else:
            draw.rectangle([pl_x, pl_y, pl_x + PL_COVER_SIZE, pl_y + PL_COVER_SIZE], fill=(50, 50, 50))
    else:
        draw.rectangle([pl_x, pl_y, pl_x + PL_COVER_SIZE, pl_y + PL_COVER_SIZE], fill=(50, 50, 50))

    # Playlist text (to the right of cover)
    txt_x = pl_x + PL_COVER_SIZE + PAD
    txt_max_w = W - txt_x - PAD

    # Platform badge (top right of header)
    badge_text = platform
    badge_bbox = draw.textbbox((0, 0), badge_text, font=font_badge)
    badge_w = badge_bbox[2] - badge_bbox[0] + 14 * S
    badge_h = badge_bbox[3] - badge_bbox[1] + 8 * S
    badge_x = W - PAD - badge_w
    badge_y = pl_y + 2 * S
    draw.rounded_rectangle([badge_x, badge_y, badge_x + badge_w, badge_y + badge_h], radius=4 * S, fill=accent)
    draw.text((badge_x + 7 * S, badge_y + 3 * S), badge_text, fill=(255, 255, 255), font=font_badge)

    # Playlist name
    pl_name_max = txt_max_w - badge_w - 12 * S
    y = pl_y + 6 * S
    draw.text((txt_x, y), truncate(playlist_name, font_pl_name, pl_name_max), fill=(255, 255, 255), font=font_pl_name)
    y += 24 * S

    # Country + followers
    meta_parts = [country]
    if followers:
        meta_parts.append(f"{followers} followers")
    meta_text = ' • '.join(meta_parts)
    draw.text((txt_x, y), truncate(meta_text, font_pl_meta, txt_max_w), fill=(160, 160, 160), font=font_pl_meta)
    y += 20 * S

    # Playlist link hint
    link = match.get('playlist_link', '')
    if link:
        draw.text((txt_x, y), truncate(link, font_pl_meta, txt_max_w), fill=(100, 100, 100), font=font_pl_meta)

    # ─── SECTION 2: Track row (highlighted) ────────────────────────
    row_y = HEADER_H
    # Highlight background
    draw.rectangle([0, row_y, W, H], fill=(40, 40, 40))
    # Left accent bar on the track row
    draw.rectangle([0, row_y, 3 * S, H], fill=accent)

    # Position number
    pos_text = f"#{position}"
    pos_bbox = draw.textbbox((0, 0), pos_text, font=font_pos)
    pos_w = pos_bbox[2] - pos_bbox[0]
    draw.text((PAD + 4 * S, row_y + (TRACK_H - 14 * S) // 2), pos_text, fill=accent, font=font_pos)

    # Track album art (small thumbnail)
    art_x = PAD + max(pos_w, 24 * S) + 12 * S
    art_y = row_y + (TRACK_H - TRACK_ART_SIZE) // 2
    if track_artwork_url:
        t_img = _fetch_image(track_artwork_url, (TRACK_ART_SIZE, TRACK_ART_SIZE))
        if t_img:
            img.paste(t_img, (art_x, art_y))
        else:
            draw.rectangle([art_x, art_y, art_x + TRACK_ART_SIZE, art_y + TRACK_ART_SIZE], fill=(60, 60, 60))
    else:
        draw.rectangle([art_x, art_y, art_x + TRACK_ART_SIZE, art_y + TRACK_ART_SIZE], fill=(60, 60, 60))

    # Track name + artist
    track_txt_x = art_x + TRACK_ART_SIZE + 12 * S
    track_max_w = W - track_txt_x - PAD
    draw.text((track_txt_x, row_y + 10 * S), truncate(track_name, font_track, track_max_w), fill=(255, 255, 255), font=font_track)
    draw.text((track_txt_x, row_y + 30 * S), truncate(artist_name, font_artist, track_max_w), fill=(180, 180, 180), font=font_artist)

    # ─── Save ──────────────────────────────────────────────────────
    def _ascii_safe(s, maxlen):
        s = unicodedata.normalize('NFKD', s).encode('ascii', 'ignore').decode()
        return re.sub(r'[^\w\s-]', '', s)[:maxlen].strip().replace(' ', '_') or 'item'
    safe_track = _ascii_safe(track_name, 30)
    safe_playlist = _ascii_safe(playlist_name, 20)
    filename = f"proof_{safe_track}_{safe_playlist}.png"
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, filename)
    img.save(filepath, 'PNG', dpi=(300, 300))
    return filepath


def generate_proof_images(results, output_dir):
    """
    Generate proof images for all DSP matches.
    Returns a dict: { artist: { title: [image_paths] } }
    """
    image_paths = {}
    for artist, releases in results.items():
        image_paths[artist] = {}
        for title, matches in releases.items():
            image_paths[artist][title] = []
            for match in matches:
                path = generate_proof_image(match, output_dir)
                if path:
                    image_paths[artist][title].append(path)
    return image_paths


def generate_dsp_docx(results, proof_image_paths, docx_path, grouping='platform'):
    """
    Generate a formatted .docx report for DSP Pickup results.
    Matches the Dorado report style: title, then per-match entries with
    platform (blue bold), country, playlist name + followers + date (bold),
    and embedded proof image.
    """
    from datetime import datetime as _dt

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("    python-docx not installed — skipping .docx generation")
        return

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    today_str = _dt.now().strftime('%b %d, %Y')

    # Title: "Streaming / Playlists Highlights" in bold red
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Streaming / Playlists Highlights')
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0xC4, 0x30, 0x30)
    title_run.font.size = Pt(14)
    title_para.paragraph_format.space_after = Pt(12)

    # Collect all matches flat, tagging each with its artist/release
    all_matches = []
    for artist, releases in results.items():
        for title, matches in releases.items():
            for m in matches:
                m['_artist'] = artist
                m['_release'] = title
                all_matches.append(m)

    if not all_matches:
        p = doc.add_paragraph()
        p.add_run('No playlist placements found.')
        doc.save(docx_path)
        return

    platform_order = ['Spotify', 'Apple Music', 'Deezer', 'Amazon Music', 'YouTube Music', 'Claro Música']

    def platform_idx(m):
        try:
            return platform_order.index(m.get('platform', ''))
        except ValueError:
            return 99

    # Helper to render a single match entry (country, playlist info, proof image)
    import unicodedata as _ud
    def _ascii_safe(s, maxlen):
        s = _ud.normalize('NFKD', s).encode('ascii', 'ignore').decode()
        return re.sub(r'[^\w\s-]', '', s)[:maxlen].strip().replace(' ', '_') or 'item'

    def _render_match(m):
        platform = m.get('platform', '')
        playlist_name = m.get('playlist_name', '')
        country = m.get('playlist_country', '')
        followers = m.get('playlist_followers', '')
        track = m.get('playlist_track', '')
        artist = m.get('playlist_artist', '')

        # Country (underlined) if present
        if country:
            cp = doc.add_paragraph()
            cr = cp.add_run(country)
            cr.underline = True
            cr.font.size = Pt(10)
            cp.paragraph_format.space_before = Pt(4)
            cp.paragraph_format.space_after = Pt(1)

        # Playlist name + followers + date (bold black)
        info_parts = [playlist_name]
        if followers:
            info_parts.append(followers)
        info_parts.append(today_str)
        info_text = ' - '.join(info_parts)

        ip = doc.add_paragraph()
        ir = ip.add_run(info_text)
        ir.bold = True
        ir.font.size = Pt(10)
        ip.paragraph_format.space_before = Pt(2)
        ip.paragraph_format.space_after = Pt(4)

        # Embed proof image if available
        safe_track = _ascii_safe(track, 30)
        safe_playlist = _ascii_safe(playlist_name, 20)
        img_filename = f"proof_{safe_track}_{safe_playlist}.png"
        img_path = os.path.join(os.path.dirname(docx_path), 'dsp_proofs', img_filename)

        if os.path.exists(img_path):
            img_para = doc.add_paragraph()
            img_para.paragraph_format.space_after = Pt(8)
            run = img_para.add_run()
            run.add_picture(img_path, width=Inches(6.2))
        else:
            # Fallback: show text description
            fp = doc.add_paragraph()
            fr = fp.add_run(f'  #{m.get("position", "?")} — {track} by {artist}')
            fr.font.size = Pt(9)
            fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            fp.paragraph_format.space_after = Pt(8)

    if grouping == 'artist':
        # Sort by artist, then release, then platform
        all_matches.sort(key=lambda m: (m['_artist'].lower(), m['_release'].lower(), platform_idx(m), m.get('playlist_name', '')))

        current_artist = None
        current_release = None
        for m in all_matches:
            # Artist header (red bold)
            if m['_artist'] != current_artist:
                if current_artist is not None:
                    doc.add_paragraph().paragraph_format.space_before = Pt(6)
                p = doc.add_paragraph()
                run = p.add_run(m['_artist'])
                run.bold = True
                run.font.color.rgb = RGBColor(0xC4, 0x30, 0x30)
                run.font.size = Pt(13)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                current_artist = m['_artist']
                current_release = None

            # Release sub-header (dark gray bold)
            if m['_release'] != current_release:
                rp = doc.add_paragraph()
                rr = rp.add_run(m['_release'])
                rr.bold = True
                rr.font.size = Pt(11)
                rr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                rp.paragraph_format.space_before = Pt(6)
                rp.paragraph_format.space_after = Pt(2)
                current_release = m['_release']

            # Platform label (blue, smaller)
            pp = doc.add_paragraph()
            pr = pp.add_run(m.get('platform', ''))
            pr.bold = True
            pr.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
            pr.font.size = Pt(10)
            pp.paragraph_format.space_before = Pt(4)
            pp.paragraph_format.space_after = Pt(1)

            _render_match(m)
    else:
        # Default: group by platform, then playlist name
        all_matches.sort(key=lambda m: (platform_idx(m), m.get('playlist_name', '')))

        current_platform = None
        for m in all_matches:
            platform = m.get('platform', '')

            # Platform header (blue bold) — only when platform changes
            if platform != current_platform:
                if current_platform is not None:
                    doc.add_paragraph().paragraph_format.space_before = Pt(6)
                p = doc.add_paragraph()
                run = p.add_run(platform)
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                current_platform = platform

            _render_match(m)

    doc.save(docx_path)
    print(f"  DSP report saved: {docx_path}")


def run_dsp_pickup(releases, playlists, output_path=None, grouping='platform'):
    """
    Main DSP pickup workflow.
    Check releases against all playlists and generate report.
    """
    spotify_playlists = [p for p in playlists if p['platform'] == 'Spotify' and p['spotify_id']]
    deezer_playlists = [p for p in playlists if p['platform'] == 'Deezer' and p['deezer_id']]
    apple_playlists = [p for p in playlists if p['platform'] == 'Apple Music' and p.get('apple_music_id')]
    amazon_playlists = [p for p in playlists if 'Amazon' in p.get('platform', '') and p.get('amazon_music_id')]
    claro_playlists = [p for p in playlists if 'Claro' in p.get('platform', '') and p.get('claro_id')]
    ytmusic_playlists = [p for p in playlists if 'YouTube' in p.get('platform', '') and p.get('ytmusic_id')]
    other_playlists = [p for p in playlists if p not in spotify_playlists + deezer_playlists + apple_playlists + amazon_playlists + claro_playlists + ytmusic_playlists]

    print(f"\nPlaylists to check:")
    print(f"  Spotify:       {len(spotify_playlists)}")
    print(f"  Deezer:        {len(deezer_playlists)}")
    print(f"  Apple Music:   {len(apple_playlists)}")
    print(f"  Amazon Music:  {len(amazon_playlists)}")
    print(f"  Claro Música:  {len(claro_playlists)}")
    print(f"  YouTube Music: {len(ytmusic_playlists)}")
    if other_playlists:
        print(f"  Other:         {len(other_playlists)} (manual check needed)")
    print(f"\nReleases to check: {len(releases)}")
    
    # Cache playlist tracks and covers to avoid re-fetching
    # { id: (tracks, cover_url) }
    playlist_cache = {}

    # Results: { artist: { release_title: [playlist_matches] } }
    results = {}
    # Dedup: skip if same artist+track already matched in same playlist
    _seen_matches = set()

    # Check Spotify playlists
    for pl in spotify_playlists:
        pl_id = pl['spotify_id']
        print(f"\n  Checking: {pl['name']} ({pl['country']}) [{pl['followers']} followers]")

        if pl_id not in playlist_cache:
            tracks, cover_url = get_spotify_playlist_tracks(pl_id)
            playlist_cache[pl_id] = (tracks, cover_url)
            print(f"    → {len(tracks)} tracks loaded")
        else:
            tracks, cover_url = playlist_cache[pl_id]

        for release in releases:
            match = check_release_in_playlist(release, tracks)
            if match:
                artist = release['artist']
                title = release['title']
                dedup_key = (artist, pl['name'], match.get('playlist_track', ''))
                if dedup_key in _seen_matches:
                    continue
                _seen_matches.add(dedup_key)

                if artist not in results:
                    results[artist] = {}
                if title not in results[artist]:
                    results[artist][title] = []

                results[artist][title].append({
                    'playlist_name': pl['name'],
                    'playlist_country': pl['country'],
                    'playlist_followers': pl['followers'],
                    'playlist_link': pl.get('link', ''),
                    'playlist_cover_url': cover_url,
                    'platform': 'Spotify',
                    **match,
                })
                print(f"    ✓ MATCH: {artist} — {title} (position #{match['position']})")

    # Check Deezer playlists
    for pl in deezer_playlists:
        pl_id = pl['deezer_id']
        print(f"\n  Checking: {pl['name']} ({pl['country']}) [Deezer]")

        tracks, cover_url = get_deezer_playlist_tracks(pl_id)
        print(f"    → {len(tracks)} tracks loaded")

        for release in releases:
            match = check_release_in_playlist(release, tracks)
            if match:
                artist = release['artist']
                title = release['title']
                dedup_key = (artist, pl['name'], match.get('playlist_track', ''))
                if dedup_key in _seen_matches:
                    continue
                _seen_matches.add(dedup_key)

                if artist not in results:
                    results[artist] = {}
                if title not in results[artist]:
                    results[artist][title] = []

                results[artist][title].append({
                    'playlist_name': pl['name'],
                    'playlist_country': pl['country'],
                    'playlist_followers': pl['followers'],
                    'playlist_link': pl.get('link', ''),
                    'playlist_cover_url': cover_url,
                    'platform': 'Deezer',
                    **match,
                })
                print(f"    ✓ MATCH: {artist} — {title} (position #{match['position']})")

    # Check Apple Music playlists
    for pl in apple_playlists:
        pl_id = pl['apple_music_id']
        print(f"\n  Checking: {pl['name']} ({pl['country']}) [Apple Music]")

        if pl_id not in playlist_cache:
            tracks, cover_url = get_apple_music_playlist_tracks(pl_id)
            playlist_cache[pl_id] = (tracks, cover_url)
            print(f"    → {len(tracks)} tracks loaded")
        else:
            tracks, cover_url = playlist_cache[pl_id]
            print(f"    → {len(tracks)} tracks (cached)")

        for release in releases:
            match = check_release_in_playlist(release, tracks)
            if match:
                artist = release['artist']
                title = release['title']
                dedup_key = (artist, pl['name'], match.get('playlist_track', ''))
                if dedup_key in _seen_matches:
                    continue
                _seen_matches.add(dedup_key)

                if artist not in results:
                    results[artist] = {}
                if title not in results[artist]:
                    results[artist][title] = []

                results[artist][title].append({
                    'playlist_name': pl['name'],
                    'playlist_country': pl['country'],
                    'playlist_followers': pl['followers'],
                    'playlist_link': pl.get('link', ''),
                    'playlist_cover_url': cover_url,
                    'platform': 'Apple Music',
                    **match,
                })
                print(f"    ✓ MATCH: {artist} — {title} (position #{match['position']})")

    # Check Amazon Music playlists
    for pl in amazon_playlists:
        pl_id = pl['amazon_music_id']
        domain = pl.get('amazon_music_domain', 'com.mx')
        print(f"\n  Checking: {pl['name']} ({pl['country']}) [Amazon Music]")

        if pl_id not in playlist_cache:
            tracks, cover_url = get_amazon_music_playlist_tracks(pl_id, domain)
            playlist_cache[pl_id] = (tracks, cover_url)
            print(f"    → {len(tracks)} tracks loaded")
        else:
            tracks, cover_url = playlist_cache[pl_id]
            print(f"    → {len(tracks)} tracks (cached)")

        for release in releases:
            match = check_release_in_playlist(release, tracks)
            if match:
                artist = release['artist']
                title = release['title']
                dedup_key = (artist, pl['name'], match.get('playlist_track', ''))
                if dedup_key in _seen_matches:
                    continue
                _seen_matches.add(dedup_key)

                if artist not in results:
                    results[artist] = {}
                if title not in results[artist]:
                    results[artist][title] = []

                results[artist][title].append({
                    'playlist_name': pl['name'],
                    'playlist_country': pl['country'],
                    'playlist_followers': pl['followers'],
                    'playlist_link': pl.get('link', ''),
                    'playlist_cover_url': cover_url,
                    'platform': 'Amazon Music',
                    **match,
                })
                print(f"    ✓ MATCH: {artist} — {title} (position #{match['position']})")

    # Check Claro Música playlists
    for pl in claro_playlists:
        pl_id = pl['claro_id']
        country = pl.get('claro_country', 'MX')
        print(f"\n  Checking: {pl['name']} ({pl['country']}) [Claro Música]")

        if pl_id not in playlist_cache:
            tracks, cover_url = get_claro_playlist_tracks(pl_id, country)
            playlist_cache[pl_id] = (tracks, cover_url)
            print(f"    → {len(tracks)} tracks loaded")
        else:
            tracks, cover_url = playlist_cache[pl_id]
            print(f"    → {len(tracks)} tracks (cached)")

        for release in releases:
            match = check_release_in_playlist(release, tracks)
            if match:
                artist = release['artist']
                title = release['title']
                dedup_key = (artist, pl['name'], match.get('playlist_track', ''))
                if dedup_key in _seen_matches:
                    continue
                _seen_matches.add(dedup_key)

                if artist not in results:
                    results[artist] = {}
                if title not in results[artist]:
                    results[artist][title] = []

                results[artist][title].append({
                    'playlist_name': pl['name'],
                    'playlist_country': pl['country'],
                    'playlist_followers': pl['followers'],
                    'playlist_link': pl.get('link', ''),
                    'playlist_cover_url': cover_url,
                    'platform': 'Claro Música',
                    **match,
                })
                print(f"    ✓ MATCH: {artist} — {title} (position #{match['position']})")

    # Check YouTube Music playlists
    for pl in ytmusic_playlists:
        pl_id = pl['ytmusic_id']
        print(f"\n  Checking: {pl['name']} ({pl['country']}) [YouTube Music]")

        if pl_id not in playlist_cache:
            tracks, cover_url = get_ytmusic_playlist_tracks(pl_id)
            playlist_cache[pl_id] = (tracks, cover_url)
            print(f"    → {len(tracks)} tracks loaded")
        else:
            tracks, cover_url = playlist_cache[pl_id]
            print(f"    → {len(tracks)} tracks (cached)")

        for release in releases:
            match = check_release_in_playlist(release, tracks)
            if match:
                artist = release['artist']
                title = release['title']
                dedup_key = (artist, pl['name'], match.get('playlist_track', ''))
                if dedup_key in _seen_matches:
                    continue
                _seen_matches.add(dedup_key)

                if artist not in results:
                    results[artist] = {}
                if title not in results[artist]:
                    results[artist][title] = []

                results[artist][title].append({
                    'playlist_name': pl['name'],
                    'playlist_country': pl['country'],
                    'playlist_followers': pl['followers'],
                    'playlist_link': pl.get('link', ''),
                    'playlist_cover_url': cover_url,
                    'platform': 'YouTube Music',
                    **match,
                })
                print(f"    ✓ MATCH: {artist} — {title} (position #{match['position']})")

    # Generate proof images for matches
    proof_image_paths = {}
    if results and output_path:
        proof_dir = str(Path(output_path).parent / 'dsp_proofs')
        print(f"\nGenerating proof images...")
        proof_image_paths = generate_proof_images(results, proof_dir)
        total_images = sum(len(imgs) for r in proof_image_paths.values() for imgs in r.values())
        print(f"  Generated {total_images} proof images in {proof_dir}")

        # Generate .docx report with embedded proof images
        docx_path = output_path.replace('.txt', '.docx')
        print(f"Generating .docx report...")
        generate_dsp_docx(results, proof_image_paths, docx_path, grouping=grouping)

    # Format output
    output_lines = ["DSP Pickup Report", "=" * 50, ""]

    # List all artists/releases that were checked
    checked_artists = {}
    for r in releases:
        a = r['artist']
        if a not in checked_artists:
            checked_artists[a] = []
        checked_artists[a].append(r['title'])

    output_lines.append(f"Releases checked ({len(releases)}):")
    for artist in sorted(checked_artists.keys()):
        titles = ', '.join(checked_artists[artist])
        output_lines.append(f"  • {artist} — {titles}")
    output_lines.append("")

    if not results:
        output_lines.append("No matches found in any checked playlists.")
    else:
        # Show which artists had no matches
        no_match_artists = [a for a in sorted(checked_artists.keys()) if a not in results]
        if no_match_artists:
            output_lines.append(f"No placements found for: {', '.join(no_match_artists)}")
            output_lines.append("")

        for artist in sorted(results.keys()):
            output_lines.append(f"\n{artist}")
            output_lines.append("-" * len(artist))

            for title, matches in sorted(results[artist].items()):
                output_lines.append(f"  {title}:")
                for m in sorted(matches, key=lambda x: x.get('playlist_name', '')):
                    followers = f" ({m['playlist_followers']} followers)" if m.get('playlist_followers') else ""
                    link = f"\n      {m['playlist_link']}" if m.get('playlist_link') else ""
                    output_lines.append(
                        f"    • {m['playlist_name']} [{m['platform']}] — "
                        f"{m['playlist_country']}{followers} — Position #{m['position']}{link}"
                    )
    
    # Note manual checks needed (only for platforms we can't scrape)
    if other_playlists:
        output_lines.append("\n" + "=" * 50)
        output_lines.append("MANUAL CHECK NEEDED:")
        for pl in other_playlists:
            output_lines.append(f"  • {pl['name']} [{pl['platform']}] — {pl['country']}")
            if pl.get('link'):
                output_lines.append(f"    {pl['link']}")
    
    output_text = '\n'.join(output_lines)
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output_text)
        print(f"\nReport saved: {output_path}")
    else:
        print("\n" + output_text)
    
    # Also save raw results as JSON for further processing
    if output_path:
        json_path = output_path.replace('.txt', '.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        print(f"Raw data saved: {json_path}")
    
    return results


def _parse_release_date(date_str, ref_year=None):
    """Parse date strings like 'Feb 23', 'Mar 6', 'April 10' into datetime."""
    if ref_year is None:
        ref_year = datetime.now().year
    date_str = date_str.strip()
    # Try common formats: "Feb 23", "Mar 6", "April 10"
    for fmt in ('%b %d', '%B %d'):
        try:
            return datetime.strptime(f"{date_str} {ref_year}", f'{fmt} %Y')
        except ValueError:
            continue
    return None


def filter_releases_by_week(releases, target_date_str):
    """
    Filter releases to a specific week block from the release schedule.
    Uses the week_block field (derived from separator rows in the Google Sheet)
    to identify which releases belong together in a week.
    Falls back to date proximity if week_block data is missing.
    """
    # Parse target date
    if target_date_str.lower() == 'current':
        target = datetime.now()
    else:
        target = datetime.strptime(target_date_str, '%Y-%m-%d')

    # Group releases by week_block
    blocks = {}
    for r in releases:
        wb = r.get('week_block', -1)
        if wb not in blocks:
            blocks[wb] = []
        blocks[wb].append(r)

    # Find the block whose date range contains the target date
    best_block = None
    best_distance = None
    for wb, block_releases in blocks.items():
        # Parse all dates in this block
        dates = []
        for r in block_releases:
            d = _parse_release_date(r.get('date', ''), target.year)
            if d:
                dates.append(d)
        if not dates:
            continue

        block_start = min(dates)
        block_end = max(dates)

        # Target falls within this block's date range
        if block_start <= target <= block_end:
            return block_releases

        # Track closest block as fallback
        dist = min(abs((target - block_start).days), abs((target - block_end).days))
        if best_distance is None or dist < best_distance:
            best_distance = dist
            best_block = wb

    # Fallback: return the closest block (within 7 days)
    if best_block is not None and best_distance <= 7:
        return blocks[best_block]

    return []


def main():
    parser = argparse.ArgumentParser(
        description='DSP Pickup Tool — Check LATAM editorial playlists for artist releases',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python dsp_pickup.py --week current
  python dsp_pickup.py --week 2026-02-21
  python dsp_pickup.py --artist "Djo"
  python dsp_pickup.py --all --output dsp_report.txt

Environment Variables:
  PLAYLIST_DB_PATH       Path to playlist database CSV
  RELEASE_SCHEDULE_URL   Published Google Sheets URL for release schedule
        """
    )
    
    parser.add_argument('--artist', help='Check specific artist across all playlists')
    parser.add_argument('--week', help='Check releases from specific week (YYYY-MM-DD or "current")')
    parser.add_argument('--all', action='store_true', help='Check all releases from schedule')
    parser.add_argument('--output', help='Output file path')
    parser.add_argument('--playlist-db', help='Path to playlist database CSV')
    parser.add_argument('--release-schedule', help='Path or URL to release schedule')
    parser.add_argument('--spotify-only', action='store_true', help='Only check Spotify playlists')
    
    args = parser.parse_args()
    
    if not args.artist and not args.week and not args.all:
        parser.error('One of --artist, --week, or --all is required')
    
    # Load databases
    pl_path = args.playlist_db or PLAYLIST_DB_PATH
    print(f"Loading playlist database from {pl_path}...")
    playlists = load_playlist_database(pl_path)
    print(f"  Loaded {len(playlists)} playlists")
    
    if args.spotify_only:
        playlists = [p for p in playlists if p['platform'] == 'Spotify']
        print(f"  Filtered to {len(playlists)} Spotify playlists")
    
    schedule_source = args.release_schedule or RELEASE_SCHEDULE_URL
    print(f"Loading release schedule from {schedule_source}...")
    releases = load_release_schedule(schedule_source)
    print(f"  Loaded {len(releases)} releases")
    
    # Filter releases
    if args.artist:
        search_lower = args.artist.lower()
        releases = [r for r in releases if search_lower in r['artist'].lower() or r['artist'].lower() in search_lower]
        print(f"  Filtered to {len(releases)} releases for {args.artist}")
    elif args.week:
        releases = filter_releases_by_week(releases, args.week)
        print(f"  Filtered to {len(releases)} releases for week of {args.week}")
    
    if not releases:
        print("No releases found matching criteria.")
        sys.exit(0)
    
    run_dsp_pickup(releases, playlists, args.output)


if __name__ == '__main__':
    main()

#!/usr/bin/env python3
"""
Press Pickup Tool
=================
Searches for Spanish-language press coverage of artists, matches results
against the press description database, and generates a formatted report.

Usage:
  python press_pickup.py --artist "Djo" --days 28
  python press_pickup.py --artist "Djo" --days 7 --output djo_press.txt
  python press_pickup.py --all --days 7  # All artists from release schedule

Requirements:
  pip install requests

Setup:
  Press articles: No API key required (Google News RSS + DuckDuckGo).
  Web search (SearXNG, self-hosted): docker run -d --name searxng -p 8888:8080 searxng/searxng
  Additional search source (optional): export TAVILY_API_KEY="tvly-..."  # Free 1000/month
  DuckDuckGo: No API key needed (pip install duckduckgo_search)

  Optional (for auto-generating missing media descriptions):
     export GROQ_API_KEY="your-groq-api-key"  # Free, no billing
"""

import argparse
import csv
from dataclasses import dataclass, field
from html import unescape
from html.parser import HTMLParser
import json
import os
import re
import sys
import unicodedata
from datetime import datetime, timedelta
from pathlib import Path
from urllib.parse import quote_plus, unquote, urljoin, urlsplit

# Add parent dir to path for shared modules
sys.path.insert(0, str(Path(__file__).parent.parent))
from shared.database import load_press_database, match_url_to_media, extract_domain

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

PRESS_DB_PATH = os.environ.get(
    'PRESS_DB_PATH',
    str(Path(__file__).parent.parent / 'data' / 'press_database.csv')
)

SOCIAL_HANDLE_REGISTRY_PATH = str(
    Path(__file__).parent.parent / 'data' / 'social_handle_registry.json'
)

RELEASE_SCHEDULE_URL = os.environ.get(
    'RELEASE_SCHEDULE_URL',
    'https://docs.google.com/spreadsheets/d/e/2PACX-1vTSd9mhkVibb7AwXtsZjRgBfuRT9sLY_qhhu-rB_P35CX2vFk_fZw_f31AJyW84KrCzWLLMUcTzzgqU/pub?gid=497066221&single=true&output=csv'
)

def _normalize_for_matching(text):
    """Strip accents/diacritics for fuzzy comparison.

    'Björk' → 'Bjork', 'Ñengo Flow' → 'Nengo Flow', 'José' → 'Jose'.
    Applied to BOTH keyword and target text so accented and unaccented forms match.
    """
    nfkd = unicodedata.normalize('NFKD', text)
    return ''.join(c for c in nfkd if not unicodedata.combining(c))


def _make_keyword_patterns(keywords):
    """Build compiled regex patterns for word-boundary matching of keywords.

    Returns a list of compiled re.Pattern objects. Creates patterns for both
    the original keyword and its accent-normalized form (if different), so
    "Björk" matches both "Björk" and "Bjork" in text.

    Each keyword is treated as a full phrase — "Bad Bunny" becomes \\bBad Bunny\\b,
    requiring the full name to match. An article mentioning only "Bunny" won't match.
    """
    patterns = []
    for kw in keywords:
        patterns.append(re.compile(r'\b' + re.escape(kw) + r'\b', re.IGNORECASE))
        normalized = _normalize_for_matching(kw)
        if normalized != kw:
            patterns.append(re.compile(r'\b' + re.escape(normalized) + r'\b', re.IGNORECASE))
    return patterns


def _any_keyword_matches(patterns, text):
    """Check if any keyword pattern matches in the given text (word-boundary).

    Checks both the original text and its accent-normalized form, so
    "Björk" in a pattern matches "Bjork" in text and vice versa.
    """
    if any(p.search(text) for p in patterns):
        return True
    normalized = _normalize_for_matching(text)
    if normalized != text and any(p.search(normalized) for p in patterns):
        return True
    return False


# Map TLDs / URL patterns to countries for grouping (longer patterns first)
DOMAIN_TO_COUNTRY = {
    '.com.ar': 'ARGENTINA', '.gob.ar': 'ARGENTINA', '.ar': 'ARGENTINA',
    '.com.br': 'BRAZIL', '.org.br': 'BRAZIL', '.br': 'BRAZIL',
    '.cl': 'CHILE',
    '.com.co': 'COLOMBIA',
    '.com.mx': 'MEXICO', '.gob.mx': 'MEXICO', '.mx': 'MEXICO',
    '.com.pe': 'PERU', '.pe': 'PERU',
    '.com.ec': 'ECUADOR', '.ec': 'ECUADOR',
    '.com.uy': 'URUGUAY', '.uy': 'URUGUAY',
    '.com.ve': 'VENEZUELA', '.ve': 'VENEZUELA',
    '.com.pa': 'PANAMA', '.pa': 'PANAMA',
    '.cr': 'COSTA RICA',
    '.com.gt': 'GUATEMALA', '.gt': 'GUATEMALA',
    '.hn': 'HONDURAS',
    '.com.sv': 'EL SALVADOR', '.sv': 'EL SALVADOR',
    '.com.ni': 'NICARAGUA', '.ni': 'NICARAGUA',
    '.com.do': 'DOMINICAN REPUBLIC', '.do': 'DOMINICAN REPUBLIC',
    '.com.py': 'PARAGUAY', '.py': 'PARAGUAY',
    '.com.bo': 'BOLIVIA', '.bo': 'BOLIVIA',
    '.cu': 'CUBA',
}

# TLD suffixes that confirm a domain is LATAM
LATAM_TLD_SUFFIXES = {
    '.ar', '.br', '.cl', '.co', '.mx', '.pe', '.ec', '.uy', '.ve', '.pa',
    '.cr', '.gt', '.hn', '.sv', '.ni', '.do', '.py', '.bo', '.cu',
}

# Social media domains to include (not skip, not press — treated specially)
SOCIAL_MEDIA_DOMAINS = {
    'instagram.com': 'Instagram',
    'facebook.com': 'Facebook',
    'x.com': 'X (Twitter)',
    'twitter.com': 'X (Twitter)',
}

# Map social media domains to url_type keys for grouping
_SOCIAL_DOMAIN_TYPE = {
    'instagram.com': 'instagram',
    'facebook.com': 'facebook',
    'x.com': 'x',
    'twitter.com': 'x',
}

# Display labels for URL types (None = no prefix)
_URL_TYPE_LABELS = {
    'article': None,
    'instagram': 'Instagram',
    'facebook': 'Facebook',
    'x': 'X',
}

# Domains to skip (streaming platforms, ticket sales, lyrics, etc.)
SKIP_DOMAINS = {
    'spotify.com', 'apple.com', 'music.apple.com', 'youtube.com',
    'youtu.be', 'tiktok.com', 'wikipedia.org', 'wikidata.org',
    'amazon.com', 'amazon.com.mx', 'amazon.com.br', 'amazon.com.ar',
    'amazon.com.co', 'amazon.com.pe', 'amazon.com.cl',
    'deezer.com', 'soundcloud.com', 'genius.com',
    'letras.com', 'letras.mus.br', 'musica.com', 'last.fm', 'discogs.com',
    'bandcamp.com', 'shazam.com', 'setlist.fm', 'songkick.com',
    'ticketmaster.com', 'ticketmaster.com.mx', 'stubhub.com', 'seatgeek.com',
    'tidal.com', 'qobuz.com', 'pandora.com', 'napster.com', 'anghami.com',
    'boomplay.com', 'claromusica.com', 'resso.com', 'jiosaavn.com',
}

# URL path segments that indicate non-article content
NON_PRESS_PATHS = (
    '/product/', '/shop/', '/cart/', '/store/', '/merch/',
    '/buy/', '/order/', '/checkout/', '/tienda/',
    '/tag/', '/tags/', '/categoria/', '/category/', '/categories/',
    '/autor/', '/author/', '/etiqueta/', '/label/',
    '/search/', '/buscar/', '/page/', '/perfil/',
    '/artist/', '/artists/', '/album/', '/albums/', '/track/', '/tracks/', '/playlist/',
)

# URL patterns that confirm LATAM relevance for .com domain articles
# Used to disambiguate multi-regional outlets (e.g. Rolling Stone US vs Mexico)
LATAM_URL_INDICATORS = re.compile(
    r'/(?:es|mx|ar|br|cl|co|pe|ec|uy|ve|pt|latam|latin|latino|latina|'
    r'en-espanol|espanol|spanish|america-latina|latinoamerica|musica)/'
    , re.IGNORECASE
)
# Spanish/Portuguese URL slug words that suggest LATAM content
# Only words that are unambiguously Spanish/Portuguese (not shared with English)
# Excluded: album, cultura, noticias — these appear in English URLs too
LATAM_SLUG_WORDS = re.compile(
    r'[-/](?:musica|música|artista|cantante|banda|disco|cancion|canción|estreno|lanzamiento|'
    r'concierto|gira|entrevista|espectaculos|espectáculos|entretenimiento|'
    r'reseña|resena|lançamento|lancamento)[-/]'
    , re.IGNORECASE
)


def _is_generic_com_domain(domain: str) -> bool:
    """Check if domain is a generic .com (not a country-specific TLD like .com.mx)."""
    if not domain:
        return False
    # Country-specific TLDs: .com.mx, .com.ar, .com.br, .co, .cl, etc.
    for suffix in LATAM_TLD_SUFFIXES:
        if domain.endswith(suffix):
            return False
    # Also check compound TLDs like .com.ar
    for tld in DOMAIN_TO_COUNTRY:
        if domain.endswith(tld):
            return False
    return domain.endswith('.com') or domain.endswith('.org') or domain.endswith('.net')


def _has_latam_url_indicators(url: str) -> bool:
    """Check if a URL contains LATAM language/region indicators in its path."""
    return bool(LATAM_URL_INDICATORS.search(url) or LATAM_SLUG_WORDS.search(url))


# Tracking query params to strip during URL normalization
_TRACKING_PARAMS = re.compile(
    r'[?&](?:utm_\w+|ref|fbclid|gclid|source|mc_cid|mc_eid|__twitter_impression'
    r'|_ga|ncid|ocid|dicbo|cmpid|cmp)=[^&#]*',
    re.IGNORECASE,
)

_NON_PRESS_SEGMENTS = {segment.strip('/') for segment in NON_PRESS_PATHS if segment.strip('/')}


def _is_skipped_domain(domain: str) -> bool:
    """Return True when a domain belongs to a skipped host family.

    This blocks both exact/subdomain matches like `open.spotify.com` and
    country-TLD variants like `music.amazon.com.mx` for skipped platforms.
    """
    if not domain:
        return False
    domain = domain.lower().strip()
    for skip in SKIP_DOMAINS:
        if (
            domain == skip
            or domain.endswith('.' + skip)
            or domain.startswith(skip + '.')
            or f'.{skip}.' in domain
        ):
            return True
    return False


def _is_non_press_url(url: str) -> bool:
    """Return True when a URL path contains a blocked non-article segment."""
    if not url:
        return False
    path = unquote(urlsplit(url).path or '').lower()
    segments = [segment for segment in path.split('/') if segment]
    return any(segment in _NON_PRESS_SEGMENTS for segment in segments)


def _normalize_url(url: str) -> str:
    """Normalize a URL for deduplication: strip tracking params, amp, www, etc."""
    u = url.strip()
    # http → https
    if u.startswith('http://'):
        u = 'https://' + u[7:]
    # Strip www.
    u = re.sub(r'^(https://)www\.', r'\1', u)
    # Strip trailing /amp/ or /amp
    u = re.sub(r'/amp/?$', '', u)
    # Strip query-string tracking params
    u = _TRACKING_PARAMS.sub('', u)
    # Clean up orphaned ? or & at end
    u = re.sub(r'[?&]+$', '', u)
    # If query string starts with & instead of ? after stripping, fix it
    u = re.sub(r'\?&', '?', u)
    # Strip trailing slash
    u = u.rstrip('/')
    return u


# Instagram paths that are content, not profile handles
_INSTAGRAM_NON_HANDLE = {'p', 'reel', 'reels', 'stories', 'tv', 'explore', 'accounts', 'direct'}
# Facebook paths that are not page names
_FACEBOOK_NON_HANDLE = {
    'sharer', 'sharer.php', 'share', 'dialog', 'plugins', 'login',
    'watch', 'groups', 'events', 'marketplace', 'gaming', 'help',
    'permalink.php', 'story.php', 'photo.php', 'video', 'pg',
    'profile.php', 'pages', 'policies', 'privacy',
}
# X/Twitter paths that are not handles
_TWITTER_NON_HANDLE = {'intent', 'share', 'search', 'explore', 'home', 'hashtag', 'i', 'settings'}


def _extract_social_handle(url: str) -> tuple[str, str] | None:
    """Extract (platform, handle) from a social media URL.

    Returns None if the handle can't be determined (e.g. instagram.com/p/... posts).
    """
    match = re.match(r'https?://(?:www\.)?([\w.]+)/([^/?&#]+)', url)
    if not match:
        return None
    host, first_segment = match.group(1).lower(), match.group(2).lower()
    # Strip @ prefix if present
    first_segment = first_segment.lstrip('@')
    if not first_segment or len(first_segment) < 2:
        return None

    if 'instagram.com' in host:
        if first_segment in _INSTAGRAM_NON_HANDLE:
            return None  # Post/reel URL — handle not determinable
        return ('instagram', first_segment)
    elif 'facebook.com' in host:
        if first_segment in _FACEBOOK_NON_HANDLE:
            return None
        return ('facebook', first_segment)
    elif 'twitter.com' in host or host == 'x.com':
        if first_segment in _TWITTER_NON_HANDLE:
            return None
        return ('twitter', first_segment)
    return None


def _handle_matches_artist_keywords(handle: str, keywords: list[str]) -> bool:
    """Conservatively match artist keywords against a social handle.

    We only treat a handle as the artist's own account when the normalized handle
    matches the full normalized artist keyword, or a very common prefixed/suffixed
    variant of it. This avoids broad substring false positives like `badgirls`
    matching the artist name `Bad`.
    """
    if not handle:
        return False

    handle_norm = _normalize_for_matching(handle.lower())
    handle_parts = [part for part in re.split(r'[^a-z0-9]+', handle_norm) if part]
    handle_compact = ''.join(handle_parts)
    if not handle_compact:
        return False

    common_prefixes = ('iam', 'soy', 'its', 'official', 'oficial')
    common_suffixes = ('official', 'oficial', 'music', 'musica', 'hq', 'tv', 'mx', 'br', 'ar', 'latam')

    for keyword in keywords:
        keyword_norm = _normalize_for_matching((keyword or '').lower())
        keyword_parts = [part for part in re.split(r'[^a-z0-9]+', keyword_norm) if part]
        keyword_compact = ''.join(keyword_parts)
        if not keyword_compact:
            continue

        if handle_compact == keyword_compact:
            return True

        if handle_parts == keyword_parts:
            return True

        if len(keyword_compact) >= 4:
            if any(handle_compact == prefix + keyword_compact for prefix in common_prefixes):
                return True
            if any(handle_compact == keyword_compact + suffix for suffix in common_suffixes):
                return True

    return False


def _keyword_match_type(patterns, title: str, snippet: str = '') -> str | None:
    """Classify where an artist keyword matched within an article candidate."""
    if title and _any_keyword_matches(patterns, title):
        return 'title'
    if snippet and _any_keyword_matches(patterns, snippet):
        return 'snippet'
    return None


_HTML_TAG_RE = re.compile(r'<[^>]+>')


def _strip_html_text(text: str) -> str:
    """Collapse basic HTML markup to text for matching and snippets."""
    return unescape(_HTML_TAG_RE.sub('', text or ''))


@dataclass(frozen=True)
class OutletAdapterSpec:
    adapter_id: str
    pattern_type: str
    outlet_name: str
    country: str
    description: str
    domain: str
    website: str
    search_url_template: str | None = None
    category_urls: tuple[str, ...] = field(default_factory=tuple)
    wp_api_url: str | None = None


_OUTLET_ADAPTER_SPECS = (
    OutletAdapterSpec(
        adapter_id='billboard-br',
        pattern_type='wordpress',
        outlet_name='BILLBOARD',
        country='BRAZIL',
        description='National edition of the international magazine, Billboard. Social media: 600K  Circulation: 40K copies.',
        domain='billboard.com.br',
        website='https://billboard.com.br',
        search_url_template='https://billboard.com.br/?s={query}',
        category_urls=('https://billboard.com.br/',),
        wp_api_url='https://billboard.com.br/wp-json/wp/v2/posts',
    ),
    OutletAdapterSpec(
        adapter_id='popline',
        pattern_type='wordpress',
        outlet_name='PORTAL POPLINE',
        country='BRAZIL',
        description='Portal focused on music news, videos, performances and everything regarding pop music. Social Media: 1M',
        domain='portalpopline.com.br',
        website='https://portalpopline.com.br',
        search_url_template='https://portalpopline.com.br/?s={query}',
        category_urls=('https://portalpopline.com.br/',),
        wp_api_url='https://portalpopline.com.br/wp-json/wp/v2/posts',
    ),
    OutletAdapterSpec(
        adapter_id='rolling-stone-mx',
        pattern_type='wordpress',
        outlet_name='Rolling Stone',
        country='MÉXICO',
        description='National edition of the renowned American music magazine. They also focus on music, film, videos, etc. Printed edition 70K Monthly Nationwide. Social Media: 2M',
        domain='rollingstone.com.mx',
        website='https://www.rollingstone.com.mx',
        search_url_template='https://www.rollingstone.com.mx/?s={query}',
        category_urls=('https://www.rollingstone.com.mx/musica/',),
        wp_api_url='https://www.rollingstone.com.mx/wp-json/wp/v2/posts',
    ),
    OutletAdapterSpec(
        adapter_id='elpais-uy',
        pattern_type='html',
        outlet_name='El Pais',
        country='URUGUAY',
        description='Biggest newspaper in Uruguay, founded in 1976, with one of the most prestigious music sections in the country. Social Media: 1.8M',
        domain='elpais.com.uy',
        website='https://www.elpais.com.uy',
        search_url_template='https://www.elpais.com.uy/busqueda?q={query}',
        category_urls=('https://www.elpais.com.uy/tvshow/musica',),
    ),
    OutletAdapterSpec(
        adapter_id='rpp',
        pattern_type='html',
        outlet_name='RPP Noticias',
        country='PERU',
        description='News from Peru and the world. The best of national and international news, sports and entertainment. Social Media: 6.9M',
        domain='rpp.pe',
        website='https://rpp.pe',
        search_url_template='https://rpp.pe/buscar?q={query}',
        category_urls=('https://rpp.pe/musica',),
    ),
    OutletAdapterSpec(
        adapter_id='biobio',
        pattern_type='html',
        outlet_name='BioBio Chile',
        country='CHILE',
        description='One of the largest independent media companies with a focus on national and international news. Social Media: 2.3M',
        domain='biobiochile.cl',
        website='https://www.biobiochile.cl',
        search_url_template='https://www.biobiochile.cl/buscar?q={query}',
        category_urls=('https://www.biobiochile.cl/lista/artes-y-cultura/musica',),
    ),
    OutletAdapterSpec(
        adapter_id='expreso',
        pattern_type='html',
        outlet_name='Diario Expreso',
        country='ECUADOR',
        description='Ecuadorian Newspaper reporting national and international news. Social Media: 102K',
        domain='expreso.ec',
        website='https://www.expreso.ec',
        search_url_template='https://www.expreso.ec/buscar?q={query}',
        category_urls=('https://www.expreso.ec/expresiones/ocio/',),
    ),
)


def _get_outlet_adapter_specs() -> tuple[OutletAdapterSpec, ...]:
    """Return the first adapter target set for the LATAM outlet framework."""
    return _OUTLET_ADAPTER_SPECS


def _build_adapter_session():
    """Create a shared HTTP session for outlet adapter fetching."""
    import requests

    session = requests.Session()
    session.headers.update({
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
    })
    return session


class _AdapterLinkParser(HTMLParser):
    """Extract anchor links with visible text from listing pages."""

    def __init__(self):
        super().__init__()
        self.links = []
        self._href = None
        self._title_attr = ''
        self._text_parts = []

    def handle_starttag(self, tag, attrs):
        if tag != 'a':
            return
        attr_map = dict(attrs)
        href = attr_map.get('href')
        if not href:
            return
        self._href = href
        self._title_attr = (attr_map.get('title') or '').strip()
        self._text_parts = []

    def handle_data(self, data):
        if self._href is not None and data.strip():
            self._text_parts.append(data.strip())

    def handle_endtag(self, tag):
        if tag != 'a' or self._href is None:
            return
        text = ' '.join(self._text_parts).strip() or self._title_attr
        self.links.append((self._href, text))
        self._href = None
        self._title_attr = ''
        self._text_parts = []


def _adapter_domain_matches(domain: str, expected_domain: str) -> bool:
    """Check whether a parsed domain belongs to the outlet's expected site."""
    if not domain or not expected_domain:
        return False
    return domain == expected_domain or domain.endswith('.' + expected_domain)


def _adapter_result(spec: OutletAdapterSpec, link: str, title: str, snippet: str, match_type: str) -> dict:
    """Build a preclassified adapter result entry."""
    return {
        'title': title.strip(),
        'link': link,
        'snippet': snippet.strip()[:300],
        'domain': extract_domain(link) or spec.domain,
        'source': spec.outlet_name,
        'feed_country': spec.country,
        'feed_description': spec.description,
        'feed_media_name': spec.outlet_name,
        '_keyword_match': match_type,
        '_source': 'adapter',
    }


def _extract_listing_links(html: str, base_url: str, expected_domain: str) -> list[tuple[str, str]]:
    """Extract likely article links and titles from a listing page."""
    parser = _AdapterLinkParser()
    parser.feed(html or '')

    seen = set()
    links = []
    for href, title in parser.links:
        if not title or len(title.strip()) < 16:
            continue
        abs_url = urljoin(base_url, href)
        domain = extract_domain(abs_url) or ''
        if not _adapter_domain_matches(domain, expected_domain):
            continue
        if _is_non_press_url(abs_url):
            continue
        norm = _normalize_url(abs_url)
        if norm in seen:
            continue
        seen.add(norm)
        links.append((abs_url, re.sub(r'\s+', ' ', title).strip()))
    return links


def _run_listing_surface(session, spec: OutletAdapterSpec, surface_url: str, kw_patterns, require_title_match=True) -> list[dict]:
    """Fetch an outlet listing/search surface and extract matching article links."""
    try:
        resp = session.get(surface_url, timeout=8)
        if resp.status_code != 200:
            return []
    except Exception:
        return []

    hits = []
    for link, title in _extract_listing_links(resp.text, surface_url, spec.domain):
        match_type = _keyword_match_type(kw_patterns, title, '')
        if require_title_match and match_type != 'title':
            continue
        if not match_type:
            continue
        hits.append(_adapter_result(spec, link, title, '', match_type))
    return hits


def _run_wordpress_adapter(session, spec: OutletAdapterSpec, keywords: list[str], kw_patterns, cutoff) -> list[dict]:
    """Search a WordPress outlet directly via REST API, with HTML fallbacks."""
    hits = []
    seen = set()
    cutoff_iso = cutoff.strftime('%Y-%m-%dT%H:%M:%S')

    if spec.wp_api_url:
        for keyword in keywords:
            try:
                resp = session.get(
                    spec.wp_api_url,
                    params={'search': keyword, 'per_page': 10, 'after': cutoff_iso},
                    timeout=8,
                )
                if resp.status_code != 200:
                    continue
                posts = resp.json()
                if not isinstance(posts, list):
                    continue
            except Exception:
                continue

            for post in posts:
                title = _strip_html_text(post.get('title', {}).get('rendered', ''))
                excerpt = _strip_html_text(post.get('excerpt', {}).get('rendered', ''))
                link = post.get('link', '')
                if not link:
                    continue
                domain = extract_domain(link) or ''
                if not _adapter_domain_matches(domain, spec.domain):
                    continue
                if _is_non_press_url(link):
                    continue
                match_type = _keyword_match_type(kw_patterns, title, excerpt)
                if not match_type:
                    continue
                norm = _normalize_url(link)
                if norm in seen:
                    continue
                seen.add(norm)
                hits.append(_adapter_result(spec, link, title, excerpt, match_type))

    # Fallback to the public search page and category listings for freshness/index lag.
    for keyword in keywords:
        if spec.search_url_template:
            search_url = spec.search_url_template.format(query=quote_plus(keyword))
            for result in _run_listing_surface(session, spec, search_url, kw_patterns, require_title_match=True):
                norm = _normalize_url(result['link'])
                if norm not in seen:
                    seen.add(norm)
                    hits.append(result)

    for category_url in spec.category_urls:
        for result in _run_listing_surface(session, spec, category_url, kw_patterns, require_title_match=True):
            norm = _normalize_url(result['link'])
            if norm not in seen:
                seen.add(norm)
                hits.append(result)

    return hits


def _run_html_adapter(session, spec: OutletAdapterSpec, keywords: list[str], kw_patterns) -> list[dict]:
    """Search a custom outlet using direct search/category listing pages."""
    hits = []
    seen = set()

    for keyword in keywords:
        if spec.search_url_template:
            search_url = spec.search_url_template.format(query=quote_plus(keyword))
            for result in _run_listing_surface(session, spec, search_url, kw_patterns, require_title_match=True):
                norm = _normalize_url(result['link'])
                if norm not in seen:
                    seen.add(norm)
                    hits.append(result)

    for category_url in spec.category_urls:
        for result in _run_listing_surface(session, spec, category_url, kw_patterns, require_title_match=True):
            norm = _normalize_url(result['link'])
            if norm not in seen:
                seen.add(norm)
                hits.append(result)

    return hits


def scan_outlet_adapters(artist_keywords, days=28, cutoff=None, adapter_specs=None):
    """Run deterministic outlet-specific retrieval adapters for priority LATAM outlets."""
    import time as _time

    if cutoff is None:
        cutoff = datetime.now().astimezone() - timedelta(days=days)
    specs = tuple(_get_outlet_adapter_specs() if adapter_specs is None else adapter_specs)
    if not specs:
        return []

    kw_patterns = _make_keyword_patterns(artist_keywords)
    session = _build_adapter_session()
    all_hits = []
    seen = set()
    start = _time.time()

    try:
        for spec in specs:
            if spec.pattern_type == 'wordpress':
                hits = _run_wordpress_adapter(session, spec, artist_keywords, kw_patterns, cutoff)
            else:
                hits = _run_html_adapter(session, spec, artist_keywords, kw_patterns)

            for hit in hits:
                norm = _normalize_url(hit['link'])
                if norm not in seen:
                    seen.add(norm)
                    all_hits.append(hit)
    finally:
        try:
            session.close()
        except Exception:
            pass

    elapsed = _time.time() - start
    print(f"  Scanned {len(specs)} outlet adapters in {elapsed:.1f}s → found {len(all_hits)} articles")
    return all_hits


def google_news_rss(query, gl='MX', hl='es-419', max_results=50, days=None, cutoff=None, decode=True):
    """
    Search Google News via free RSS feed. No API key required.

    Returns list of { title, link, snippet, domain, source }.
    If decode=True, links are decoded from Google News redirects to actual article URLs.
    If decode=False, returns raw google_link in the 'link' field for later batch decoding.
    """
    import requests
    import xml.etree.ElementTree as ET
    from email.utils import parsedate_to_datetime
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from googlenewsdecoder import new_decoderv1

    ceid = f'{gl}:{hl.split("-")[0]}'
    rss_url = f'https://news.google.com/rss/search?q={requests.utils.quote(query)}&hl={hl}&gl={gl}&ceid={ceid}'

    try:
        resp = requests.get(rss_url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }, timeout=15)
        resp.raise_for_status()
        root = ET.fromstring(resp.content)
    except Exception as e:
        print(f"    RSS fetch failed: {e}")
        return []

    # Date cutoff for filtering
    if cutoff is None and days:
        cutoff = datetime.now().astimezone() - timedelta(days=days)

    # ── Phase 1: Parse RSS items ──────────────────────────────────────────
    pending = []  # (index, title, google_link, source_name, snippet)
    items = root.findall('.//item')[:max_results]

    for item in items:
        title_el = item.find('title')
        link_el = item.find('link')
        source_el = item.find('source')
        desc_el = item.find('description')
        pub_el = item.find('pubDate')

        title = title_el.text if title_el is not None else ''
        google_link = link_el.text if link_el is not None else ''
        source_name = source_el.text if source_el is not None else ''
        snippet = desc_el.text if desc_el is not None else ''

        # Filter by date if cutoff is set
        if cutoff and pub_el is not None and pub_el.text:
            try:
                pub_dt = parsedate_to_datetime(pub_el.text)
                if pub_dt < cutoff:
                    continue
            except Exception:
                pass

        # Strip source name suffix from title
        if source_name and title.endswith(f' - {source_name}'):
            title = title[: -len(f' - {source_name}')]

        pending.append((len(pending), title, google_link, source_name, snippet))

    if not pending:
        return []

    if not decode:
        # Return raw items for later batch decoding
        return [
            {
                'title': p[1],
                'link': p[2],
                'snippet': p[4],
                'domain': 'news.google.com',
                'source': p[3],
                '_source_name_match': p[3],
            }
            for p in pending
        ]

    # ── Phase 2: Decode all Google News URLs concurrently ─────────────────
    DECODE_TIMEOUT = 20  # seconds total for all URL decodes

    def _decode(google_link):
        try:
            decoded = new_decoderv1(google_link)
            if decoded.get('status'):
                return decoded['decoded_url']
        except Exception:
            pass
        # Fallback: follow the redirect chain directly
        try:
            head_resp = requests.head(google_link, allow_redirects=True, timeout=5,
                                      headers={'User-Agent': 'Mozilla/5.0'})
            if head_resp.url and 'news.google.com' not in head_resp.url:
                return head_resp.url
        except Exception:
            pass
        return google_link

    decoded_urls = {}  # index → decoded_url
    executor = ThreadPoolExecutor(max_workers=40)
    future_map = {
        executor.submit(_decode, p[2]): p[0] for p in pending
    }
    decode_stats = {'decoded': 0, 'fallback': 0, 'name_matched': 0, 'failed': 0}
    try:
        for future in as_completed(future_map, timeout=DECODE_TIMEOUT):
            idx = future_map[future]
            try:
                url = future.result(timeout=1)
                decoded_urls[idx] = url
                if 'news.google.com' not in url:
                    decode_stats['decoded'] += 1
                else:
                    decode_stats['failed'] += 1
            except Exception:
                decoded_urls[idx] = pending[idx][2]
                decode_stats['failed'] += 1
    except TimeoutError:
        pass  # global timeout hit — use raw google links for remaining
    finally:
        executor.shutdown(wait=False, cancel_futures=True)

    # Fill in any undecoded URLs with the original google links
    for idx, _title, google_link, _src, _snip in pending:
        if idx not in decoded_urls:
            decoded_urls[idx] = google_link
            decode_stats['failed'] += 1

    # Phase 2b: For URLs still pointing to news.google.com, try source-name
    # matching against the press database to preserve the hit
    _source_name_matches = {}  # index → source_name (for results that failed URL decode)
    for idx, _title, google_link, source_name, _snip in pending:
        url = decoded_urls.get(idx, google_link)
        if 'news.google.com' in url and source_name:
            _source_name_matches[idx] = source_name
            decode_stats['failed'] -= 1
            decode_stats['name_matched'] += 1

    if any(v > 0 for v in decode_stats.values()):
        print(f"    URL decode: {decode_stats['decoded']} decoded, "
              f"{decode_stats['name_matched']} name-matched, "
              f"{decode_stats['failed']} dropped")

    # ── Phase 3: Build results ────────────────────────────────────────────
    results = []
    seen_urls = set()

    for idx, title, google_link, source_name, snippet in pending:
        link = decoded_urls.get(idx, google_link)

        # For unresolved Google News URLs, preserve the hit with source-name metadata
        # so downstream matching can use the outlet name instead of the URL
        is_name_matched = idx in _source_name_matches
        if 'news.google.com' in link and not is_name_matched:
            continue  # Truly unresolvable — no URL and no source name

        domain = extract_domain(link) or ''

        if _is_skipped_domain(domain):
            continue

        if _is_non_press_url(link):
            continue

        norm = _normalize_url(link)
        if norm not in seen_urls:
            seen_urls.add(norm)
            entry = {
                'title': title,
                'link': link,
                'snippet': snippet,
                'domain': domain,
                'source': source_name,
            }
            # Tag source-name-matched results so downstream can match by outlet name
            if is_name_matched:
                entry['_source_name_match'] = source_name
            results.append(entry)

    return results


def _search_searxng(query, num_results=30, categories='general'):
    """
    Search via local SearXNG instance (metasearch: Google, Bing, DuckDuckGo, etc).
    categories: 'general' for web results, 'news' for news articles.
    Returns list of { title, link, snippet, domain }.
    """
    import requests

    searxng_url = os.environ.get('SEARXNG_URL', 'http://localhost:8888')
    params = {
        'q': query,
        'format': 'json',
        'categories': categories,
        'language': 'es',
        'pageno': 1,
    }

    try:
        resp = requests.get(
            f'{searxng_url}/search',
            params=params,
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        print(f"    SearXNG search failed: {e}")
        return []

    raw_items = data.get('results', [])

    results = []
    for item in raw_items[:num_results]:
        link = item.get('url', '')
        domain = extract_domain(link) or ''

        if _is_skipped_domain(domain):
            continue

        if _is_non_press_url(link):
            continue

        results.append({
            'title': item.get('title', ''),
            'link': link,
            'snippet': item.get('content', ''),
            'domain': domain,
        })

    return results


def detect_country_from_url(url):
    """Detect country from URL's TLD. Returns None for non-LATAM domains."""
    domain = extract_domain(url) or ''
    # Check specific TLDs (longer first to avoid false matches)
    for tld, country in sorted(DOMAIN_TO_COUNTRY.items(), key=lambda x: -len(x[0])):
        if domain.endswith(tld):
            return country
    return None


def is_latam_domain(domain):
    """Check if a domain has a LATAM country-code TLD."""
    for suffix in LATAM_TLD_SUFFIXES:
        if domain.endswith(suffix):
            return True
    return False


def normalize_country(name):
    """Normalize country names to a consistent format."""
    mapping = {
        'MÉXICO': 'MEXICO',
        'PERÚ': 'PERU',
        'PANAMÁ': 'PANAMA',
        'REPÚBLICA DOMINICANA': 'DOMINICAN REPUBLIC',
    }
    return mapping.get(name, name)


def _extract_json_array(text: str):
    """Robustly extract a JSON array from LLM output.

    Handles markdown fences, leading prose, trailing commentary, etc.
    Returns the parsed list or None if no valid array is found.
    """
    # Strip markdown code fences
    cleaned = text.strip()
    if cleaned.startswith('```'):
        cleaned = cleaned.split('\n', 1)[1] if '\n' in cleaned else cleaned[3:]
    if cleaned.endswith('```'):
        cleaned = cleaned[:-3].strip()

    # Try direct parse first (fast path)
    try:
        result = json.loads(cleaned)
        if isinstance(result, list):
            return result
    except (json.JSONDecodeError, ValueError):
        pass

    # Find the outermost [...] boundaries
    start = cleaned.find('[')
    if start == -1:
        return None
    # Find matching closing bracket (handle nested arrays/objects)
    depth = 0
    end = None
    for i in range(start, len(cleaned)):
        if cleaned[i] == '[':
            depth += 1
        elif cleaned[i] == ']':
            depth -= 1
            if depth == 0:
                end = i
                break
    if end is None:
        return None

    try:
        result = json.loads(cleaned[start:end + 1])
        if isinstance(result, list):
            return result
    except (json.JSONDecodeError, ValueError):
        pass

    return None


def _groq_enrich_descriptions(outlets_to_enrich, log_fn=print):
    """
    Batch-enrich new outlet descriptions using Groq Llama 3.3 70B (free).
    Each item in outlets_to_enrich is a dict with media_name, url, snippet.
    Returns a dict mapping media_name → description.
    """
    import requests as _req

    api_key = os.environ.get('GROQ_API_KEY')
    if not api_key:
        return {}

    BATCH_SIZE = 10
    results = {}

    for batch_start in range(0, len(outlets_to_enrich), BATCH_SIZE):
        batch = outlets_to_enrich[batch_start:batch_start + BATCH_SIZE]
        log_fn(f"  Generating AI descriptions for {len(batch)} new outlet(s)...")

        outlet_list = []
        for i, o in enumerate(batch):
            outlet_list.append(
                f'{i+1}. Name: {o["media_name"]}\n'
                f'   URL: {o["url"]}\n'
                f'   Context: {o["snippet"][:150]}'
            )

        prompt = f"""You are helping a Latin American music marketing team describe new media outlets for their contact database.

For each outlet below, write a one-sentence description. Be factual and concise.

OUTLETS:
{chr(10).join(outlet_list)}

Respond with a JSON array (no markdown, no code fences). Each element must have:
- "index": the outlet number (1-based)
- "description": one sentence, like "Digital platform focused on music, cinema, shows and culture news." or "Brazilian electronic music blog covering DJs, festivals and new releases."

Output ONLY the JSON array, nothing else."""

        try:
            resp = _req.post(
                'https://api.groq.com/openai/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {api_key}',
                    'Content-Type': 'application/json',
                },
                json={
                    'model': 'llama-3.3-70b-versatile',
                    'messages': [{'role': 'user', 'content': prompt}],
                    'max_tokens': 1024,
                    'temperature': 0.3,
                },
                timeout=30,
            )

            if resp.status_code == 200:
                content = resp.json()['choices'][0]['message']['content'].strip()
                enrichments = _extract_json_array(content)
                if enrichments is not None:
                    for item in enrichments:
                        idx = item.get('index', 0) - 1
                        if 0 <= idx < len(batch) and item.get('description'):
                            results[batch[idx]['media_name']] = item['description']
                else:
                    log_fn(f"  AI description: could not parse JSON from response")
            else:
                log_fn(f"  Groq API error: {resp.status_code}")

        except Exception as e:
            log_fn(f"  AI description batch failed: {e}")

    return results


def _groq_filter_relevance(all_results, artist, keywords, releases=None, log_fn=print):
    """
    Use Groq Llama 3.3 70B to filter out false positives from search results.

    Two-tier approach:
    - Articles with the artist keyword in the TITLE are auto-confirmed (high confidence).
    - Articles where the keyword only appears in the snippet (not the title) are sent
      to Groq for AI review — these are often false positives from sidebar mentions,
      related-article widgets, or tag clouds.

    Returns a filtered list. If Groq is unavailable, snippet-only articles are kept
    and tagged as unreviewed so recall is preserved.
    """
    import requests as _req

    if not all_results:
        return all_results

    kw_patterns = _make_keyword_patterns(keywords)

    def _story_cluster_key(title: str) -> str:
        normalized = re.sub(r'\s+', ' ', _normalize_title(title or ''))
        if not normalized:
            return normalized
        trimmed = re.split(r'\s(?:\||-)\s', normalized, maxsplit=1)[0]
        return trimmed if len(trimmed) >= 24 else normalized

    def _cluster_results(results):
        clusters = {}
        cluster_order = []
        for result in results:
            key = _story_cluster_key(result.get('title', '')) or _normalize_url(result.get('link', ''))
            if key not in clusters:
                clusters[key] = []
                cluster_order.append(key)
            clusters[key].append(result)
        return [(key, clusters[key]) for key in cluster_order]

    # Split into title-confirmed vs snippet-only
    title_confirmed = []
    snippet_only = []
    for r in all_results:
        title = r.get('title') or ''
        if _any_keyword_matches(kw_patterns, title):
            title_confirmed.append(r)
        else:
            snippet_only.append(r)

    if not snippet_only:
        log_fn(f"  AI relevance filter: all {len(all_results)} articles have keyword in title")
        return all_results

    snippet_clusters = _cluster_results(snippet_only)
    snippet_representatives = [cluster[0] for _key, cluster in snippet_clusters]
    clustered_count = len(snippet_representatives)

    cluster_note = ''
    if clustered_count != len(snippet_only):
        cluster_note = f" across {clustered_count} story clusters"
    log_fn(f"  AI relevance filter: {len(title_confirmed)} title-confirmed, "
           f"{len(snippet_only)} snippet-only{cluster_note} → sending to Groq...")

    api_key = os.environ.get('GROQ_API_KEY')
    if not api_key:
        # No API key — keep snippet-only articles rather than silently dropping them.
        for result in snippet_only:
            result['_groq_unreviewed'] = True
        log_fn(f"  No GROQ_API_KEY — keeping {len(snippet_only)} snippet-only articles unreviewed")
        return all_results

    # Build release context string
    release_context = "No recent releases found."
    if releases:
        parts = []
        for r in releases[:3]:
            parts.append(f"\"{r['title']}\" ({r['format']}, {r['date']})")
        release_context = "Recent releases: " + ", ".join(parts)

    BATCH_SIZE = 15
    keep_flags = [False] * len(snippet_representatives)  # Default: reject snippet-only

    for batch_start in range(0, len(snippet_representatives), BATCH_SIZE):
        batch = snippet_representatives[batch_start:batch_start + BATCH_SIZE]

        article_list = []
        for i, r in enumerate(batch):
            title = (r.get('title') or '')[:120]
            domain = r.get('domain', '')
            article_list.append(f'{i+1}. Title: "{title}" | Domain: {domain}')

        prompt = f"""You are filtering press search results for a Latin American music marketing report.
Artist: {artist}
{release_context}

IMPORTANT: Judge ONLY on the article title. The article snippet is not shown because it may contain the artist name from unrelated page elements like sidebars, related articles, or tag clouds.

Based ONLY on the article title, is this article actually about {artist}?

An article IS about the artist if the title mentions them by name in a meaningful way (review, interview, feature, tour announcement, album release, concert coverage, etc.).

An article is NOT about the artist if:
- The title is about a different artist entirely
- The artist name appears only as a substring of another word (e.g. if the artist is "Metric", reject articles about "biometric", "parametric", "Biometrica", etc.)
- The title is a generic lineup/playlist listing many artists and {artist} isn't the focus
- The title is about an event from more than 1 year ago
- The title doesn't mention {artist} at all (it appeared only in page metadata)
- The FULL artist name is not referenced — an article mentioning only part of a multi-word name does not count (e.g. "Bunny" alone is NOT coverage for "Bad Bunny")

Special attention for ambiguous artist names:
- If the artist name is a common English or Spanish word (like Metric, Future, Calle, Sol, Air, Her, Low), be EXTRA strict: the article must clearly be about a musical artist, not using the word in its ordinary meaning
- If the artist name is very short (2-3 characters like LP, IU, AI), require strong musical context in the title
- Articles about "the metric system", "biometric data", "future trends", "solar energy", "air quality" etc. are NOT about the corresponding musical artists

Articles:
{chr(10).join(article_list)}

Respond with ONLY a JSON array of booleans, e.g. [true, false, true, ...]. No other text."""

        try:
            resp = _req.post(
                'https://api.groq.com/openai/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {api_key}',
                    'Content-Type': 'application/json',
                },
                json={
                    'model': 'llama-3.3-70b-versatile',
                    'messages': [{'role': 'user', 'content': prompt}],
                    'max_tokens': 256,
                    'temperature': 0.1,
                },
                timeout=15,
            )

            if resp.status_code == 200:
                content = resp.json()['choices'][0]['message']['content'].strip()
                verdicts = _extract_json_array(content)
                if isinstance(verdicts, list) and len(verdicts) == len(batch):
                    for i, keep in enumerate(verdicts):
                        if keep:
                            keep_flags[batch_start + i] = True
                # If length mismatch or parse failure, keep none (fail closed)

        except Exception as e:
            log_fn(f"  AI relevance filter batch failed: {e}")  # Fail closed

    groq_kept = []
    groq_removed = 0
    for (_cluster_key, cluster), keep in zip(snippet_clusters, keep_flags):
        if keep:
            for result in cluster:
                result['_story_cluster_size'] = len(cluster)
                groq_kept.append(result)
        else:
            groq_removed += len(cluster)

    filtered = title_confirmed + groq_kept
    total_removed = len(all_results) - len(filtered)

    if total_removed > 0:
        log_fn(f"  AI relevance filter: kept {len(filtered)}/{len(all_results)} articles "
               f"(removed {total_removed}: {groq_removed} snippet-only rejected by Groq)")
    else:
        log_fn(f"  AI relevance filter: all {len(all_results)} articles confirmed relevant")

    return filtered


def _serper_date_within(date_str, cutoff):
    """Check if a Serper date string falls within the cutoff.
    Handles Spanish relative dates ('hace 3 días', 'hace 1 semana') and
    absolute dates ('8 oct 2025', '11 feb 2026').
    Returns True if the date is recent enough or can't be parsed.
    """
    if not date_str:
        return True  # Include if no date

    date_str = date_str.lower().strip()

    # Relative dates: "hace X días/horas/semanas/meses"
    match = re.match(r'hace\s+(\d+)\s+(hora|día|semana|mes|min)', date_str)
    if match:
        num = int(match.group(1))
        unit = match.group(2)
        if unit.startswith('min') or unit.startswith('hora'):
            return True  # Always recent
        elif unit.startswith('día'):
            article_date = datetime.now().astimezone() - timedelta(days=num)
        elif unit.startswith('semana'):
            article_date = datetime.now().astimezone() - timedelta(weeks=num)
        elif unit.startswith('mes'):
            article_date = datetime.now().astimezone() - timedelta(days=num * 30)
        else:
            return True
        return article_date >= cutoff

    # Absolute dates: "8 oct 2025", "11 feb 2026"
    MONTHS_ES = {
        'ene': 1, 'feb': 2, 'mar': 3, 'abr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'ago': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dic': 12,
    }
    match = re.match(r'(\d{1,2})\s+(\w{3})\s+(\d{4})', date_str)
    if match:
        day, month_str, year = int(match.group(1)), match.group(2), int(match.group(3))
        month = MONTHS_ES.get(month_str[:3])
        if month:
            try:
                article_date = datetime(year, month, day).astimezone()
                return article_date >= cutoff
            except ValueError:
                pass

    return True  # Include if can't parse


def scan_outlet_feeds(artist_keywords, days=28, feed_registry_path=None, cutoff=None):
    """
    Scan known outlet RSS feeds and WordPress APIs for artist coverage.
    Returns results in the same format as the other search functions:
    list of { title, link, snippet, domain, source, feed_country, feed_description, feed_media_name }

    The extra feed_* fields carry the outlet metadata from the registry so that
    the caller doesn't need to re-match against the press database.
    """
    import time as _time
    import requests
    import feedparser
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from html import unescape

    registry_path = feed_registry_path or str(
        Path(__file__).parent.parent / 'data' / 'feed_registry.json'
    )

    if not os.path.exists(registry_path):
        print("  Feed registry not found — run discover_feeds.py for better results")
        return []

    with open(registry_path) as f:
        registry = json.load(f)

    outlets = registry.get('outlets', {})
    rss_outlets = []
    wp_outlets = []
    for domain, info in outlets.items():
        if info.get('feed_type') == 'rss' and info.get('feed_url'):
            rss_outlets.append((domain, info))
        elif info.get('feed_type') == 'wordpress' and info.get('wp_api_url'):
            wp_outlets.append((domain, info))

    if not rss_outlets and not wp_outlets:
        return []

    kw_patterns = _make_keyword_patterns(artist_keywords)
    if cutoff is None:
        cutoff = datetime.now().astimezone() - timedelta(days=days)
    cutoff_iso = cutoff.strftime('%Y-%m-%dT%H:%M:%S')

    FEED_TIMEOUT = 8
    FEED_UA = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'

    _html_tag_re = re.compile(r'<[^>]+>')

    def _strip_html(text):
        return unescape(_html_tag_re.sub('', text or ''))

    results = []
    results_lock = __import__('threading').Lock()
    failed_count = 0
    failed_lock = __import__('threading').Lock()

    # ── RSS scanning ──────────────────────────────────────────────────────

    def _scan_rss(domain_info):
        nonlocal failed_count
        domain, info = domain_info
        try:
            feed = feedparser.parse(
                info['feed_url'],
                agent=FEED_UA,
                request_headers={'User-Agent': FEED_UA},
            )
            # feedparser doesn't natively support timeout, so we set socket default
            if feed.bozo and not feed.entries:
                return []

            hits = []
            for entry in feed.entries:
                # Check publication date — skip entries with no parseable date
                published = entry.get('published_parsed') or entry.get('updated_parsed')
                if published:
                    from calendar import timegm
                    entry_ts = timegm(published)
                    entry_dt = datetime.fromtimestamp(entry_ts).astimezone()
                    if entry_dt < cutoff:
                        continue
                else:
                    # No date available — skip rather than risk old articles
                    continue

                title = entry.get('title', '')
                summary = _strip_html(entry.get('summary', '') or entry.get('description', ''))

                match_type = _keyword_match_type(kw_patterns, title, summary)
                if not match_type:
                    continue

                link = entry.get('link', '')
                entry_domain = extract_domain(link) if link else domain

                hits.append({
                    'title': title,
                    'link': link,
                    'snippet': summary[:300],
                    'domain': entry_domain or domain,
                    'source': info.get('name', domain),
                    'feed_country': info.get('country', ''),
                    'feed_description': info.get('description', ''),
                    'feed_media_name': info.get('name', domain),
                    '_keyword_match': match_type,
                })
            return hits

        except Exception:
            with failed_lock:
                nonlocal failed_count
                failed_count += 1
            return []

    # ── WordPress API scanning ────────────────────────────────────────────

    def _scan_wp(domain_info):
        nonlocal failed_count
        domain, info = domain_info
        hits = []
        session = requests.Session()
        session.headers.update({'User-Agent': FEED_UA})

        for kw in artist_keywords:
            try:
                url = info['wp_api_url']
                params = {
                    'search': kw,
                    'per_page': 10,
                    'after': cutoff_iso,
                }
                resp = session.get(url, params=params, timeout=FEED_TIMEOUT)
                if resp.status_code != 200:
                    continue

                posts = resp.json()
                if not isinstance(posts, list):
                    continue

                for post in posts:
                    title = _strip_html(post.get('title', {}).get('rendered', ''))
                    link = post.get('link', '')
                    excerpt = _strip_html(post.get('excerpt', {}).get('rendered', ''))

                    match_type = _keyword_match_type(kw_patterns, title, excerpt)
                    if not match_type:
                        continue

                    entry_domain = extract_domain(link) if link else domain

                    hits.append({
                        'title': title,
                        'link': link,
                        'snippet': excerpt[:300],
                        'domain': entry_domain or domain,
                        'source': info.get('name', domain),
                        'feed_country': info.get('country', ''),
                        'feed_description': info.get('description', ''),
                        'feed_media_name': info.get('name', domain),
                        '_keyword_match': match_type,
                    })

            except Exception:
                with failed_lock:
                    failed_count += 1

        return hits

    # ── Run concurrently ──────────────────────────────────────────────────

    # Set socket timeout for feedparser (it doesn't have its own timeout param)
    import socket
    old_timeout = socket.getdefaulttimeout()
    socket.setdefaulttimeout(FEED_TIMEOUT)

    start = _time.time()
    all_hits = []

    try:
        with ThreadPoolExecutor(max_workers=50) as executor:
            futures = {}
            for item in rss_outlets:
                futures[executor.submit(_scan_rss, item)] = item[0]
            for item in wp_outlets:
                futures[executor.submit(_scan_wp, item)] = item[0]

            for future in as_completed(futures):
                try:
                    hits = future.result()
                    if hits:
                        all_hits.extend(hits)
                except Exception:
                    pass
    finally:
        socket.setdefaulttimeout(old_timeout)

    elapsed = _time.time() - start
    print(f"  Scanned {len(rss_outlets)} RSS feeds + {len(wp_outlets)} WordPress APIs "
          f"in {elapsed:.1f}s → found {len(all_hits)} articles"
          + (f" ({failed_count} feeds failed)" if failed_count else ""))

    return all_hits


def mine_outlet_sitemaps(artist_keywords, days=28, feed_registry_path=None, cutoff=None, end_date_dt=None):
    """
    Mine XML sitemaps of outlets that have no RSS feed or WordPress API.
    Two-phase approach for speed:
      Phase 1: Fetch all root sitemaps concurrently (107 outlets × 1 URL each)
      Phase 2: Fetch relevant sub-sitemaps concurrently (only the handful that
               came back as sitemap indexes with recent/news sub-sitemaps)
    Returns results in the same format as scan_outlet_feeds().
    """
    import time as _time
    import requests
    import xml.etree.ElementTree as ET
    from concurrent.futures import ThreadPoolExecutor, as_completed

    registry_path = feed_registry_path or str(
        Path(__file__).parent.parent / 'data' / 'feed_registry.json'
    )

    if not os.path.exists(registry_path):
        return []

    with open(registry_path) as f:
        registry = json.load(f)

    # Only target outlets with no feed
    outlets = []
    for domain, info in registry.get('outlets', {}).items():
        if info.get('feed_type') is None:
            outlets.append((domain, info))

    if not outlets:
        return []

    kw_patterns = _make_keyword_patterns(artist_keywords)
    # Build URL-slug variants: "Bad Bunny" → ["bad-bunny", "bad_bunny", "badbunny"]
    # Slug matching stays substring-based since URL slugs use delimiters as word boundaries
    slug_variants = []
    for kw in artist_keywords:
        kw_l = kw.lower()
        slug_variants.append(kw_l.replace(' ', '-'))
        slug_variants.append(kw_l.replace(' ', '_'))
        slug_variants.append(kw_l.replace(' ', ''))

    if cutoff is None:
        cutoff = datetime.now().astimezone() - timedelta(days=days)
    end_dt = end_date_dt if end_date_dt is not None else datetime.now().astimezone()

    # Month strings to look for in sitemap index URLs (e.g. "2026-02", "2026-01")
    relevant_months = set()
    d = cutoff
    while d <= end_dt:
        relevant_months.add(d.strftime('%Y-%m'))
        d += timedelta(days=28)
    relevant_months.add(end_dt.strftime('%Y-%m'))

    TIMEOUT = 4
    UA = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    WORKERS = 50
    # XML namespaces
    SM_NS = {'sm': 'http://www.sitemaps.org/schemas/sitemap/0.9'}
    NEWS_NS = {'news': 'http://www.google.com/schemas/sitemap-news/0.9'}

    # Shared session with connection pooling
    session = requests.Session()
    adapter = requests.adapters.HTTPAdapter(pool_connections=WORKERS, pool_maxsize=WORKERS)
    session.mount('https://', adapter)
    session.mount('http://', adapter)
    session.headers.update({'User-Agent': UA})

    def _fetch_xml(url):
        """Fetch and parse an XML URL. Returns ElementTree root or None."""
        try:
            resp = session.get(url, timeout=TIMEOUT, allow_redirects=True)
            if resp.status_code != 200:
                return None
            ct = resp.headers.get('Content-Type', '').lower()
            if not any(x in ct for x in ('xml', 'text/html')):
                return None
            text = resp.text[:200].strip()
            if not text.startswith('<?xml') and not text.startswith('<'):
                return None
            return ET.fromstring(resp.content)
        except Exception:
            return None

    def _urls_from_sitemap(root):
        """Extract <loc> URLs from a sitemap XML root."""
        urls = []
        for url_el in root.findall('.//sm:url', SM_NS):
            loc = url_el.find('sm:loc', SM_NS)
            if loc is not None and loc.text:
                lastmod_el = url_el.find('sm:lastmod', SM_NS)
                lastmod = lastmod_el.text if lastmod_el is not None else None
                title = None
                news_title = url_el.find('.//news:title', NEWS_NS)
                if news_title is not None:
                    title = news_title.text
                urls.append({'loc': loc.text.strip(), 'lastmod': lastmod, 'title': title})
        return urls

    # Year/month patterns to check in URLs when lastmod is missing
    now = datetime.now().astimezone()
    _current_year = str(now.year)
    _recent_year_months = set()
    _d = cutoff
    while _d <= now.astimezone():
        _recent_year_months.add(_d.strftime('%Y/%m'))
        _recent_year_months.add(_d.strftime('%Y-%m'))
        _recent_year_months.add(str(_d.year))
        _d += timedelta(days=28)
    _recent_year_months.add(now.strftime('%Y/%m'))
    _recent_year_months.add(now.strftime('%Y-%m'))
    _recent_year_months.add(_current_year)

    def _is_recent(lastmod_str, url=None):
        """Check if a lastmod date is within our search range.
        If no lastmod, check URL for recent year/month patterns.
        Skip undated entries that don't have a recent date in the URL."""
        if lastmod_str:
            try:
                date_part = lastmod_str[:10]
                dt = datetime.strptime(date_part, '%Y-%m-%d').astimezone()
                return dt >= cutoff
            except Exception:
                pass  # Fall through to URL check

        # No parseable date — check URL for recent year/month patterns
        if url:
            url_lower = url.lower()
            for pattern in _recent_year_months:
                if pattern in url_lower:
                    return True

        return False  # Skip undated content without recent URL patterns

    def _url_matches_keywords(url_entry):
        """Check if a URL TITLE contains artist keywords (word-boundary matching).
        For sitemap results, only match on title or news:title — not URL slug alone,
        as slug matches without title confirmation catch tag/category pages."""
        title = url_entry.get('title') or ''
        if title and _any_keyword_matches(kw_patterns, title):
            return True
        # Fall back to URL slug matching only if there's no title to check
        if not title:
            loc = url_entry['loc'].lower()
            path = loc.split('/', 3)[-1] if loc.count('/') >= 3 else loc
            return any(slug in path for slug in slug_variants)
        return False

    def _match_entries(entries):
        """Filter sitemap entries to recent keyword matches."""
        return [e for e in entries if _is_recent(e['lastmod'], e['loc']) and _url_matches_keywords(e)]

    def _base_url_for(info, domain):
        website = info.get('website', '') or f'https://{domain}'
        if not website.startswith('http'):
            website = f'https://{website}'
        return website.rstrip('/')

    def _is_relevant_sub(sub_url):
        """Check if a sub-sitemap URL is worth fetching (recent month or news)."""
        sub_lower = sub_url.lower()
        for month in relevant_months:
            if month in sub_lower or month.replace('-', '') in sub_lower:
                return True
        return any(x in sub_lower for x in
                   ('news', 'post-sitemap', 'article', 'noticias', 'contenido'))

    # ── Phase 1: Fetch root sitemaps concurrently ─────────────────────────
    # Build all root sitemap URLs to fetch (one per outlet, prioritized)
    # We try news-sitemap.xml first since it has titles and is usually small.
    # If that fails we fall through to sitemap.xml then sitemap_index.xml,
    # but we submit ALL of them upfront and short-circuit on first success.

    start = _time.time()

    # Build fetch tasks: (url, domain, info, priority)
    # Lower priority number = preferred. We try all 3 concurrently per outlet
    # and pick the first successful one.
    fetch_tasks = []
    for domain, info in outlets:
        base = _base_url_for(info, domain)
        fetch_tasks.append((f'{base}/news-sitemap.xml', domain, info, 0))
        fetch_tasks.append((f'{base}/sitemap.xml', domain, info, 1))
        fetch_tasks.append((f'{base}/sitemap_index.xml', domain, info, 2))

    # Fetch all root sitemaps concurrently
    # Results: domain → list of (root, priority) for successful fetches
    outlet_roots = {}  # domain → [(root, priority, url)]
    with ThreadPoolExecutor(max_workers=WORKERS) as executor:
        future_map = {}
        for url, domain, info, prio in fetch_tasks:
            f = executor.submit(_fetch_xml, url)
            future_map[f] = (url, domain, info, prio)

        for future in as_completed(future_map):
            url, domain, info, prio = future_map[future]
            try:
                root = future.result()
                if root is not None:
                    outlet_roots.setdefault(domain, []).append((root, prio, url, info))
            except Exception:
                pass

    phase1_time = _time.time() - start

    # ── Process roots: extract direct hits + collect sub-sitemap URLs ─────
    all_hits = []          # final result entries
    sub_fetch_tasks = []   # (sub_url, domain, info) for phase 2

    for domain, roots in outlet_roots.items():
        # Sort by priority — prefer news-sitemap over sitemap over index
        roots.sort(key=lambda x: x[1])
        found_direct = False

        for root, prio, url, info in roots:
            # Check if sitemap index (has <sitemap> children)
            sub_sitemaps = root.findall('.//sm:sitemap', SM_NS)
            if sub_sitemaps:
                # Sitemap index — collect relevant sub-sitemap URLs for phase 2
                for sm_el in sub_sitemaps:
                    loc_el = sm_el.find('sm:loc', SM_NS)
                    if loc_el is not None and loc_el.text:
                        sub_url = loc_el.text.strip()
                        if _is_relevant_sub(sub_url):
                            sub_fetch_tasks.append((sub_url, domain, info))
            else:
                # Regular sitemap — scan entries directly
                entries = _urls_from_sitemap(root)
                hits = _match_entries(entries)
                if hits:
                    for h in hits:
                        loc = h['loc']
                        title = h.get('title') or ''
                        if not title:
                            path = loc.rstrip('/').rsplit('/', 1)[-1]
                            title = path.replace('-', ' ').replace('_', ' ').title()
                        all_hits.append({
                            'title': title, 'link': loc, 'snippet': '',
                            'domain': extract_domain(loc) or domain,
                            'source': info.get('name', domain),
                            'feed_country': info.get('country', ''),
                            'feed_description': info.get('description', ''),
                            'feed_media_name': info.get('name', domain),
                        })
                    found_direct = True
                    break  # Got hits from this sitemap, skip lower-priority ones

        # If we already got direct hits, no need for sub-sitemaps from this domain
        if found_direct:
            sub_fetch_tasks = [t for t in sub_fetch_tasks if t[1] != domain]

    # ── Phase 2: Fetch sub-sitemaps concurrently ──────────────────────────
    if sub_fetch_tasks:
        with ThreadPoolExecutor(max_workers=WORKERS) as executor:
            future_map = {}
            for sub_url, domain, info in sub_fetch_tasks:
                f = executor.submit(_fetch_xml, sub_url)
                future_map[f] = (domain, info)

            for future in as_completed(future_map):
                domain, info = future_map[future]
                try:
                    sub_root = future.result()
                    if sub_root is None:
                        continue
                    entries = _urls_from_sitemap(sub_root)
                    hits = _match_entries(entries)
                    for h in hits:
                        loc = h['loc']
                        title = h.get('title') or ''
                        if not title:
                            path = loc.rstrip('/').rsplit('/', 1)[-1]
                            title = path.replace('-', ' ').replace('_', ' ').title()
                        all_hits.append({
                            'title': title, 'link': loc, 'snippet': '',
                            'domain': extract_domain(loc) or domain,
                            'source': info.get('name', domain),
                            'feed_country': info.get('country', ''),
                            'feed_description': info.get('description', ''),
                            'feed_media_name': info.get('name', domain),
                        })
                except Exception:
                    pass

    elapsed = _time.time() - start
    print(f"  Mined {len(outlets)} outlet sitemaps in {elapsed:.1f}s → found {len(all_hits)} articles"
          f" (phase1: {phase1_time:.1f}s, phase2: {elapsed - phase1_time:.1f}s)")

    return all_hits


def parse_search_terms(raw_input):
    """Split free-text input into individual artist keywords.

    Multi-word names are preserved as full phrases — "Bad Bunny" stays as one keyword,
    NOT split into ["Bad", "Bunny"]. Only explicit separators trigger a split:
      'PNAU, Meduza'       → ['PNAU', 'Meduza']
      'PNAU ft. Meduza'    → ['PNAU', 'Meduza']
      'Bad Bunny & J Balvin'→ ['Bad Bunny', 'J Balvin']
      'Bad Bunny'           → ['Bad Bunny']  (no split — no separator found)

    The matching step (_make_keyword_patterns) then creates \\bBad Bunny\\b which
    requires the full phrase, so "Bunny" alone won't match.
    """
    # Split on common separators (comma, ampersand, slash, ft., feat., etc.)
    terms = re.split(r'[,&/]\s*|\s+(?:ft\.?|feat\.?|featuring|x|w/)\s+', raw_input, flags=re.IGNORECASE)
    terms = [t.strip() for t in terms if t.strip()]
    # If no explicit separator was found, treat the entire input as one term
    return terms if terms else [raw_input.strip()]


_AMBIGUOUS_ARTIST_TERMS = {
    'air', 'future', 'her', 'metric', 'motel', 'sol', 'calle', 'low',
}


def _is_ambiguous_artist_query(keywords: list[str]) -> bool:
    """Detect artist queries that need stricter music-context shaping."""
    if not keywords:
        return False

    normalized = [
        re.sub(r'[^a-z0-9]', '', _normalize_for_matching((kw or '').lower()))
        for kw in keywords
    ]
    normalized = [kw for kw in normalized if kw]
    if not normalized:
        return False

    if len(normalized) == 1:
        kw = normalized[0]
        if kw in _AMBIGUOUS_ARTIST_TERMS:
            return True
        if len(kw) <= 4:
            return True

    return False


def _build_enriched_queries(keywords, release_schedule_url=None):
    """
    Build enriched search queries using release schedule context.
    Returns a dict with query sets for different search sources:
    {
      'google_news': [(query, gl, hl), ...],   # 5 region-specific queries
      'web_news':  [query, ...],
      'web_search':   [query, ...],
      'tavily_news': query_str,
      'tavily_web':  query_str,
      'ddg':         [query, ...],
    }
    Falls back to basic artist-name queries if no release data found.
    """
    from shared.database import load_release_schedule

    artist_base = ' OR '.join(f'"{kw}"' for kw in keywords)

    # Try to load release schedule and find recent releases for this artist
    releases = []
    try:
        schedule_source = release_schedule_url or RELEASE_SCHEDULE_URL
        all_releases = load_release_schedule(schedule_source)

        # Match releases to this artist (word-boundary matching)
        release_patterns = _make_keyword_patterns(keywords)
        cutoff_days = 60  # Look back 60 days for release context
        now = datetime.now()

        for rel in all_releases:
            if not _any_keyword_matches(release_patterns, rel['artist']):
                continue

            # Parse release date (format: "Jan 5", "Feb 14" — no year, assume current year)
            if rel['date']:
                try:
                    rel_date = datetime.strptime(rel['date'] + f' {now.year}', '%b %d %Y')
                    # Handle year boundary (e.g. Dec releases when we're in Jan)
                    if rel_date > now + timedelta(days=30):
                        rel_date = rel_date.replace(year=now.year - 1)
                    days_ago = (now - rel_date).days
                    if days_ago <= cutoff_days:
                        releases.append({
                            'title': rel['title'],
                            'format': rel.get('format', 'Single'),
                            'date': rel['date'],
                            'days_ago': days_ago,
                        })
                except ValueError:
                    pass
    except Exception:
        pass  # No release schedule available — fall back to basic queries

    # Sort by most recent first
    releases.sort(key=lambda r: r['days_ago'])

    # ── Build query variants ──────────────────────────────────────────────

    regions = [
        ('MX', 'es-419'),  # Mexico
        ('AR', 'es-419'),  # Argentina
        ('BR', 'pt-BR'),   # Brazil
        ('CL', 'es-419'),  # Chile
        ('CO', 'es-419'),  # Colombia
        ('PE', 'es-419'),  # Peru
        ('EC', 'es-419'),  # Ecuador
        ('UY', 'es-419'),  # Uruguay
        ('VE', 'es-419'),  # Venezuela
    ]

    # Split regions: core (high-priority, get 2 queries) vs extended (get 1 query)
    core_regions = regions[:5]      # MX, AR, BR, CL, CO
    extended_regions = regions[5:]  # PE, EC, UY, VE

    is_ambiguous = _is_ambiguous_artist_query(keywords)

    if not releases:
        # No release context — keep the free Google pass bounded.
        # Core regions get 2 shapes; extended regions get 1.
        q_base = artist_base
        q_music = f'{artist_base} música OR musica'
        q_music_strict_es = f'{artist_base} música OR musica OR banda OR cantante OR disco OR álbum OR album OR sencillo OR canción OR cancion'
        q_music_strict_pt = f'{artist_base} música OR banda OR cantor OR cantora OR disco OR álbum OR album OR single OR canção OR cancao'
        q_press_es = f'{artist_base} estreno OR lanzamiento OR reseña OR entrevista OR concierto'
        q_press_pt = f'{artist_base} lançamento OR resenha OR entrevista OR show'

        google_queries = []
        for gl, hl in core_regions:
            if is_ambiguous:
                google_queries.append((q_music_strict_pt if hl == 'pt-BR' else q_music_strict_es, gl, hl))
                google_queries.append((q_press_pt if hl == 'pt-BR' else q_press_es, gl, hl))
            else:
                google_queries.append((q_base, gl, hl))
                google_queries.append((q_music if hl != 'pt-BR' else f'{artist_base} música', gl, hl))
        for gl, hl in extended_regions:
            if is_ambiguous:
                google_queries.append((q_press_pt if hl == 'pt-BR' else q_press_es, gl, hl))
            else:
                google_queries.append((q_base, gl, hl))

        web_news = [f'"{kw}"' for kw in keywords]
        if is_ambiguous:
            web_search = [f'"{kw}" música banda entrevista lanzamiento' for kw in keywords]
            tavily_news = q_press_es
            tavily_web = q_music_strict_es
            ddg = [f'{kw} música banda entrevista' for kw in keywords]
        else:
            web_search = [f'"{kw}" música' for kw in keywords]
            tavily_news = artist_base
            tavily_web = f'{artist_base} música'
            ddg = [f'{kw} música' for kw in keywords]

        return {
            'google_news': google_queries,
            'web_news': web_news,
            'web_search': web_search,
            'tavily_news': tavily_news,
            'tavily_web': tavily_web,
            'ddg': ddg,
            'releases': [],
        }

    # We have release context — build enriched queries
    latest = releases[0]
    release_title = latest['title']
    release_format = (latest['format'] or 'Single').strip()

    # Release-type keywords (Spanish / Portuguese)
    format_kw_es = {
        'Single': 'nuevo sencillo',
        'Album': 'nuevo álbum',
        'EP': 'nuevo EP',
    }.get(release_format, 'nuevo sencillo')

    format_kw_pt = {
        'Single': 'novo single',
        'Album': 'novo álbum',
        'EP': 'novo EP',
    }.get(release_format, 'novo single')

    # Query variants (most specific → broadest)
    q_release = f'{artist_base} "{release_title}"'                       # Exact release match
    q_format_es = f'{artist_base} {format_kw_es}'                        # Spanish format keyword
    q_format_pt = f'{artist_base} {format_kw_pt}'                        # Portuguese format keyword
    q_broad_es = f'{artist_base} estreno OR lanzamiento OR reseña OR entrevista'

    # Google News: keep the query matrix bounded for performance.
    # Core regions get 2 shapes; extended regions get 1.
    google_queries = []
    for gl, hl in core_regions:
        if hl == 'pt-BR':
            google_queries.append((q_format_pt, gl, hl))
            google_queries.append((f'{artist_base} lançamento OR entrevista OR resenha', gl, hl))
        else:
            google_queries.append((q_release, gl, hl))
            google_queries.append((q_broad_es, gl, hl))
    for gl, hl in extended_regions:
        if hl == 'pt-BR':
            google_queries.append((q_format_pt, gl, hl))
        else:
            google_queries.append((q_release, gl, hl))

    # SearXNG: release title + format keywords
    web_news = [f'"{kw}"' for kw in keywords]  # Keep broad for news (it's already filtered by recency)
    web_search = [
        f'{artist_base} "{release_title}"',
        f'{artist_base} {format_kw_es}',
    ]

    # Tavily: use release context
    tavily_news = q_release
    tavily_web = f'{artist_base} {format_kw_es}'

    # DDG: release title + format
    ddg = [
        f'{keywords[0]} "{release_title}"',
        f'{keywords[0]} {format_kw_es}',
    ]

    return {
        'google_news': google_queries,
        'web_news': web_news,
        'web_search': web_search,
        'tavily_news': tavily_news,
        'tavily_web': tavily_web,
        'ddg': ddg,
        'releases': releases,
    }


# Unicode quote variants to normalize for title comparison
_QUOTE_MAP = str.maketrans({
    '\u2018': "'", '\u2019': "'",  # curly single quotes
    '\u201c': '"', '\u201d': '"',  # curly double quotes
    '\u00ab': '"', '\u00bb': '"',  # guillemets
    '\u2013': '-', '\u2014': '-',  # en/em dashes
})


def _normalize_title(title: str) -> str:
    """Normalize a title for dedup: lowercase, strip, normalize quotes."""
    return title.lower().strip().translate(_QUOTE_MAP)


def _match_source_name_to_media(source_name: str, press_index):
    """Best-effort outlet lookup for Google News RSS source labels."""
    if not source_name:
        return None

    source_key = source_name.lower().strip()
    if source_key in press_index:
        return press_index[source_key]

    source_norm = re.sub(r'[^a-z0-9]', '', _normalize_for_matching(source_key))
    if not source_norm:
        return None

    unique_entries = []
    seen_names = set()
    for entry in press_index.values():
        entry_name = (entry.get('name') or '').lower().strip()
        if not entry_name or entry_name in seen_names:
            continue
        seen_names.add(entry_name)
        entry_norm = re.sub(r'[^a-z0-9]', '', _normalize_for_matching(entry_name))
        if entry_norm:
            unique_entries.append((entry_norm, entry))

    exact_matches = [entry for entry_norm, entry in unique_entries if entry_norm == source_norm]
    if len(exact_matches) == 1:
        return exact_matches[0]
    if len(exact_matches) > 1:
        return None

    relaxed_matches = []
    for entry_norm, entry in unique_entries:
        shorter = min(len(source_norm), len(entry_norm))
        longer = max(len(source_norm), len(entry_norm))
        if shorter / longer < 0.55:
            continue
        if source_norm.startswith(entry_norm) or entry_norm.startswith(source_norm):
            relaxed_matches.append(entry)

    if len(relaxed_matches) == 1:
        return relaxed_matches[0]
    return None


def _group_entries_by_outlet(entries):
    """Group results by outlet name, merging multiple URLs per outlet."""
    grouped = {}
    for entry in entries:
        name = entry['media_name']
        if name not in grouped:
            grouped[name] = {
                'media_name': name,
                'description': entry['description'],
                'urls': [],
                'in_database': entry['in_database'],
            }
        # Don't add duplicate URLs or duplicate titles within the same outlet
        existing_normalized = {_normalize_url(u['url']) for u in grouped[name]['urls']}
        existing_titles = {_normalize_title(u['title']) for u in grouped[name]['urls'] if u.get('title', '').strip()}
        entry_title = entry.get('title', '').strip()
        if (_normalize_url(entry['url']) not in existing_normalized
                and (not entry_title or _normalize_title(entry_title) not in existing_titles)):
            grouped[name]['urls'].append({
                'url': entry['url'],
                'type': entry.get('url_type', 'article'),
                'title': entry.get('title', ''),
            })
        # Prefer DB info over placeholder
        if entry['in_database'] and not grouped[name]['in_database']:
            grouped[name]['in_database'] = True
            grouped[name]['description'] = entry['description']
    return list(grouped.values())


def run_press_pickup(artist, days=28, output_path=None, press_db_path=None, start_date=None, end_date=None):
    """
    Main press pickup workflow for a single artist.
    """
    # Load press database
    db_path = press_db_path or PRESS_DB_PATH
    if os.path.exists(db_path):
        print(f"Loading press database from {db_path}...")
        press_index, press_entries = load_press_database(db_path)
        print(f"  Loaded {len(press_entries)} media outlets")
    else:
        print(f"Warning: Press database not found at {db_path}")
        press_index, press_entries = {}, []

    # Load social handle registry for social media classification
    social_handle_lookup = None  # None = registry not available
    if os.path.exists(SOCIAL_HANDLE_REGISTRY_PATH):
        try:
            with open(SOCIAL_HANDLE_REGISTRY_PATH) as f:
                _social_data = json.load(f)
            social_handle_lookup = _social_data.get('handle_to_outlet', {})
            _social_outlets = _social_data.get('stats', {}).get('with_any_social', '?')
            print(f"  Loaded social handle registry ({_social_outlets} outlets with handles)")
        except Exception:
            social_handle_lookup = None
    if social_handle_lookup is None:
        print("  Social handle registry not found — run discover_social_handles.py for better social media classification.")

    keywords = parse_search_terms(artist)

    # Compute cutoff datetime — either from explicit dates or days preset
    cutoff = None
    end_date_dt = None
    if start_date and end_date:
        from datetime import timezone
        cutoff = datetime.fromisoformat(start_date).replace(tzinfo=timezone.utc)
        end_date_dt = datetime.fromisoformat(end_date).replace(hour=23, minute=59, second=59, tzinfo=timezone.utc)
        days = (end_date_dt - cutoff).days or 1
        print(f"\nSearching press for: {artist} ({start_date} to {end_date})")
    else:
        cutoff = datetime.now().astimezone() - timedelta(days=days)
        print(f"\nSearching press for: {artist} (last {days} days)")

    # Build enriched queries from release schedule context
    queries = _build_enriched_queries(keywords)
    if queries['releases']:
        latest = queries['releases'][0]
        print(f"  Release context: \"{latest['title']}\" ({latest['format']}, {latest['date']}, "
              f"{latest['days_ago']}d ago)")
    else:
        print(f"  No recent releases found — using basic queries")

    all_results = []
    seen_urls = set()

    # Source tracking for breakdown summary
    source_counts = {
        'feeds': 0, 'sitemaps': 0, 'adapter': 0, 'google_news': 0,
        'searxng': 0, 'serper': 0, 'tavily': 0, 'ddg': 0,
    }

    # 0) Feed scan — RSS feeds + WordPress APIs from known outlets (instant, free)
    #    Uses raw keywords only — feeds are already targeted to the right outlets
    print(f"\n  Scanning outlet feeds...")
    feed_results = scan_outlet_feeds(keywords, days=days, cutoff=cutoff)
    for r in feed_results:
        if _normalize_url(r['link']) not in seen_urls:
            seen_urls.add(_normalize_url(r['link']))
            r['_source'] = 'feeds'
            all_results.append(r)
            source_counts['feeds'] += 1
    if feed_results:
        print(f"  Found {len(feed_results)} results from outlet feeds")

    # 0b) Sitemap mining — scan outlets with no RSS/WP for URL matches
    print(f"  Mining outlet sitemaps...")
    sitemap_results = mine_outlet_sitemaps(keywords, days=days, cutoff=cutoff, end_date_dt=end_date_dt)
    for r in sitemap_results:
        if _normalize_url(r['link']) not in seen_urls:
            seen_urls.add(_normalize_url(r['link']))
            r['_source'] = 'sitemaps'
            all_results.append(r)
            source_counts['sitemaps'] += 1

    # 0c) Outlet adapters — direct retrieval for priority LATAM outlets without strong feeds
    print(f"  Scanning outlet adapters...")
    adapter_results = scan_outlet_adapters(keywords, days=days, cutoff=cutoff)
    for r in adapter_results:
        norm_url = _normalize_url(r['link'])
        if norm_url not in seen_urls:
            seen_urls.add(norm_url)
            r['_source'] = 'adapter'
            all_results.append(r)
            source_counts['adapter'] += 1
    if adapter_results:
        print(f"  Found {len(adapter_results)} results from outlet adapters")

    # 1) Google News RSS — free, unlimited, enriched queries per region
    #    Fetch all regions (raw links) first, deduplicate, then decode in a single pool.
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from googlenewsdecoder import new_decoderv1
    gn_queries = queries['google_news']
    gn_regions = sorted(set(gl for _, gl, _ in gn_queries))
    print(f"  Google News [{'/'.join(gn_regions)}]: {len(gn_queries)} queries across {len(gn_regions)} regions...")

    def _run_gn_fetch(args):
        query, gl, hl = args
        return google_news_rss(query, gl=gl, hl=hl, days=days, cutoff=cutoff, decode=False)

    raw_items = []
    if gn_queries:
        with ThreadPoolExecutor(max_workers=min(len(gn_queries), 20)) as executor:
            for results in executor.map(_run_gn_fetch, gn_queries):
                raw_items.extend(results)

    # Deduplicate unique Google News redirect URLs before expensive decoding
    unique_google_links = {}  # link -> item template
    for item in raw_items:
        glink = item['link']
        if glink not in unique_google_links:
            unique_google_links[glink] = item

    def _safe_decode(glink):
        # 1. Try googlenewsdecoder (can hang, but pool is now bounded)
        try:
            decoded = new_decoderv1(glink)
            if decoded.get('status'):
                return decoded['decoded_url']
        except Exception:
            pass
        # 2. Fallback: follow redirect chain manually with strict timeout
        try:
            import requests as _reqs
            resp = _reqs.head(glink, allow_redirects=True, timeout=5, headers={'User-Agent': 'Mozilla/5.0'})
            if resp.status_code == 200 and 'news.google.com' not in resp.url:
                return resp.url
        except Exception:
            pass
        return glink

    if unique_google_links:
        print(f"  Decoding {len(unique_google_links)} unique Google News URLs (from {len(raw_items)} total regional hits)...")

        decoded_map = {}
        # Single bounded pool for all decodes
        with ThreadPoolExecutor(max_workers=50) as executor:
            fmap = {executor.submit(_safe_decode, gl): gl for gl in unique_google_links}
            try:
                for f in as_completed(fmap, timeout=40):
                    glink = fmap[f]
                    try:
                        decoded_map[glink] = f.result(timeout=1)
                    except Exception:
                        decoded_map[glink] = glink
            except TimeoutError:
                pass  # Use raw links for remaining

        # Build final results using the decoded URLs
        for glink, item in unique_google_links.items():
            final_link = decoded_map.get(glink, glink)
            norm = _normalize_url(final_link)
            if norm not in seen_urls:
                seen_urls.add(norm)
                item['link'] = final_link
                # Source name fallback logic (matching news.google.com results to DB)
                # is handled in Phase 2 of google_news_rss or downstream.
                # Here we just ensure we preserve the item.
                item['_source'] = 'google_news'
                all_results.append(item)
                source_counts['google_news'] += 1

    print(f"  Found {len(all_results)} results from feeds + Google News")

    # 2) SearXNG metasearch — self-hosted, aggregates Google/Bing/DuckDuckGo (free, no API key)
    searxng_available = False
    try:
        import requests as _req_check
        _req_check.get(os.environ.get('SEARXNG_URL', 'http://localhost:8888'), timeout=3)
        searxng_available = True
    except Exception:
        print("  SearXNG not available — skipping web search source")

    if searxng_available:
        # SearXNG News — enriched queries
        for query in queries['web_news']:
            print(f"  Web Search News: {query[:80]}")
            results = _search_searxng(query, num_results=20, categories='news')
            for r in results:
                link = r.get('link') or ''
                domain = r.get('domain') or extract_domain(link) or ''
                if _is_skipped_domain(domain) or _is_non_press_url(link):
                    continue
                norm = _normalize_url(link)
                if norm not in seen_urls:
                    seen_urls.add(norm)
                    r['_source'] = 'searxng'
                    all_results.append(r)
                    source_counts['searxng'] += 1

        # SearXNG Web — enriched queries with release context
        for query in queries['web_search']:
            print(f"  Web Search: {query[:80]}")
            results = _search_searxng(query, num_results=20, categories='general')
            for r in results:
                link = r.get('link') or ''
                domain = r.get('domain') or extract_domain(link) or ''
                if _is_skipped_domain(domain) or _is_non_press_url(link):
                    continue
                norm = _normalize_url(link)
                if norm not in seen_urls:
                    seen_urls.add(norm)
                    r['_source'] = 'searxng'
                    all_results.append(r)
                    source_counts['searxng'] += 1

    # 3) Serper — targeted site: queries against high-priority outlets with no results yet
    #    Instead of broad queries (old approach), we identify which known outlets from the
    #    feed registry haven't returned results via feeds or Google News, then use Serper's
    #    site: operator to search them directly. ~3 credits per artist (25 domains per query).
    serper_key = os.environ.get('SERPER_API_KEY')
    if serper_key:
        import requests as _requests

        if days <= 1:
            tbs = 'qdr:d'
        elif days <= 7:
            tbs = 'qdr:w'
        elif days <= 30:
            tbs = 'qdr:m'
        else:
            tbs = 'qdr:y'

        # Load feed registry to get the outlet list with metadata
        _registry_path = str(Path(__file__).parent.parent / 'data' / 'feed_registry.json')
        _registry_outlets = {}
        if os.path.exists(_registry_path):
            with open(_registry_path) as _f:
                _registry_outlets = json.load(_f).get('outlets', {})

        # Identify which outlet domains already have results
        domains_with_results = set()
        for r in all_results:
            d = r.get('domain', '')
            if d:
                domains_with_results.add(d)

        # Collect high-priority outlets with no results: have a real country
        # and were not already covered by feed scan or Google News
        _SKIP_TERRITORIES = {'LATAM', 'PENDING', 'CANCELLED', ''}
        no_result_outlets = []
        for domain, info in _registry_outlets.items():
            country = (info.get('country') or '').upper().strip()
            # Skip outlets with no real country — check each part of multi-value territories
            country_parts = {p.strip() for p in country.split(',')}
            if not country_parts - _SKIP_TERRITORIES:
                continue
            if domain not in domains_with_results:
                no_result_outlets.append((domain, info))

        if no_result_outlets:
            query_parts = [f'"{kw}"' for kw in keywords]
            artist_query = ' OR '.join(query_parts)

            # Group unmatched outlets by country for locale-aware Serper queries
            _COUNTRY_LOCALE = {
                'BRAZIL': ('br', 'pt'),
                'ARGENTINA': ('ar', 'es'),
                'MEXICO': ('mx', 'es'),
                'CHILE': ('cl', 'es'),
                'COLOMBIA': ('co', 'es'),
                'PERU': ('pe', 'es'),
                'ECUADOR': ('ec', 'es'),
                'URUGUAY': ('uy', 'es'),
                'VENEZUELA': ('ve', 'es'),
                'SPAIN': ('es', 'es'),
            }

            # Sort outlets: group by country for locale-aware batching
            def _outlet_locale(item):
                domain, info = item
                country = (info.get('country') or '').upper().strip().split(',')[0].strip()
                return _COUNTRY_LOCALE.get(country, ('mx', 'es'))

            no_result_outlets.sort(key=lambda x: _outlet_locale(x)[0])

            # Batch domains into groups of 15 for site: queries
            # (25 was too large — Serper returns 400; 15 is the tested sweet spot)
            BATCH_SIZE = 15
            serper_credits_used = 0
            MAX_CREDITS = 3

            for batch_start in range(0, len(no_result_outlets), BATCH_SIZE):
                if serper_credits_used >= MAX_CREDITS:
                    break

                batch = no_result_outlets[batch_start:batch_start + BATCH_SIZE]
                site_parts = ' OR '.join(f'site:{d}' for d, _ in batch)
                query = f'{artist_query} ({site_parts})'

                # Use locale of the majority country in this batch
                locale_counts = {}
                for d, info in batch:
                    loc = _outlet_locale((d, info))
                    locale_counts[loc] = locale_counts.get(loc, 0) + 1
                batch_locale = max(locale_counts, key=locale_counts.get)

                print(f"  Serper targeted [{len(batch)} outlets, gl={batch_locale[0]}]: {artist_query} + {len(batch)} site: filters")
                try:
                    # Note: tbs (time filter) is incompatible with multi-site OR
                    # queries on Serper — returns 400. We omit it here; recency
                    # is still handled by the article processing / date checks.
                    payload = {
                        'q': query,
                        'gl': batch_locale[0],
                        'hl': batch_locale[1],
                        'num': 20,
                    }
                    resp = _requests.post(
                        'https://google.serper.dev/search',
                        headers={'X-API-KEY': serper_key, 'Content-Type': 'application/json'},
                        json=payload,
                    )
                    resp.raise_for_status()
                    serper_credits_used += 1

                    added = 0
                    for item in resp.json().get('organic', []):
                        link = item.get('link', '')
                        domain = extract_domain(link) or ''
                        if _is_skipped_domain(domain):
                            continue
                        if _normalize_url(link) not in seen_urls:
                            seen_urls.add(_normalize_url(link))

                            # Look up registry metadata for this domain
                            reg_info = _registry_outlets.get(domain)
                            result_entry = {
                                'title': item.get('title', ''),
                                'link': link,
                                'snippet': item.get('snippet', ''),
                                'domain': domain,
                            }

                            # Pre-classify with registry metadata if we have it
                            if reg_info:
                                result_entry['feed_media_name'] = reg_info.get('name', domain)
                                result_entry['feed_country'] = reg_info.get('country', '')
                                result_entry['feed_description'] = reg_info.get('description', '')
                                result_entry['source'] = reg_info.get('name', domain)

                            result_entry['_source'] = 'serper'
                            all_results.append(result_entry)
                            added += 1
                            source_counts['serper'] += 1

                    print(f"    → {added} new results")

                except Exception as e:
                    print(f"    Serper targeted failed: {e}")
                    serper_credits_used += 1  # Count failed attempts too
        else:
            print(f"  Serper: all registry outlets already covered — skipping")

    # ── Old Serper approach (broad queries) ────────────────────────────────
    # Replaced by targeted site: queries above. The old approach used 3 broad
    # queries per artist (1 news + 2 organic with music keywords), which burned
    # the same ~3 credits but returned mostly results we'd already found via
    # Google News RSS, plus lots of non-LATAM noise to filter out.
    #
    # if serper_key:
    #     import requests as _requests
    #     query_parts = [f'"{kw}"' for kw in keywords]
    #     base = ' OR '.join(query_parts)
    #     serper_calls = [
    #         ('news',   base),
    #         ('search', f'{base} música'),
    #         ('search', f'{base} lanzamiento OR álbum OR disco OR entrevista'),
    #     ]
    #     for search_type, query in serper_calls:
    #         label = 'News' if search_type == 'news' else 'Web'
    #         print(f"  Serper {label}: {query}")
    #         try:
    #             payload = {'q': query, 'gl': 'mx', 'hl': 'es', 'num': 20}
    #             if search_type == 'search':
    #                 payload['tbs'] = tbs
    #             resp = _requests.post(
    #                 f'https://google.serper.dev/{search_type}',
    #                 headers={'X-API-KEY': serper_key, 'Content-Type': 'application/json'},
    #                 json=payload,
    #             )
    #             resp.raise_for_status()
    #             result_key = 'news' if search_type == 'news' else 'organic'
    #             cutoff = datetime.now().astimezone() - timedelta(days=days)
    #             for item in resp.json().get(result_key, []):
    #                 if search_type == 'news':
    #                     date_str = item.get('date', '')
    #                     if not _serper_date_within(date_str, cutoff):
    #                         continue
    #                 link = item.get('link', '')
    #                 domain = extract_domain(link) or ''
    #                 if any(skip in domain for skip in SKIP_DOMAINS):
    #                     continue
    #                 if link not in seen_urls:
    #                     seen_urls.add(link)
    #                     all_results.append({
    #                         'title': item.get('title', ''),
    #                         'link': link,
    #                         'snippet': item.get('snippet', ''),
    #                         'domain': domain,
    #                     })
    #         except Exception as e:
    #             print(f"    Serper failed: {e}")

    # 4) Tavily — 2 credits per artist: 1 news + 1 general (free tier: 1000/month recurring)
    tavily_key = os.environ.get('TAVILY_API_KEY')
    if tavily_key:
        import requests as _requests

        if days <= 1:
            time_range = 'day'
        elif days <= 7:
            time_range = 'week'
        elif days <= 30:
            time_range = 'month'
        else:
            time_range = 'year'

        tavily_calls = [
            # News search — enriched with release context
            {'query': queries['tavily_news'], 'topic': 'news', 'time_range': time_range,
             'max_results': 20, 'search_depth': 'basic'},
            # General search with country=mexico — enriched with release keywords
            {'query': queries['tavily_web'], 'topic': 'general', 'country': 'mexico',
             'time_range': time_range, 'max_results': 20, 'search_depth': 'basic'},
        ]

        for payload in tavily_calls:
            label = 'News' if payload['topic'] == 'news' else 'Web'
            print(f"  Tavily {label}: {payload['query'][:80]}")
            try:
                resp = _requests.post(
                    'https://api.tavily.com/search',
                    headers={'Authorization': f'Bearer {tavily_key}', 'Content-Type': 'application/json'},
                    json=payload,
                    timeout=15,
                )
                resp.raise_for_status()
                tavily_results = resp.json().get('results', [])
                added = 0
                for item in tavily_results:
                    link = item.get('url', '')
                    domain = extract_domain(link) or ''
                    if _is_skipped_domain(domain):
                        continue
                    if _normalize_url(link) not in seen_urls:
                        seen_urls.add(_normalize_url(link))
                        all_results.append({
                            'title': item.get('title', ''),
                            'link': link,
                            'snippet': item.get('content', ''),
                            'domain': domain,
                            '_source': 'tavily',
                        })
                        added += 1
                        source_counts['tavily'] += 1
                print(f"    → {added} new results")
            except Exception as e:
                print(f"    Tavily failed: {e}")

    # 5) DuckDuckGo News — free, unlimited, unofficial (could break if DDG changes backend)
    #    Only news search is used; text search in this library version returns garbage results.
    try:
        from duckduckgo_search import DDGS

        for query in queries['ddg']:
            print(f"  DuckDuckGo News: {query}")
            try:
                ddg = DDGS()
                raw = ddg.news(query, max_results=20)
                added = 0
                for item in raw:
                    link = item.get('url', '')
                    domain = extract_domain(link) or ''
                    if _is_skipped_domain(domain):
                        continue
                    if _normalize_url(link) not in seen_urls:
                        seen_urls.add(_normalize_url(link))
                        all_results.append({
                            'title': item.get('title', ''),
                            'link': link,
                            'snippet': item.get('body', ''),
                            'domain': domain,
                            '_source': 'ddg',
                        })
                        added += 1
                        source_counts['ddg'] += 1
                print(f"    → {added} new results")
            except Exception as e:
                print(f"    DuckDuckGo failed: {e}")
    except ImportError:
        print("  DuckDuckGo: skipped (duckduckgo_search not installed)")

    # 6) Free targeted retries — use Google News RSS with site: queries for
    #    high-priority outlets that returned 0 results from all previous sources.
    #    This is completely free (Google News RSS is unlimited).
    _registry_path = str(Path(__file__).parent.parent / 'data' / 'feed_registry.json')
    _retry_registry = {}
    if os.path.exists(_registry_path):
        try:
            with open(_registry_path) as _f:
                _retry_registry = json.load(_f).get('outlets', {})
        except Exception:
            pass

    if _retry_registry:
        domains_with_results = {r.get('domain', '') for r in all_results if r.get('domain')}
        _SKIP_TERRITORIES = {'LATAM', 'PENDING', 'CANCELLED', ''}
        unmatched = []
        for domain, info in _retry_registry.items():
            country = (info.get('country') or '').upper().strip()
            country_parts = {p.strip() for p in country.split(',')}
            if not country_parts - _SKIP_TERRITORIES:
                continue
            if domain not in domains_with_results:
                unmatched.append((domain, info))

        if unmatched:
            artist_query = ' OR '.join(f'"{kw}"' for kw in keywords)
            SITE_BATCH = 8
            MAX_RETRY_OUTLETS = 40  # Cap to prevent excessive requests
            # Run targeted site: queries across 2 key regions (MX covers most of LATAM, BR for Portuguese)
            retry_regions = [('MX', 'es-419'), ('BR', 'pt-BR')]
            retry_added = 0
            capped_unmatched = unmatched[:MAX_RETRY_OUTLETS]

            for batch_start in range(0, len(capped_unmatched), SITE_BATCH):
                batch = capped_unmatched[batch_start:batch_start + SITE_BATCH]
                site_filter = ' OR '.join(f'site:{d}' for d, _ in batch)
                site_query = f'{artist_query} ({site_filter})'

                def _run_retry_fetch(args):
                    q, gl, hl = args
                    return google_news_rss(q, gl=gl, hl=hl, days=days, cutoff=cutoff, decode=False)

                retry_raw = []
                with ThreadPoolExecutor(max_workers=len(retry_regions)) as ex:
                    for results in ex.map(_run_retry_fetch, [(site_query, gl, hl) for gl, hl in retry_regions]):
                        retry_raw.extend(results)

                # Deduplicate unique redirect links from retry pass
                unique_retry = {}
                for item in retry_raw:
                    glink = item['link']
                    if glink not in unique_retry:
                        unique_retry[glink] = item

                if unique_retry:
                    # Reuse simple decode for retry pass
                    with ThreadPoolExecutor(max_workers=20) as ex:
                        fmap = {ex.submit(_safe_decode, gl): gl for gl in unique_retry}
                        for f in as_completed(fmap, timeout=20):
                            glink = fmap[f]
                            try:
                                final_link = f.result()
                                norm = _normalize_url(final_link)
                                if norm not in seen_urls:
                                    seen_urls.add(norm)
                                    item = unique_retry[glink]
                                    item['link'] = final_link
                                    item['_source'] = 'google_news'
                                    all_results.append(item)
                                    source_counts['google_news'] += 1
                                    retry_added += 1
                            except Exception:
                                pass

            if retry_added:
                print(f"  Free targeted retries: {retry_added} new results from {len(capped_unmatched)} outlets (of {len(unmatched)} unmatched)")

    # Final global safety filter pass before AI processing
    final_safety = []
    leaked_count = 0
    for r in all_results:
        link = r.get('link') or ''
        domain = r.get('domain') or extract_domain(link) or ''
        if _is_skipped_domain(domain) or _is_non_press_url(link):
            leaked_count += 1
            continue
        final_safety.append(r)
    all_results = final_safety
    if leaked_count > 0:
        print(f"  Safety filter: removed {leaked_count} non-press or skipped results")

    print(f"\nFound {len(all_results)} total unique results")

    # AI relevance filter — remove false positives using Groq (free, optional)
    all_results = _groq_filter_relevance(
        all_results, artist, keywords,
        releases=queries.get('releases'),
        log_fn=print,
    )

    # Match against press database and group by country
    country_results = {}
    skipped = 0
    kw_patterns = _make_keyword_patterns(keywords)
    new_outlets_to_enrich = []  # Collect new outlets for batch AI enrichment
    social_stats = {'known_outlet': 0, 'artist_account': 0, 'unknown': 0}

    for result in all_results:
        domain = result['domain']

        # ── Filter tag/category/index pages (all sources) ──
        if _is_non_press_url(result.get('link', '')):
            skipped += 1
            continue

        # ── Feed-sourced results: pre-classified, skip DB matching logic ──
        if result.get('feed_media_name'):
            source = result.get('_source', '')
            match_type = result.get('_keyword_match', 'title')
            if source in ('feeds', 'adapter'):
                if match_type not in ('title', 'snippet'):
                    skipped += 1
                    continue
            elif source == 'sitemaps':
                if not _any_keyword_matches(kw_patterns, result['title']):
                    skipped += 1
                    continue

            media_name = result['feed_media_name']
            description = result.get('feed_description', '')
            country = normalize_country(result.get('feed_country', 'LATAM').upper())
            # Handle multi-country territories — use first real country
            if ',' in country:
                parts = [p.strip() for p in country.split(',')]
                country = next(
                    (p for p in parts if p not in ('PENDING', 'CANCELLED', '')),
                    'LATAM'
                )
            if not country or country in ('PENDING', 'CANCELLED'):
                country = 'LATAM'

            country = normalize_country(country)
            if country not in country_results:
                country_results[country] = []
            country_results[country].append({
                'media_name': media_name,
                'description': description,
                'url': result['link'],
                'title': result['title'],
                'snippet': result['snippet'],
                'in_database': True,  # Feed outlets are from our registry
                'url_type': 'article',
            })
            continue

        # ── Standard results: keyword check + DB matching ──
        # The Groq filter upstream already split title-confirmed vs snippet-only
        # and rejected irrelevant snippet-only articles. Here we just verify at
        # least one keyword appears somewhere (title or snippet) as a safety net.
        title_text = result['title']
        snippet_text = result.get('snippet') or ''
        if not (_any_keyword_matches(kw_patterns, title_text) or _any_keyword_matches(kw_patterns, snippet_text)):
            skipped += 1
            continue

        country = detect_country_from_url(result['link'])
        source_name_match = result.get('_source_name_match')
        media_entry = match_url_to_media(result['link'], press_index)
        if not media_entry and source_name_match:
            media_entry = _match_source_name_to_media(source_name_match, press_index)

        effective_domain = domain
        if media_entry and source_name_match and domain == 'news.google.com':
            effective_domain = extract_domain(media_entry.get('website', '')) or domain

        if media_entry:
            # Known outlet from DB — always include
            description = media_entry['description']
            media_name = media_entry['name']
            if media_entry['territory'] and media_entry['territory'] not in ('PENDING', 'CANCELLED'):
                territory = media_entry['territory']
                if ',' not in territory:
                    country = normalize_country(territory.upper())
            if not country:
                country = 'LATAM'
            # Multi-regional outlet check: article is on generic .com but
            # DB entry has a country-specific domain (.com.mx, .com.ar, etc.)
            # → likely the US/international edition, not the LATAM one.
            # Exclude entirely unless the URL contains LATAM language indicators.
            db_domain = extract_domain(media_entry.get('website', ''))
            if (country != 'LATAM'
                    and _is_generic_com_domain(effective_domain)
                    and db_domain and not _is_generic_com_domain(db_domain)
                    and not _has_latam_url_indicators(result['link'])):
                print(f"  Skipped: {media_name} (US edition) — no LATAM indicators in URL")
                continue
        elif domain in SOCIAL_MEDIA_DOMAINS:
            # ── Social media classification ──
            platform_label = SOCIAL_MEDIA_DOMAINS[domain]
            handle_info = _extract_social_handle(result['link'])

            if social_handle_lookup is not None and handle_info:
                platform_key, handle = handle_info

                # Check if it's the artist's own account
                if _handle_matches_artist_keywords(handle, keywords):
                    social_stats['artist_account'] += 1
                    skipped += 1
                    continue

                # Check if handle belongs to a known outlet
                platform_handles = social_handle_lookup.get(platform_key, {})
                outlet_info = platform_handles.get(handle)
                if outlet_info:
                    media_name = outlet_info['name']
                    # Look up full description from press DB
                    outlet_entry = press_index.get(outlet_info['name'].lower().strip())
                    if not outlet_entry:
                        outlet_entry = press_index.get(outlet_info.get('domain', ''))
                    description = (outlet_entry or {}).get('description', '') or result['title']
                    country = normalize_country(outlet_info.get('country', 'LATAM').upper())
                    if not country or country in ('PENDING', 'CANCELLED'):
                        country = 'LATAM'
                    social_stats['known_outlet'] += 1
                else:
                    # Unknown handle — exclude
                    social_stats['unknown'] += 1
                    skipped += 1
                    continue
            elif social_handle_lookup is not None:
                # URL pattern not parseable (e.g. instagram.com/p/...) — exclude
                social_stats['unknown'] += 1
                skipped += 1
                continue
            else:
                # No registry — fall back to old behavior
                media_name = platform_label
                description = result['title']
                if not country:
                    country = 'LATAM'
        elif is_latam_domain(domain):
            # New outlet with LATAM TLD — placeholder description, enrich later
            media_name = domain.split('.')[0].title() if domain else 'Unknown'
            description = 'Online media outlet covering entertainment and music news.'
            if not country:
                country = 'LATAM'
            print(f"  New outlet (not in DB): {media_name} ({domain})")
            new_outlets_to_enrich.append({
                'media_name': media_name,
                'url': result['link'],
                'snippet': result.get('snippet', ''),
            })
        else:
            # Non-LATAM domain, not in DB — skip (likely US/UK/Spain)
            skipped += 1
            continue

        country = normalize_country(country)

        if country not in country_results:
            country_results[country] = []

        country_results[country].append({
            'media_name': media_name,
            'description': description,
            'url': result['link'],
            'title': result['title'],
            'snippet': result['snippet'],
            'in_database': media_entry is not None,
            'url_type': _SOCIAL_DOMAIN_TYPE.get(domain, 'article'),
        })

    # Batch-enrich new outlet descriptions with Groq AI (free)
    if new_outlets_to_enrich:
        # Deduplicate by media_name before enriching
        seen_names = set()
        unique_to_enrich = []
        for o in new_outlets_to_enrich:
            if o['media_name'] not in seen_names:
                seen_names.add(o['media_name'])
                unique_to_enrich.append(o)

        enriched = _groq_enrich_descriptions(unique_to_enrich, log_fn=print)
        if enriched:
            # Apply enriched descriptions back to country_results
            for country_entries in country_results.values():
                for entry in country_entries:
                    if not entry['in_database'] and entry['media_name'] in enriched:
                        entry['description'] = enriched[entry['media_name']]

    # Group results by outlet within each country (merges multiple URLs per outlet)
    for country in country_results:
        country_results[country] = _group_entries_by_outlet(country_results[country])

    if skipped:
        print(f"  Filtered out {skipped} non-LATAM or irrelevant results")

    # Social media classification summary
    if any(social_stats.values()):
        parts = []
        if social_stats['known_outlet']:
            parts.append(f"{social_stats['known_outlet']} from known outlets")
        if social_stats['artist_account']:
            parts.append(f"{social_stats['artist_account']} artist accounts (excluded)")
        if social_stats['unknown']:
            parts.append(f"{social_stats['unknown']} unknown (excluded)")
        print(f"  Social media: {', '.join(parts)}")

    # Source breakdown summary
    source_labels = {
        'feeds': 'Outlet Feeds (RSS/WP)',
        'sitemaps': 'Sitemap Mining',
        'adapter': 'Outlet Adapters',
        'google_news': 'Google News',
        'searxng': 'Web Search',
        'serper': 'Serper (targeted)',
        'tavily': 'Tavily',
        'ddg': 'DuckDuckGo',
    }
    active_sources = [(source_labels[k], v) for k, v in source_counts.items() if v > 0]
    if active_sources:
        parts = [f"{count} from {label}" for label, count in active_sources]
        source_summary = "Sources: " + ", ".join(parts)
        print(f"\n  {source_summary}")
    else:
        source_summary = ""

    # Format output
    output_lines = [f"Press Pickup — {artist}\n"]

    for country in sorted(country_results.keys()):
        entries = country_results[country]
        output_lines.append(f"\n{country}")

        for entry in entries:
            output_lines.append(f"{entry['media_name']}: {entry['description']}")
            urls = entry['urls']
            if len(urls) == 1:
                u = urls[0]
                label = _URL_TYPE_LABELS.get(u['type'])
                title = u.get('title', '').strip()
                if label:
                    line = f"{label}: "
                else:
                    line = ""
                if title:
                    line += f"{title} — {u['url']}"
                else:
                    line += u['url']
                output_lines.append(line)
            else:
                for u in urls:
                    label = _URL_TYPE_LABELS.get(u['type'])
                    title = u.get('title', '').strip()
                    if label:
                        prefix = f"• {label}: "
                    else:
                        prefix = "• "
                    if title:
                        output_lines.append(f"{prefix}{title} — {u['url']}")
                    else:
                        output_lines.append(f"{prefix}{u['url']}")
            output_lines.append("")

    if source_summary:
        output_lines.append("")
        output_lines.append(source_summary)

    output_text = '\n'.join(output_lines)
    
    # Save or print
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output_text)
        print(f"\nReport saved: {output_path}")

        # Generate formatted .docx alongside the .txt
        try:
            docx_path = output_path.replace('.txt', '.docx') if output_path.endswith('.txt') else output_path + '.docx'
            _generate_press_docx(artist, country_results, docx_path)
            print(f"Word report saved: {docx_path}")
        except Exception as e:
            print(f"Warning: Could not generate .docx: {e}")
    else:
        print("\n" + "=" * 60)
        print(output_text)

    return country_results


def _generate_press_docx(artist, country_results, docx_path):
    """Generate a formatted .docx Press Pickup report."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Override docDefaults to remove built-in spacing (200 twips after, 1.15 line)
    from docx.oxml.ns import qn
    styles_el = doc.styles.element
    for child in styles_el:
        if child.tag.endswith('docDefaults'):
            for ppr_default in child.iter(qn('w:pPrDefault')):
                for ppr in ppr_default.iter(qn('w:pPr')):
                    for spacing in ppr.iter(qn('w:spacing')):
                        spacing.set(qn('w:after'), '0')
                        spacing.set(qn('w:before'), '0')
                        spacing.set(qn('w:line'), '240')
                        spacing.set(qn('w:lineRule'), 'auto')

    # Set default font and zero paragraph spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Title: "Press pickup" in bold red
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Press pickup')
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0xC4, 0x30, 0x30)
    title_run.font.size = Pt(13)
    title_para.paragraph_format.space_after = Pt(0)

    for country in sorted(country_results.keys()):
        entries = country_results[country]

        # Country header: underlined, not bold
        country_para = doc.add_paragraph()
        country_run = country_para.add_run(country)
        country_run.underline = True
        country_run.font.size = Pt(11)
        country_para.paragraph_format.space_before = Pt(0)
        country_para.paragraph_format.space_after = Pt(0)

        for entry in entries:
            # Media entry: bold name + normal description
            media_para = doc.add_paragraph()
            name_run = media_para.add_run(f"{entry['media_name']}: ")
            name_run.bold = True
            name_run.font.size = Pt(11)
            desc_text = entry['description']
            desc_run = media_para.add_run(desc_text)
            desc_run.font.size = Pt(11)
            media_para.paragraph_format.space_before = Pt(0)
            media_para.paragraph_format.space_after = Pt(0)

            # URLs — display as raw clickable links
            urls = entry.get('urls', [])
            if len(urls) == 1:
                u = urls[0]
                url_para = doc.add_paragraph()
                label = _URL_TYPE_LABELS.get(u['type'])
                if label:
                    prefix_run = url_para.add_run(f"{label}: ")
                    prefix_run.font.size = Pt(11)
                _add_hyperlink(url_para, u['url'], u['url'])
                url_para.paragraph_format.space_before = Pt(0)
                url_para.paragraph_format.space_after = Pt(0)
            else:
                for u in urls:
                    url_para = doc.add_paragraph()
                    label = _URL_TYPE_LABELS.get(u['type'])
                    if label:
                        prefix_run = url_para.add_run(f"• {label}: ")
                    else:
                        prefix_run = url_para.add_run("• ")
                    prefix_run.font.size = Pt(11)
                    _add_hyperlink(url_para, u['url'], u['url'])
                    url_para.paragraph_format.space_before = Pt(0)
                    url_para.paragraph_format.space_after = Pt(0)

            # Blank line separator after each entry
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(0)
            sep.paragraph_format.space_after = Pt(0)
            sep.paragraph_format.line_spacing = 1.0

    doc.save(docx_path)


def _add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph in a .docx document."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2E74B5')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')  # 11pt = 22 half-points
    rPr.append(sz)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def main():
    parser = argparse.ArgumentParser(
        description='Press Pickup Tool — Find and format press coverage for artists',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python press_pickup.py --artist "Djo" --days 28
  python press_pickup.py --artist "Djo" --days 7 --output djo_press.txt
  python press_pickup.py --all --days 7

Environment Variables:
  GROQ_API_KEY       Groq API key (optional, free — for generating missing descriptions)
  TAVILY_API_KEY     Tavily API key (optional, free 1000 credits/month — additional search source)
  PRESS_DB_PATH      Path to press description database CSV
  RELEASE_SCHEDULE_URL  Published Google Sheets URL for release schedule
        """
    )
    
    parser.add_argument('--artist', help='Artist name to search for')
    parser.add_argument('--all', action='store_true', help='Process all artists from release schedule')
    parser.add_argument('--days', type=int, default=28, help='Number of days to search back (default: 28)')
    parser.add_argument('--output', help='Output file path (default: print to stdout)')
    parser.add_argument('--press-db', help='Path to press description database CSV')
    parser.add_argument('--release-schedule', help='Path or URL to release schedule CSV')
    
    args = parser.parse_args()
    
    if not args.artist and not args.all:
        parser.error('Either --artist or --all is required')
    
    if args.all:
        # Load release schedule and process all artists
        schedule_source = args.release_schedule or RELEASE_SCHEDULE_URL
        print(f"Loading release schedule from {schedule_source}...")
        
        sys.path.insert(0, str(Path(__file__).parent.parent))
        from shared.database import load_release_schedule
        
        releases = load_release_schedule(schedule_source)
        artists = sorted(set(r['artist'] for r in releases))
        print(f"Found {len(artists)} unique artists")
        
        for artist in artists:
            print(f"\n{'=' * 60}")
            output_path = None
            if args.output:
                safe_name = re.sub(r'[^a-zA-Z0-9]', '_', artist).lower()
                output_path = f"{args.output}/{safe_name}_press.txt"
            
            run_press_pickup(artist, args.days, output_path, args.press_db)
    else:
        run_press_pickup(args.artist, args.days, args.output, args.press_db)


if __name__ == '__main__':
    main()

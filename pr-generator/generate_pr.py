#!/usr/bin/env python3
"""
Press Release Translator
========================
Translates press releases from any language into Spanish and/or Brazilian
Portuguese for LATAM distribution.

Accepts either pasted text or an uploaded .docx file.

Two translation modes:
  - Google Translate (default, free, no API key)
  - Gemini AI (optional, higher quality for music-industry tone, requires GEMINI_API_KEY)

When a .docx file is uploaded, the output .docx preserves ALL formatting from
the original document (fonts, bold, italic, alignment, sizes, paragraph spacing).

Called from the web UI via /api/pr/translate.
"""

import os
import re
import shutil
import sys
from pathlib import Path

ROOT_DIR = Path(__file__).parent.parent.resolve()
sys.path.insert(0, str(ROOT_DIR))


# ═════════════════════════════════════════════════════════════════════════
#  DOCX TEXT EXTRACTION
# ═════════════════════════════════════════════════════════════════════════

def extract_docx_text(docx_path):
    """Extract text from a .docx file, preserving paragraph breaks."""
    from docx import Document
    doc = Document(docx_path)
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
    return '\n\n'.join(paragraphs)


# ═════════════════════════════════════════════════════════════════════════
#  LANGUAGE DETECTION
# ═════════════════════════════════════════════════════════════════════════

# ISO codes for translation engines
LANG_CODES = {
    'English': 'en',
    'French': 'fr',
    'Spanish': 'es',
    'Portuguese': 'pt',
}

def _detect_language(text):
    """Detect likely source language from common word patterns. Returns label string."""
    sample = text[:3000].lower()

    # Word-boundary markers for more accurate detection
    en_markers = [' the ', ' and ', ' is ', ' for ', ' with ', ' this ', ' that ', ' from ', ' are ', ' will ', ' their ', ' has ', ' been ']
    fr_markers = [' le ', ' la ', ' les ', ' des ', ' est ', ' une ', ' dans ', ' pour ', ' avec ', ' qui ', ' sur ', ' sont ', ' du ']
    es_markers = [' el ', ' los ', ' las ', ' del ', ' con ', ' por ', ' una ', ' está ', ' para ', ' sus ', ' más ', ' será ']
    pt_markers = [' os ', ' das ', ' dos ', ' uma ', ' com ', ' para ', ' está ', ' não ', ' seu ', ' sua ', ' mais ', ' será ']

    scores = {
        'English': sum(sample.count(m) for m in en_markers),
        'French': sum(sample.count(m) for m in fr_markers),
        'Spanish': sum(sample.count(m) for m in es_markers),
        'Portuguese': sum(sample.count(m) for m in pt_markers),
    }

    return max(scores, key=scores.get)


# ═════════════════════════════════════════════════════════════════════════
#  GOOGLE TRANSLATE — PLAIN TEXT (for pasted text input)
# ═════════════════════════════════════════════════════════════════════════

def _translate_google(text, target_lang, source_lang=None, log_fn=print):
    """Translate plain text using Google Translate (free, no API key)."""
    from deep_translator import GoogleTranslator

    src_code = LANG_CODES.get(source_lang, 'auto')
    tgt_code = 'es' if target_lang == 'es' else 'pt'

    paragraphs = text.split('\n')
    chunks = []
    current_chunk = []
    current_len = 0

    for para in paragraphs:
        para_len = len(para) + 1
        if current_len + para_len > 4500 and current_chunk:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [para]
            current_len = para_len
        else:
            current_chunk.append(para)
            current_len += para_len

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    translator = GoogleTranslator(source=src_code, target=tgt_code)

    translated_chunks = []
    for i, chunk in enumerate(chunks):
        if not chunk.strip():
            translated_chunks.append(chunk)
            continue
        try:
            result = translator.translate(chunk)
            translated_chunks.append(result)
            if len(chunks) > 1:
                log_fn(f'  Chunk {i + 1}/{len(chunks)} translated.')
        except Exception as e:
            log_fn(f'  Warning: chunk {i + 1} failed ({e}), keeping original.')
            translated_chunks.append(chunk)

    return '\n'.join(translated_chunks)


# ═════════════════════════════════════════════════════════════════════════
#  GEMINI AI — PLAIN TEXT (for pasted text input)
# ═════════════════════════════════════════════════════════════════════════

def _translate_gemini(text, target_lang, source_lang=None, notes='', log_fn=print):
    """Translate plain text using Google Gemini Flash."""
    api_key = os.environ.get('GEMINI_API_KEY', '')
    if not api_key:
        return None

    import requests as req

    lang_name = 'Latin American Spanish' if target_lang == 'es' else 'Brazilian Portuguese'
    src_label = source_lang or 'the source language'

    system_prompt = (
        f"You are an expert music industry translator specializing in press releases for Latin American markets. "
        f"Translate the following press release from {src_label} into {lang_name}.\n\n"
        f"CRITICAL RULES:\n"
        f"- Preserve ALL formatting: paragraph breaks, section headers, bullet points, tracklists, tour dates, links.\n"
        f"- Keep artist names, song/album titles, venue names, label names, and proper nouns EXACTLY as-is (do NOT translate them).\n"
        f"- Keep all URLs and email addresses exactly as-is.\n"
        f"- Translate quotes naturally — they should sound like the person said them in {lang_name}.\n"
        f"- Use professional music press release tone appropriate for {lang_name}-speaking media outlets.\n"
        f"- Localize date formats (e.g., 'June 5' → '5 de junio' for Spanish, '5 de junho' for Portuguese).\n"
        f"- If the PR contains sections in multiple languages, translate ALL sections into {lang_name}.\n"
        f"- Output ONLY the translated press release. No preamble, no notes, no commentary.\n"
    )

    if notes:
        system_prompt += f"\nAdditional instructions from the user: {notes}\n"

    try:
        url = f'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}'

        resp = req.post(
            url,
            headers={'content-type': 'application/json'},
            json={
                'system_instruction': {
                    'parts': [{'text': system_prompt}],
                },
                'contents': [{
                    'parts': [{'text': text}],
                }],
                'generationConfig': {
                    'maxOutputTokens': 8192,
                    'temperature': 0.3,
                },
            },
            timeout=120,
        )

        if resp.status_code == 200:
            data = resp.json()
            candidates = data.get('candidates', [])
            if candidates:
                parts = candidates[0].get('content', {}).get('parts', [])
                if parts:
                    return parts[0].get('text', '').strip()
            log_fn('  Gemini returned empty response.')
            return None
        else:
            log_fn(f'  Gemini API error: {resp.status_code} — {resp.text[:200]}')
            return None
    except Exception as e:
        log_fn(f'  Gemini translation error: {e}')
        return None


# ═════════════════════════════════════════════════════════════════════════
#  DOCX FORMAT-PRESERVING TRANSLATION
# ═════════════════════════════════════════════════════════════════════════

def _collect_run_info(para):
    """Collect text and formatting metadata for each run in a paragraph."""
    runs_info = []
    for run in para.runs:
        runs_info.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'font_size': run.font.size,
        })
    return runs_info


def _all_same_format(runs_info):
    """Check if all runs share the same bold/italic/size formatting."""
    if len(runs_info) <= 1:
        return True
    ref = runs_info[0]
    return all(
        ri['bold'] == ref['bold'] and
        ri['italic'] == ref['italic'] and
        ri['font_size'] == ref['font_size']
        for ri in runs_info[1:]
    )


def _snap_to_word_boundary(text, start, target):
    """Snap a character position to the nearest word boundary (space)."""
    if target >= len(text) or target <= start:
        return target
    for d in range(1, min(20, len(text) - start)):
        fwd = target + d
        bwd = target - d
        if fwd < len(text) and text[fwd] == ' ':
            return fwd + 1
        if bwd > start and text[bwd] == ' ':
            return bwd + 1
    return target


# ── Apply translation to paragraph runs ───────────────────────────────

def _apply_translation_proportional(para, translated_text):
    """
    Apply translated text using proportional redistribution.

    Groups consecutive same-format runs, distributes text proportionally.
    """
    if not translated_text:
        return

    runs = list(para.runs)
    if not runs:
        return

    runs_info = _collect_run_info(para)

    if len(runs) == 1:
        runs[0].text = translated_text
        return

    if _all_same_format(runs_info):
        runs[0].text = translated_text
        for r in runs[1:]:
            r.text = ''
        return

    # Group consecutive same-format runs
    groups = []  # [(start_idx, end_idx, combined_text), ...]
    i = 0
    while i < len(runs_info):
        ri = runs_info[i]
        fmt_key = (ri['bold'], ri['italic'], ri['font_size'])
        combined = ri['text']
        j = i + 1
        while j < len(runs_info):
            rj = runs_info[j]
            if (rj['bold'], rj['italic'], rj['font_size']) == fmt_key:
                combined += rj['text']
                j += 1
            else:
                break
        groups.append((i, j, combined))
        i = j

    if len(groups) == 1:
        runs[0].text = translated_text
        for r in runs[1:]:
            r.text = ''
        return

    total_orig = sum(len(g[2]) for g in groups)
    if total_orig == 0:
        runs[0].text = translated_text
        for r in runs[1:]:
            r.text = ''
        return

    total_trans = len(translated_text)
    char_pos = 0

    for gi, (start_idx, end_idx, orig_text) in enumerate(groups):
        if gi == len(groups) - 1:
            segment = translated_text[char_pos:]
        else:
            proportion = len(orig_text) / total_orig
            target_end = int(char_pos + proportion * total_trans)
            target_end = _snap_to_word_boundary(translated_text, char_pos, target_end)
            segment = translated_text[char_pos:target_end]
            char_pos = target_end

        runs[start_idx].text = segment
        for k in range(start_idx + 1, end_idx):
            if k < len(runs):
                runs[k].text = ''




# ── Batch translation functions ───────────────────────────────────────

def _translate_batch_google(texts, target_lang, source_lang, log_fn):
    """Translate a list of paragraph texts using Google Translate."""
    from deep_translator import GoogleTranslator

    src_code = LANG_CODES.get(source_lang, 'auto')
    tgt_code = 'es' if target_lang == 'es' else 'pt'
    translator = GoogleTranslator(source=src_code, target=tgt_code)

    translated = []
    for i, text in enumerate(texts):
        if not text.strip():
            translated.append(text)
            continue
        try:
            if len(text) > 4500:
                parts = []
                sentences = re.split(r'(?<=[.!?])\s+', text)
                current = ''
                for sent in sentences:
                    if len(current) + len(sent) + 1 > 4500 and current:
                        r = translator.translate(current)
                        parts.append(r if r else current)
                        current = sent
                    else:
                        current += (' ' if current else '') + sent
                if current:
                    r = translator.translate(current)
                    parts.append(r if r else current)
                translated.append(' '.join(parts))
            else:
                r = translator.translate(text)
                translated.append(r if r else text)
        except Exception as e:
            log_fn(f'  Warning: paragraph {i+1} failed ({e}), keeping original.')
            translated.append(text)

        if len(texts) > 10 and (i + 1) % 10 == 0:
            log_fn(f'  {i+1}/{len(texts)} paragraphs translated...')

    return translated


def _gemini_api_call(system_prompt, user_text, max_tokens=16384, log_fn=print):
    """Make a Gemini API call. Returns response text or None."""
    api_key = os.environ.get('GEMINI_API_KEY', '')
    if not api_key:
        return None

    import requests as req

    try:
        url = f'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}'
        resp = req.post(
            url,
            headers={'content-type': 'application/json'},
            json={
                'system_instruction': {'parts': [{'text': system_prompt}]},
                'contents': [{'parts': [{'text': user_text}]}],
                'generationConfig': {
                    'maxOutputTokens': max_tokens,
                    'temperature': 0.3,
                },
            },
            timeout=180,
        )

        if resp.status_code == 200:
            data = resp.json()
            candidates = data.get('candidates', [])
            if candidates:
                parts = candidates[0].get('content', {}).get('parts', [])
                if parts:
                    return parts[0].get('text', '').strip()
            log_fn('  Gemini returned empty response.')
        else:
            log_fn(f'  Gemini API error: {resp.status_code} \u2014 {resp.text[:200]}')
    except Exception as e:
        log_fn(f'  Gemini error: {e}')

    return None


def _translate_batch_gemini(texts, target_lang, source_lang, notes, log_fn):
    """
    Translate paragraphs using Gemini, chunked for reliability.

    Sends paragraphs in chunks of ~15 to avoid marker loss with large documents.
    Falls back to Google Translate for any chunk that fails.
    """
    lang_name = 'Latin American Spanish' if target_lang == 'es' else 'Brazilian Portuguese'
    src_label = source_lang or 'the source language'

    PARA_MARKER = '|||PARA|||'
    CHUNK_SIZE = 15  # paragraphs per API call

    all_translated = []

    for chunk_start in range(0, len(texts), CHUNK_SIZE):
        chunk = texts[chunk_start:chunk_start + CHUNK_SIZE]
        combined = f'\n{PARA_MARKER}\n'.join(chunk)

        system_prompt = (
            f"You are an expert music industry translator for Latin American press releases.\n"
            f"Translate from {src_label} into {lang_name}.\n\n"
            f"CRITICAL RULES:\n"
            f"- The input has {len(chunk)} text sections separated by {PARA_MARKER} markers.\n"
            f"- Your output MUST have exactly {len(chunk) - 1} {PARA_MARKER} markers (preserving all {len(chunk)} sections).\n"
            f"- Keep artist names, song/album titles, venue names, label names, and proper nouns EXACTLY as-is (do NOT translate them).\n"
            f"- Keep all URLs and email addresses exactly as-is.\n"
            f"- Localize date formats (June 5 \u2192 5 de junio / 5 de junho).\n"
            f"- Professional music press release tone.\n"
            f"- Output ONLY the translated text with {PARA_MARKER} markers. No preamble, no commentary.\n"
        )
        if notes:
            system_prompt += f"\nAdditional instructions: {notes}\n"

        result_text = _gemini_api_call(system_prompt, combined, max_tokens=8192, log_fn=log_fn)

        if result_text:
            parts = result_text.split(PARA_MARKER)
            parts = [p.strip() for p in parts]

            if len(parts) == len(chunk):
                all_translated.extend(parts)
                if len(texts) > CHUNK_SIZE:
                    log_fn(f'  Chunk {chunk_start // CHUNK_SIZE + 1}/{(len(texts) + CHUNK_SIZE - 1) // CHUNK_SIZE} translated ({len(chunk)} paragraphs).')
                continue
            else:
                log_fn(f'  Gemini marker mismatch ({len(parts)} vs {len(chunk)}), falling back for this chunk...')

        # Fallback for this chunk
        fallback = _translate_batch_google(chunk, target_lang, source_lang, log_fn)
        all_translated.extend(fallback)

    log_fn(f'  Gemini translated {len(all_translated)} paragraphs.')
    return all_translated


# ── Main .docx translation orchestrator ───────────────────────────────

def _translate_docx(docx_path, target_lang, source_lang, use_ai, notes, log_fn, output_dir):
    """
    Translate a .docx preserving all formatting (alignment, bold, italic,
    font sizes, paragraph spacing, images, headers, footers).

    For Gemini AI: uses format markers so bold/italic boundaries are
    preserved accurately through translation.

    For Google Translate: uses proportional redistribution of translated
    text across original run boundaries.

    Returns path to the translated .docx file.
    """
    from docx import Document

    lang_suffix = 'ES' if target_lang == 'es' else 'PT-BR'
    base = Path(docx_path)
    output_path = Path(output_dir) / f'{base.stem}_{lang_suffix}{base.suffix}'
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Copy original .docx (preserves all styles, images, headers, footers)
    shutil.copy2(docx_path, str(output_path))
    doc = Document(str(output_path))

    # Collect all paragraphs with translatable run text (body + tables)
    para_entries = []  # [(paragraph_obj, plain_text), ...]

    def _collect(para):
        run_text = ''.join(r.text for r in para.runs)
        if run_text.strip():
            para_entries.append((para, run_text))

    for para in doc.paragraphs:
        _collect(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _collect(para)

    log_fn(f'  {len(para_entries)} paragraphs to translate.')

    # Get translations for all paragraphs
    orig_texts = [text for _, text in para_entries]

    if use_ai:
        translations = _translate_batch_gemini(orig_texts, target_lang, source_lang, notes, log_fn)
    else:
        translations = _translate_batch_google(orig_texts, target_lang, source_lang, log_fn)

    # Apply translations using proportional redistribution (preserves run formatting)
    for (para, _), trans_text in zip(para_entries, translations):
        _apply_translation_proportional(para, trans_text)

    doc.save(str(output_path))
    log_fn(f'  Saved: {output_path.name}')
    return str(output_path)


# ═════════════════════════════════════════════════════════════════════════
#  MAIN FUNCTION
# ═════════════════════════════════════════════════════════════════════════

def translate_pr(
    text='',
    docx_path='',
    target_es=True,
    target_pt=True,
    use_ai=False,
    notes='',
    output_dir='',
    log_fn=None,
):
    """
    Translate a press release into Spanish and/or Portuguese.

    Provide either `text` (pasted PR content) or `docx_path` (uploaded .docx).

    When a .docx is provided, the output preserves all original formatting.

    Returns dict with:
      - source_text: the original text used
      - source_lang: detected source language
      - es_text / pt_text: plain text translations
      - es_docx_path / pt_docx_path: paths to translated .docx files (when input was .docx)
      - engine: 'google' or 'gemini'
    """
    if log_fn is None:
        log_fn = print

    result = {
        'source_text': '',
        'source_lang': '',
        'es_text': '',
        'pt_text': '',
        'es_docx_path': '',
        'pt_docx_path': '',
        'engine': '',
    }

    # ── Extract source text ───────────────────────────────────
    if docx_path:
        log_fn('Extracting text from .docx...')
        try:
            text = extract_docx_text(docx_path)
            log_fn(f'  Extracted {len(text):,} characters from document.')
        except Exception as e:
            raise ValueError(f'Failed to read .docx file: {e}')

    if not text or not text.strip():
        raise ValueError('No text provided. Paste the PR content or upload a .docx file.')

    result['source_text'] = text.strip()

    # ── Detect source language ────────────────────────────────
    source_lang = _detect_language(text)
    result['source_lang'] = source_lang
    log_fn(f'Detected source language: {source_lang}')

    # ── Choose engine ─────────────────────────────────────────
    if use_ai:
        api_key = os.environ.get('GEMINI_API_KEY', '')
        if not api_key:
            log_fn('GEMINI_API_KEY not set — falling back to Google Translate.')
            use_ai = False
        else:
            log_fn('Using Gemini Flash AI for translation.')

    if not use_ai:
        log_fn('Using Google Translate (free).')

    result['engine'] = 'gemini' if use_ai else 'google'

    # ── Translate to Spanish ──────────────────────────────────
    if target_es:
        if source_lang == 'Spanish':
            log_fn('Source is already Spanish — skipping ES translation.')
            result['es_text'] = text.strip()
            if docx_path:
                result['es_docx_path'] = docx_path
        else:
            log_fn('Translating to Spanish...')
            if docx_path:
                es_path = _translate_docx(docx_path, 'es', source_lang, use_ai, notes, log_fn, output_dir)
                result['es_docx_path'] = es_path
                result['es_text'] = extract_docx_text(es_path)
                log_fn(f'  Spanish translation complete ({len(result["es_text"]):,} chars).')
            else:
                if use_ai:
                    es = _translate_gemini(text, 'es', source_lang, notes, log_fn)
                else:
                    es = _translate_google(text, 'es', source_lang, log_fn)
                if es:
                    result['es_text'] = es
                    log_fn(f'  Spanish translation complete ({len(es):,} chars).')
                else:
                    raise ValueError('Spanish translation failed.')

    # ── Translate to Portuguese ───────────────────────────────
    if target_pt:
        if source_lang == 'Portuguese':
            log_fn('Source is already Portuguese — skipping PT translation.')
            result['pt_text'] = text.strip()
            if docx_path:
                result['pt_docx_path'] = docx_path
        else:
            log_fn('Translating to Portuguese...')
            if docx_path:
                pt_path = _translate_docx(docx_path, 'pt', source_lang, use_ai, notes, log_fn, output_dir)
                result['pt_docx_path'] = pt_path
                result['pt_text'] = extract_docx_text(pt_path)
                log_fn(f'  Portuguese translation complete ({len(result["pt_text"]):,} chars).')
            else:
                if use_ai:
                    pt = _translate_gemini(text, 'pt', source_lang, notes, log_fn)
                else:
                    pt = _translate_google(text, 'pt', source_lang, log_fn)
                if pt:
                    result['pt_text'] = pt
                    log_fn(f'  Portuguese translation complete ({len(pt):,} chars).')
                else:
                    raise ValueError('Portuguese translation failed.')

    log_fn('Translation complete!')
    return result


if __name__ == '__main__':
    sample = """BYE PARULA ANNOUNCES "SOMETHING OUT OF NOTHING"
A BRAND-NEW ALBUM OUT JUNE 5 VIA SECRET CITY RECORDS

LISTEN TO THE FIRST SINGLE "KISSBURN" AVAILABLE NOW

Bye Parula is thrilled to announce their brand-new album "Something Out Of Nothing" set to release on June 5 via Secret City Records in both LP and CD formats."""

    result = translate_pr(text=sample, target_es=True, target_pt=True)
    print(f'=== DETECTED: {result["source_lang"]} (engine: {result["engine"]}) ===')
    print()
    print('=== SPANISH ===')
    print(result['es_text'])
    print()
    print('=== PORTUGUESE ===')
    print(result['pt_text'])

#!/usr/bin/env python3
"""
Report Compiler
===============
Compiles a full client-facing report by orchestrating Radio, DSP, and Press
tools and combining their outputs into a single branded .docx document.

Called from the web UI via /api/report/compile.
"""

import importlib.util
import io
import os
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

ROOT_DIR = Path(__file__).parent.parent.resolve()
sys.path.insert(0, str(ROOT_DIR))

REPORT_DIR = ROOT_DIR / 'reports'
REPORT_DIR.mkdir(exist_ok=True)

RELEASE_SCHEDULE_URL = os.environ.get(
    'RELEASE_SCHEDULE_URL',
    'https://docs.google.com/spreadsheets/d/e/2PACX-1vTSd9mhkVibb7AwXtsZjRgBfuRT9sLY_qhhu-rB_P35CX2vFk_fZw_f31AJyW84KrCzWLLMUcTzzgqU/pub?gid=497066221&single=true&output=csv'
)

PLAY_KEY_MAP = {
    '7d':   'weekly_plays',
    '28d':  'plays_28d',
    '1y':   'yearly_plays',
}

SORT_COL_MAP = {
    '7d':  'weeklyPlaysCount',
    '28d': 'monthlyPlaysCount',
    '1y':  'yearlyPlaysCount',
}


def compile_report(
    artist,
    days=28,
    press_days=None,
    press_start_date=None,
    press_end_date=None,
    radio_region='latam',
    radio_time_range='28d',
    radio_start_date=None,
    radio_end_date=None,
    efforts_text='',
    output_path=None,
    log_fn=None,
    include_radio=True,
    include_dsp=True,
    include_press=True,
):
    """
    Compile a full client report combining radio, DSP, and press data.

    Returns dict with:
      - radio_data: list of airplay entries (or None)
      - press_data: country_results dict (or None)
      - dsp_data: results dict (or None)
      - releases: list of artist releases from schedule
      - output_path: path to generated .docx
    """
    if log_fn is None:
        log_fn = print

    safe_artist = artist.lower().replace(' ', '_')
    if not output_path:
        output_path = str(REPORT_DIR / f'{safe_artist}_full_report.docx')

    result = {
        'radio_data': None,
        'press_data': None,
        'dsp_data': None,
        'releases': [],
        'output_path': output_path,
    }

    # ─── Load release schedule ─────────────────────────────────────
    log_fn('Loading release schedule...')
    from shared.database import load_release_schedule, load_playlist_database

    try:
        all_releases = load_release_schedule(RELEASE_SCHEDULE_URL)
    except Exception as e:
        log_fn(f'  Release schedule unavailable: {e}')
        all_releases = []
    search_lower = artist.lower()
    artist_releases = [
        r for r in all_releases
        if search_lower in r['artist'].lower() or r['artist'].lower() in search_lower
    ]
    result['releases'] = artist_releases
    log_fn(f'  Found {len(artist_releases)} releases for {artist}')

    # ─── 1. Radio Report (Soundcharts) ────────────────────────────
    radio_data = None
    if include_radio:
        log_fn('\n── Radio Report ──')
        try:
            from shared.soundcharts import search_artist, fetch_airplay_data, get_token

            token = get_token()
            if not token:
                log_fn('  Soundcharts credentials not configured — skipping radio.')
            else:
                region_label = 'LATAM' if radio_region == 'latam' else 'all countries'
                log_fn(f'  Searching Soundcharts for "{artist}" ({region_label})...')
                match = search_artist(artist, token=token)
                if not match:
                    log_fn(f'  Artist "{artist}" not found on Soundcharts.')
                elif radio_time_range == 'custom' and radio_start_date and radio_end_date:
                    # Custom date range — fetch per-song via SongBroadcastTopBroadcastPlayList
                    from shared.soundcharts import fetch_song_custom_range, LATAM_CODES
                    log_fn(f'  Found: {match["name"]} (UUID: {match["uuid"]})')
                    log_fn(f'  Fetching airplay data ({radio_start_date} to {radio_end_date})...')

                    # First get songs via standard fetch to discover song UUIDs
                    airplay_preview = fetch_airplay_data(
                        match['uuid'], token,
                        sort_by='monthlyPlaysCount',
                        region=radio_region if radio_region != 'all' else None,
                        log_fn=log_fn,
                    )
                    # Collect unique song UUIDs
                    song_uuids = {}
                    for entry in (airplay_preview or []):
                        sname = entry.get('song', '')
                        suuid = entry.get('song_uuid')
                        if sname and suuid and sname not in song_uuids:
                            song_uuids[sname] = suuid

                    if not song_uuids:
                        log_fn('  No songs found for this artist.')
                    else:
                        country_filter = LATAM_CODES if radio_region == 'latam' else None
                        all_custom = []
                        for song_name, song_uuid in song_uuids.items():
                            log_fn(f'  Custom range: "{song_name}" ({radio_start_date} to {radio_end_date})...')
                            items = fetch_song_custom_range(
                                song_uuid, token, radio_start_date, radio_end_date,
                                country_codes=country_filter, log_fn=log_fn,
                            )
                            if items:
                                total_plays = sum(i['plays'] for i in items)
                                log_fn(f'    → {len(items)} stations, {total_plays} total plays')
                                for item in items:
                                    all_custom.append({
                                        'song': song_name,
                                        'station': item['station'],
                                        'plays_28d': item['plays'],
                                        'country': item['country'],
                                    })

                        if all_custom:
                            radio_data = all_custom
                            result['radio_data'] = radio_data
                            log_fn(f'  Total: {len(radio_data)} station entries')
                        else:
                            log_fn('  No airplay data found in custom range.')
                else:
                    log_fn(f'  Found: {match["name"]} (UUID: {match["uuid"]})')
                    log_fn(f'  Fetching airplay data...')

                    sort_col = SORT_COL_MAP.get(radio_time_range, 'monthlyPlaysCount')

                    airplay = fetch_airplay_data(
                        match['uuid'], token,
                        sort_by=sort_col,
                        region=radio_region if radio_region != 'all' else None,
                        log_fn=log_fn,
                    )
                    if airplay:
                        radio_data = airplay
                        result['radio_data'] = radio_data
                        log_fn(f'  Total: {len(radio_data)} station entries')
                    else:
                        log_fn('  No airplay data found.')
        except Exception as e:
            log_fn(f'  Radio fetch failed: {e}')

    # ─── 2. Press Pickup ──────────────────────────────────────────
    press_data = None
    if include_press:
        log_fn('\n── Press Pickup ──')
        try:
            from shared.capture import capture_stdout
            spec_path = ROOT_DIR / 'press-pickup' / 'press_pickup.py'
            spec = importlib.util.spec_from_file_location('press_pickup', str(spec_path))
            mod = importlib.util.module_from_spec(spec)

            with capture_stdout() as buf:
                spec.loader.exec_module(mod)
                press_output = str(REPORT_DIR / f'{safe_artist}_press.txt')
                _press_days = press_days if press_days is not None else days
                _press_kwargs = {}
                if press_start_date and press_end_date:
                    _press_kwargs['start_date'] = press_start_date
                    _press_kwargs['end_date'] = press_end_date
                press_data = mod.run_press_pickup(artist, _press_days, press_output, **_press_kwargs)
                result['press_data'] = press_data

            for line in buf.getvalue().splitlines():
                log_fn(line)

            total_press = sum(len(v) for v in press_data.values()) if press_data else 0
            log_fn(f'  Found {total_press} press results across {len(press_data) if press_data else 0} countries')
        except Exception as e:
            log_fn(f'  Press pickup failed: {e}')

    # ─── 3. DSP Pickup ────────────────────────────────────────────
    dsp_data = None
    if include_dsp and artist_releases:
        log_fn('\n── DSP Pickup ──')
        try:
            pl_path = os.environ.get(
                'PLAYLIST_DB_PATH',
                str(ROOT_DIR / 'data' / 'playlist_database.csv')
            )
            playlists = load_playlist_database(pl_path)
            log_fn(f'  Loaded {len(playlists)} playlists')

            # Clear previous proof images
            proof_dir = REPORT_DIR / 'dsp_proofs'
            if proof_dir.exists():
                shutil.rmtree(proof_dir, ignore_errors=True)

            dsp_output = str(REPORT_DIR / f'{safe_artist}_dsp.txt')

            with capture_stdout() as buf:
                spec_path = ROOT_DIR / 'dsp-pickup' / 'dsp_pickup.py'
                spec = importlib.util.spec_from_file_location('dsp_pickup_run', str(spec_path))
                mod = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(mod)
                dsp_data = mod.run_dsp_pickup(artist_releases, playlists, dsp_output)
                result['dsp_data'] = dsp_data

            for line in buf.getvalue().splitlines():
                log_fn(line)

            total_dsp = sum(
                len(matches)
                for rel in dsp_data.values()
                for matches in rel.values()
            ) if dsp_data else 0
            log_fn(f'  Found {total_dsp} playlist placements')
        except Exception as e:
            log_fn(f'  DSP pickup failed: {e}')
    elif include_dsp:
        log_fn('\n── DSP Pickup ──')
        log_fn('  No releases found for this artist — skipping DSP check.')

    # ─── Save snapshot to dashboard history ──────────────────────
    play_key = PLAY_KEY_MAP.get(radio_time_range, 'plays_28d')
    try:
        from shared.history import save_snapshot
        save_snapshot(artist, radio_data=radio_data, press_data=press_data,
                      dsp_data=dsp_data, play_key=play_key, source='report')
        log_fn('  Snapshot saved to dashboard history')
    except Exception as e:
        log_fn(f'  (Snapshot save skipped: {e})')

    # ─── Generate combined .docx ──────────────────────────────────
    log_fn('\n── Generating Full Report ──')
    _generate_full_docx(
        artist=artist,
        days=days,
        radio_time_range=radio_time_range,
        radio_start_date=radio_start_date,
        radio_end_date=radio_end_date,
        play_key=play_key,
        releases=artist_releases,
        radio_data=radio_data,
        press_data=press_data,
        dsp_data=dsp_data,
        efforts_text=efforts_text,
        output_path=output_path,
    )
    log_fn(f'  Report saved: {output_path}')

    return result


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def _generate_full_docx(
    artist, days, radio_time_range, play_key,
    releases, radio_data, press_data, dsp_data,
    efforts_text, output_path,
    radio_start_date=None, radio_end_date=None,
):
    """Generate the combined client report .docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Override docDefaults to remove built-in spacing (200 twips after, 1.15 line)
    from docx.oxml.ns import qn as _qn
    for child in doc.styles.element:
        if child.tag.endswith('docDefaults'):
            for ppr_default in child.iter(_qn('w:pPrDefault')):
                for ppr in ppr_default.iter(_qn('w:pPr')):
                    for spacing in ppr.iter(_qn('w:spacing')):
                        spacing.set(_qn('w:after'), '0')
                        spacing.set(_qn('w:before'), '0')
                        spacing.set(_qn('w:line'), '240')
                        spacing.set(_qn('w:lineRule'), 'auto')

    # Default styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    RED = RGBColor(0xC4, 0x30, 0x30)
    BLUE = RGBColor(0x00, 0x56, 0xD2)
    GRAY = RGBColor(0x66, 0x66, 0x66)

    RANGE_LABELS = {
        '7d': 'Last 7 Days',
        '28d': 'Last 28 Days',
        '1y': 'Last Year',
    }
    if radio_time_range == 'custom' and radio_start_date and radio_end_date:
        period_label = f'{radio_start_date} — {radio_end_date}'
    else:
        period_label = RANGE_LABELS.get(radio_time_range, f'Last {days} days')

    # ─── Title ────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(artist)
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(20)

    sub_para = doc.add_paragraph()
    run = sub_para.add_run(f'Report — {period_label}')
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    sub_para.paragraph_format.space_after = Pt(16)

    # ─── Release Timeline ─────────────────────────────────────
    if releases:
        _add_section_header(doc, 'Release Timeline', RED)
        for r in releases:
            p = doc.add_paragraph()
            title_run = p.add_run(r['title'])
            title_run.bold = True
            title_run.font.size = Pt(10)

            parts = []
            if r.get('date'):
                parts.append(r['date'])
            if r.get('format'):
                parts.append(r['format'])
            if r.get('label'):
                parts.append(r['label'])
            if parts:
                detail_run = p.add_run(f" — {' / '.join(parts)}")
                detail_run.font.size = Pt(10)
                detail_run.font.color.rgb = GRAY
            p.paragraph_format.space_before = Pt(2)

    # ─── Overall Efforts ──────────────────────────────────────
    if efforts_text and efforts_text.strip():
        _add_section_header(doc, 'Overall Efforts', RED)
        for line in efforts_text.strip().split('\n'):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            if line.startswith(('-', '\u2022', '*')):
                line = line.lstrip('-\u2022* ')
                run = p.add_run(f'\u2022 {line}')
            else:
                run = p.add_run(line)
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(2)

    # ─── Radio Plays ──────────────────────────────────────────
    if radio_data:
        _add_section_header(doc, 'Radio Plays', RED)

        # Group by country -> station -> songs
        country_stations = {}
        for entry in radio_data:
            country = entry.get('country', 'UNKNOWN')
            station = entry.get('station', 'Unknown')
            song = entry.get('song', '')
            plays = entry.get(play_key, 0) or entry.get('plays_28d', 0)

            if not plays or not song:
                continue

            if country not in country_stations:
                country_stations[country] = {}
            if station not in country_stations[country]:
                country_stations[country][station] = []

            existing = [s for s in country_stations[country][station] if s['song'] == song]
            if existing:
                existing[0]['plays'] = max(existing[0]['plays'], plays)
            else:
                country_stations[country][station].append({'song': song, 'plays': plays})

        for country in sorted(country_stations.keys()):
            cp = doc.add_paragraph()
            cr = cp.add_run(country)
            cr.underline = True
            cr.font.size = Pt(10)
            cp.paragraph_format.space_before = Pt(12)
            cp.paragraph_format.space_after = Pt(2)

            stations = country_stations[country]
            for station_name in sorted(stations.keys()):
                songs = stations[station_name]
                songs.sort(key=lambda s: s['plays'], reverse=True)

                sp = doc.add_paragraph()
                sr = sp.add_run(station_name)
                sr.bold = True
                sr.font.size = Pt(10)

                for song_info in songs:
                    sr = sp.add_run(f'\n     \u2022 {song_info["song"]} ({song_info["plays"]}x)')
                    sr.font.size = Pt(10)

                sp.paragraph_format.space_before = Pt(6)
                sp.paragraph_format.space_after = Pt(2)

    # ─── Streaming / Playlists Highlights ─────────────────────
    if dsp_data:
        _add_section_header(doc, 'Streaming / Playlists Highlights', RED)

        all_matches = []
        for a, releases_dict in dsp_data.items():
            for title, matches in releases_dict.items():
                for m in matches:
                    all_matches.append(m)

        if all_matches:
            platform_order = [
                'Spotify', 'Apple Music', 'Deezer',
                'Amazon Music', 'YouTube Music', 'Claro M\u00fasica',
            ]
            def sort_key(m):
                try:
                    idx = platform_order.index(m.get('platform', ''))
                except ValueError:
                    idx = 99
                return (idx, m.get('playlist_name', ''))
            all_matches.sort(key=sort_key)

            current_platform = None
            today_str = datetime.now().strftime('%b %d, %Y')

            for m in all_matches:
                platform = m.get('platform', '')

                # Platform header
                if platform != current_platform:
                    if current_platform is not None:
                        doc.add_paragraph().paragraph_format.space_before = Pt(6)
                    p = doc.add_paragraph()
                    run = p.add_run(platform)
                    run.bold = True
                    run.font.color.rgb = BLUE
                    run.font.size = Pt(12)
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(2)
                    current_platform = platform

                # Country
                country = m.get('playlist_country', '')
                if country:
                    cp = doc.add_paragraph()
                    cr = cp.add_run(country)
                    cr.underline = True
                    cr.font.size = Pt(10)
                    cp.paragraph_format.space_before = Pt(4)
                    cp.paragraph_format.space_after = Pt(1)

                # Playlist info line
                info_parts = [m.get('playlist_name', '')]
                if m.get('playlist_followers'):
                    info_parts.append(m['playlist_followers'])
                info_parts.append(today_str)

                ip = doc.add_paragraph()
                ir = ip.add_run(' - '.join(info_parts))
                ir.bold = True
                ir.font.size = Pt(10)
                ip.paragraph_format.space_before = Pt(2)
                ip.paragraph_format.space_after = Pt(4)

                # Embed proof image
                img_path = _find_proof_image(
                    m.get('playlist_track', ''),
                    m.get('playlist_name', ''),
                    output_path,
                )
                if img_path:
                    img_para = doc.add_paragraph()
                    img_para.paragraph_format.space_after = Pt(8)
                    run = img_para.add_run()
                    run.add_picture(str(img_path), width=Inches(6.2))
                else:
                    fp = doc.add_paragraph()
                    track = m.get('playlist_track', '')
                    fr = fp.add_run(
                        f'  #{m.get("position", "?")} \u2014 {track} by {m.get("playlist_artist", "")}'
                    )
                    fr.font.size = Pt(9)
                    fr.font.color.rgb = GRAY
                    fp.paragraph_format.space_after = Pt(8)

    # ─── Press pickup ─────────────────────────────────────────
    if press_data:
        _add_section_header(doc, 'Press pickup', RED)

        for country in sorted(press_data.keys()):
            entries = press_data[country]

            cp = doc.add_paragraph()
            cr = cp.add_run(country)
            cr.underline = True
            cr.font.size = Pt(11)
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)

            for entry in entries:
                mp = doc.add_paragraph()
                nr = mp.add_run(f"{entry['media_name']}: ")
                nr.bold = True
                nr.font.size = Pt(11)

                dr = mp.add_run(entry['description'])
                dr.font.size = Pt(11)
                mp.paragraph_format.space_before = Pt(0)
                mp.paragraph_format.space_after = Pt(0)

                # Entries are grouped by outlet — each has a 'urls' list
                # Display article title as clickable hyperlink
                urls = entry.get('urls', [])
                if urls:
                    if len(urls) == 1:
                        u = urls[0]
                        url_para = doc.add_paragraph()
                        display = u.get('title', '').strip() or u['url']
                        _add_hyperlink(url_para, u['url'], display)
                        url_para.paragraph_format.space_before = Pt(0)
                        url_para.paragraph_format.space_after = Pt(0)
                    else:
                        for u in urls:
                            url_para = doc.add_paragraph()
                            bullet_run = url_para.add_run('\u2022 ')
                            bullet_run.font.size = Pt(11)
                            display = u.get('title', '').strip() or u['url']
                            _add_hyperlink(url_para, u['url'], display)
                            url_para.paragraph_format.space_before = Pt(0)
                            url_para.paragraph_format.space_after = Pt(0)
                elif entry.get('url'):
                    # Fallback for ungrouped entries
                    url_para = doc.add_paragraph()
                    _add_hyperlink(url_para, entry['url'], entry.get('title', '').strip() or entry['url'])
                    url_para.paragraph_format.space_before = Pt(0)
                    url_para.paragraph_format.space_after = Pt(0)

                # Blank line separator after each entry
                sep = doc.add_paragraph()
                sep.paragraph_format.space_before = Pt(0)
                sep.paragraph_format.space_after = Pt(0)
                sep.paragraph_format.line_spacing = 1.0

    doc.save(output_path)


def _add_section_header(doc, text, color):
    """Add a styled section header."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)


def _find_proof_image(track_name, playlist_name, output_path):
    """Find the proof image file for a given track + playlist combo."""
    import unicodedata

    def _ascii_safe(s, maxlen):
        s = unicodedata.normalize('NFKD', s).encode('ascii', 'ignore').decode()
        return re.sub(r'[^\w\s-]', '', s)[:maxlen].strip().replace(' ', '_') or 'item'

    safe_track = _ascii_safe(track_name, 30)
    safe_playlist = _ascii_safe(playlist_name, 20)
    img_filename = f'proof_{safe_track}_{safe_playlist}.png'

    proof_dir = Path(output_path).parent / 'dsp_proofs'
    img_path = proof_dir / img_filename
    return img_path if img_path.exists() else None


def _add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2E74B5')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')  # 11pt = 22 half-points
    rPr.append(sz)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
